#!/usr/bin/env python3
"""Extract plain text from generated/**/*.docx into audit_text/ (one .txt per file).

Naming: audit_text/generated__<folder>__<stem>.docx.txt
Removes .txt files with no matching .docx under generated/.
"""

from __future__ import annotations

import argparse
from pathlib import Path

from docx import Document
from docx.opc.exceptions import PackageNotFoundError


def _skip_docx_path(path: Path) -> bool:
    """Microsoft Word creates ``~$name.docx`` lock files while a file is open; they are not valid OOXML packages."""
    return path.name.startswith("~$")


def extract_docx_text(docx_path: Path) -> str:
    doc = Document(str(docx_path))
    lines = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            lines.append("\t".join(cell.text for cell in row.cells))
    return "\n".join(lines) + "\n"


def main() -> None:
    p = argparse.ArgumentParser(description=__doc__)
    p.add_argument(
        "--repo-root",
        type=Path,
        default=Path(__file__).resolve().parent.parent,
        help="Repository root (default: parent of scripts/)",
    )
    args = p.parse_args()
    root: Path = args.repo_root
    gen = root / "generated"
    audit = root / "audit_text"
    audit.mkdir(exist_ok=True)

    expected: set[str] = set()
    skipped = 0
    # Include all generated .docx, including compiled books under generated/books/.
    for docx_path in sorted(gen.rglob("*.docx")):
        if _skip_docx_path(docx_path):
            skipped += 1
            continue
        folder = docx_path.parent.name
        stem = docx_path.stem
        name = f"generated__{folder}__{stem}.docx.txt"
        try:
            text = extract_docx_text(docx_path)
        except PackageNotFoundError:
            print(f"skip (not a valid docx package): {docx_path.relative_to(root)}", flush=True)
            skipped += 1
            continue
        expected.add(name)
        (audit / name).write_text(text, encoding="utf-8")

    removed = 0
    for f in audit.glob("*.txt"):
        if f.name not in expected:
            f.unlink()
            removed += 1

    msg = f"wrote {len(expected)} extracts, removed {removed} stale file(s)"
    if skipped:
        msg += f", skipped {skipped} path(s)"
    print(msg)


if __name__ == "__main__":
    main()
