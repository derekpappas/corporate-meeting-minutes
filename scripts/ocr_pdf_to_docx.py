#!/usr/bin/env python3
from __future__ import annotations

import argparse
from dataclasses import dataclass
from pathlib import Path
import shutil
import subprocess
import tempfile
from typing import Iterable


def _require_imports() -> tuple[object, object, object, object]:
    """
    Import heavy/optional deps lazily so the script can print a helpful error.

    Returns: (fitz, Vision, Quartz, Document)
    """
    try:
        import fitz  # type: ignore[import-not-found]
    except Exception as e:  # pragma: no cover
        raise RuntimeError(
            "Missing dependency: PyMuPDF.\n"
            "Install with: `poetry add pymupdf`"
        ) from e

    try:
        import Vision  # type: ignore[import-not-found]
        import Quartz  # type: ignore[import-not-found]
        import Foundation  # type: ignore[import-not-found]
    except Exception as e:  # pragma: no cover
        raise RuntimeError(
            "Missing dependency: macOS Vision bindings (pyobjc).\n"
            "Install with:\n"
            "  `poetry add pyobjc-core pyobjc-framework-Vision pyobjc-framework-Quartz pyobjc-framework-Cocoa`\n"
            "\n"
            "Note: This OCR implementation is macOS-only."
        ) from e

    try:
        from docx import Document  # type: ignore[import-not-found]
    except Exception as e:  # pragma: no cover
        raise RuntimeError(
            "Missing dependency: python-docx.\n"
            "Install with: `poetry add python-docx`"
        ) from e

    return fitz, Vision, Quartz, Foundation, Document


@dataclass(frozen=True)
class OcrOptions:
    dpi: int = 300
    language_hints: tuple[str, ...] = ("en-US",)
    accurate: bool = True
    engine: str = "auto"  # auto | vision | tesseract


def _pdf_page_png_bytes(fitz: object, pdf_path: Path, page_index: int, dpi: int) -> bytes:
    doc = fitz.open(str(pdf_path))
    try:
        page = doc.load_page(page_index)
        zoom = dpi / 72.0
        mat = fitz.Matrix(zoom, zoom)
        pix = page.get_pixmap(matrix=mat, alpha=False)
        return pix.tobytes("png")
    finally:
        doc.close()


def _png_bytes_to_cgimage(Quartz: object, Foundation: object, png_bytes: bytes) -> object:
    data = Foundation.NSData.dataWithBytes_length_(png_bytes, len(png_bytes))
    src = Quartz.CGImageSourceCreateWithData(data, None)
    if src is None:
        raise RuntimeError("Failed to create image source for OCR.")
    img = Quartz.CGImageSourceCreateImageAtIndex(src, 0, None)
    if img is None:
        raise RuntimeError("Failed to create CGImage for OCR.")
    return img


def _ocr_png_bytes(Vision: object, Foundation: object, png_bytes: bytes, opts: OcrOptions) -> str:
    lines: list[str] = []
    handler_error_messages: list[str] = []

    def handler(request, error):  # noqa: ANN001
        if error is not None:
            handler_error_messages.append(str(error))
            return
        results = request.results() or []
        for obs in results:
            candidates = obs.topCandidates_(1)
            if candidates and len(candidates) > 0:
                s = candidates[0].string()
                if s:
                    lines.append(str(s))

    req = Vision.VNRecognizeTextRequest.alloc().initWithCompletionHandler_(handler)
    req.setRecognitionLanguages_(list(opts.language_hints))
    req.setUsesLanguageCorrection_(True)
    req.setRecognitionLevel_(
        Vision.VNRequestTextRecognitionLevelAccurate
        if opts.accurate
        else Vision.VNRequestTextRecognitionLevelFast
    )

    data = Foundation.NSData.dataWithBytes_length_(png_bytes, len(png_bytes))
    img_handler = Vision.VNImageRequestHandler.alloc().initWithData_options_(data, None)
    ok, err = img_handler.performRequests_error_([req], None)
    if not ok:
        details = []
        if err is not None:
            details.append(str(err))
        details.extend(handler_error_messages)
        raise RuntimeError("OCR failed." + (f" Details: {' | '.join(details)}" if details else ""))

    return "\n".join(lines).strip()


def _ocr_with_tesseract(png_bytes: bytes) -> str:
    exe = shutil.which("tesseract")
    if not exe:
        raise RuntimeError(
            "Tesseract not found on PATH. Install it (e.g. `brew install tesseract`) "
            "or use the Vision engine on macOS."
        )

    with tempfile.TemporaryDirectory() as td:
        img_path = Path(td) / "page.png"
        img_path.write_bytes(png_bytes)
        # Output to stdout to avoid side-effect files.
        proc = subprocess.run(
            [exe, str(img_path), "stdout", "-l", "eng"],
            check=False,
            capture_output=True,
            text=True,
        )
        if proc.returncode != 0:
            raise RuntimeError(
                "Tesseract OCR failed."
                + (f" stderr: {proc.stderr.strip()}" if proc.stderr.strip() else "")
            )
        return proc.stdout.strip()


def ocr_pdf_to_docx(pdf_path: Path, out_docx: Path, opts: OcrOptions) -> None:
    fitz, Vision, Quartz, Foundation, Document = _require_imports()

    doc = fitz.open(str(pdf_path))
    try:
        page_count = doc.page_count
    finally:
        doc.close()

    word_doc = Document()
    word_doc.add_heading(pdf_path.name, level=1)

    for i in range(page_count):
        png = _pdf_page_png_bytes(fitz, pdf_path, i, opts.dpi)
        # Create CGImage as a sanity check that the PNG is well-formed; OCR uses the bytes path.
        _ = _png_bytes_to_cgimage(Quartz, Foundation, png)
        if opts.engine not in ("auto", "vision", "tesseract"):
            raise RuntimeError(f"Unknown OCR engine: {opts.engine}")

        text = ""
        if opts.engine in ("auto", "vision"):
            try:
                text = _ocr_png_bytes(Vision, Foundation, png, opts)
            except Exception:
                if opts.engine == "vision":
                    raise

        if not text and opts.engine in ("auto", "tesseract"):
            text = _ocr_with_tesseract(png)

        word_doc.add_heading(f"Page {i + 1}", level=2)
        word_doc.add_paragraph(text if text else "")
        if i != page_count - 1:
            word_doc.add_page_break()

    out_docx.parent.mkdir(parents=True, exist_ok=True)
    word_doc.save(str(out_docx))


def _iter_pdfs(inputs: Iterable[Path]) -> list[Path]:
    pdfs: list[Path] = []
    for p in inputs:
        if p.is_dir():
            pdfs.extend(sorted(p.glob("*.pdf")))
        else:
            pdfs.append(p)
    return pdfs


def main() -> int:
    ap = argparse.ArgumentParser(
        description="OCR scanned PDFs (image-only) into a new .docx (one section per page)."
    )
    ap.add_argument("inputs", nargs="+", help="PDF file(s) and/or directories containing PDFs.")
    ap.add_argument(
        "--out-dir",
        default="ocr_output",
        help="Output directory (default: ocr_output)",
    )
    ap.add_argument("--dpi", type=int, default=300, help="Render DPI for OCR (default: 300)")
    ap.add_argument(
        "--engine",
        default="auto",
        choices=["auto", "vision", "tesseract"],
        help="OCR engine (default: auto)",
    )
    args = ap.parse_args()

    in_paths = [Path(s).expanduser().resolve() for s in args.inputs]
    out_dir = Path(args.out_dir).expanduser().resolve()
    opts = OcrOptions(dpi=args.dpi, engine=args.engine)

    pdfs = _iter_pdfs(in_paths)
    if not pdfs:
        raise SystemExit("No PDFs found.")

    for pdf in pdfs:
        out_docx = out_dir / f"{pdf.stem}.ocr.docx"
        ocr_pdf_to_docx(pdf, out_docx, opts)
        print(f"Wrote {out_docx}")

    return 0


if __name__ == "__main__":
    raise SystemExit(main())
