import calendar
import csv
import glob
import hashlib
import json
import os
import random
import shutil
import subprocess
import sys
from collections import defaultdict
import re
import warnings
from datetime import date, datetime, timedelta
from zoneinfo import ZoneInfo

from docx import Document

# 1. DATA STRUCTURES
locations_timeline = [
    ("2018-07-01", "2021-08-15", "Palo Alto, California"),
    ("2021-08-15", "2021-10-15", "Amstelveen, Holland"),
    ("2021-10-15", "2021-11-15", "Barcelona, Spain"),
    ("2022-03-15", "2022-05-15", "Sunny Isles, Florida"),
    ("2022-05-15", "2022-07-05", "Wayne, Pennsylvania"),
    ("2022-07-05", "2022-07-15", "Portland, Oregon"),
    ("2022-07-15", "2023-09-27", "Wayne, Pennsylvania"),
    ("2023-09-27", "2023-10-09", "Knoxville, TN"),
    ("2023-10-09", "2023-10-19", "Austin, TX"),
    ("2023-10-19", "2023-10-25", "Palo Alto, California"),
    ("2023-10-25", "2023-10-29", "Surrey, Canada"),
    ("2023-10-29", "2023-11-07", "Knoxville, TN"),
    ("2023-11-08", "2023-11-27", "Belgrade, Serbia"),
    ("2023-11-27", "2024-01-13", "Thessaloniki, Greece"),
    ("2024-01-13", "2024-01-27", "Valencia, Spain"),
    ("2024-01-28", "2024-02-14", "Knoxville, TN"),
    ("2024-02-14", "2024-04-21", "Belgrade, Serbia"),
    ("2024-04-21", "2024-05-17", "Thessaloniki, Greece"),
    ("2024-05-17", "2024-07-06", "Belgrade, Serbia"),
    ("2024-07-06", "2024-07-12", "Athens, Greece"),
    ("2024-07-12", "2024-07-19", "Istanbul, Turkey"),
    ("2024-07-19", "2024-08-19", "Birmingham, England"),
    ("2024-08-19", "2024-10-31", "Hoegaarden, Belgium"),
    ("2024-10-31", "2024-11-18", "Birmingham, England"),
    ("2024-11-18", "2025-02-15", "Denver, NC"),
    ("2025-02-15", "2025-03-31", "Lantana, FL"),
    ("2025-04-01", "2026-04-30", "Wayne, Pennsylvania"),
    ("2026-05-01", "2026-12-31", "Wayne, Pennsylvania")
]

# Stockholder-side instruments:
# - "written_consent": sole stockholder action by written consent under DGCL § 228 (charter must authorize written consent).
# - "annual_meeting_stockholders": formal minutes of an annual meeting of stockholders (e.g. multiple stockholders).
#
# Location in minutes:
# - use_timeline_place: include principal-operational place from locations_timeline for the meeting date (all corps here).
# - virtual_ok: Delaware permits remote meetings when notice/boilerplate matches your bylaws; minutes use "via digital
#   communication" alongside the timeline place for consistency. Set False only if you intentionally omit remote language.
#
# Annual scheduling constraints (sole director across multiple corporations):
# - annual_day_offset staggers annual meetings on consecutive weekdays in December
#   (computed as Monday-based business days starting from the first Monday on or after December 8).
# - All annual meetings for a single corporation occur on the same day, commencing at 1:00 PM (sequence is documented in minutes).
# Optional per company: voting_shares_description — phrase after "holding" in §228 written consent if not all voting power is one class.
# minutes_display_name — legal name as it should appear in minutes (e.g. "Hippo, Inc.") when the dict key omits a period.
# stockholders_roll_call — for annual_meeting_stockholders only: list of {"name": str, "presence": str} for roll call.
# stockholders_quorum_collective_sentence — follows roll call; explains majority quorum.
# stockholders_absent_line — text under Stockholders Absent (default "None.").
# annual_board_resolution_blocks — optional list of markdown strings (each: optional **title** line + RESOLVED, …). If key
#   omitted, AGM uses the default three resolutions. If [] (empty list), minutes use “no resolutions” language.
# quarterly_resolution_blocks — same for quarterly board minutes; default one ratification. [] for no resolutions.
# agm_president_report_product_line — optional AGM sentence(s) after the standard ops/dev locations paragraph: factual product
#   scope (e.g. API / web / mobile). If omitted, a short generic “continued development…” sentence is used.
# agm_president_report_infrastructure_line — optional sentence on hosting / infrastructure (e.g. named providers).
# agm_president_report_operating_exhibit_label — if set (e.g. "Exhibit B"), minutes may reference a written addendum with deeper
#   specs, roadmaps, KPIs, and diagrams. When `audit_reports/all_corp_accomplishments_2021-2025.json` lists detailed
#   `annual_report` bullets for the year, an operating addendum .docx is generated and the minutes reference **detailed accomplishments**.
#   If detailed bullets exist but this key is omitted, **Exhibit B** is used by default for that addendum.
#
# Cookie-cutter controls (optional; each company in `company_information` should set these for distinct minutes):
# - development_centers_line — semicolon-separated regions where contractor/consultant personnel supporting the
#   Corporation are based (President’s Report / quarterlies; not an assertion that the Corporation owns those premises).
# - primary_banking_institution — bank name in the default banking RESOLVED (e.g. "JPMorgan Chase Bank, N.A.").
# - agm_president_report_product_line — President’s Report product summary: a **str** (optional `{year}` / `{next_year}`) or a **dict** keyed by
#   calendar year strings (e.g. `"2023"`) plus optional `"default"` for other years.
# - agm_discussion_items_line — AGM §VI paragraph; may include `{next_year}` / `{year}` placeholders.
# - special_meeting_purpose — one-line Purpose field in the annual special board meeting.
# - special_meeting_ratification_resolution_markdown — first special-meeting resolution block; `{year}` / `{next_year}` allowed.
# - treasurer_contingent_obligations_clause — sentence fragment after "certain outstanding obligations, " in Treasurer’s Report.
# - treasurer_report_minutes_paragraph — optional full Treasurer’s Report paragraph; `{issued}` and `{par}` placeholders.
# - quarterly_default_ratification_resolution — default quarterly RESOLVED when quarterly_resolution_blocks omitted.
# - agm_ip_affirmation_sentence — optional closing sentence for President’s Report IP affirmation.
# - signature_block_print_signing_lines — optional True: emit a rule line plus **Name:** / **Title:** / **Date:** for wet-ink execution.
# - signature_block_signing_rule_line — optional underscore rule (default long underscore rule).
# - minutes_principal_address_note — optional markdown line/paragraph after **Principal Address:** (overrides the default note). Set to "" to suppress.
# - board_meeting_remote_presence_markdown — optional sentence for remote participation = presence (AGM/special/quarterly); default is shared boilerplate. `""` omits.
# - board_meeting_reliance_markdown — optional reliance paragraph after Treasurer’s Report; default is `reliance_standard(co)` (same for all DE corps). `""` omits.
# - annual_stockholder_director_election_votes — optional for `annual_meeting_stockholders` corps: list of
#   `{"name": str, "vote": "FOR"|"WITHHOLD"|"ABSTAIN", "shares": str optional}`; when set, minutes include a **VI-A. Vote Tabulation** block after §VI.
# - board_roll_quorum_layout — int 0–3: sole-director roll/quorum/notice/remote block shape (differs by company); meeting-to-meeting wording still rotates.
# - board_directors — optional list of `{"name": ..., "title": "Director"}`; when **two or more**, board minutes use full-board roll call / quorum,
#   “Board adopted” resolutions, Chair call-to-order (`board_meeting_chair_name`), and dual signatures. Omit for sole-director companies.
# - board_sole_director_first_chronological_meeting — optional **True**: when `board_directors` lists two or more directors, the **first** board meeting
#   in calendar order (across the minute series) is minuted as **sole director**; every later meeting uses the full `board_directors` list (e.g. appointment
#   effective after the initial meeting). Waiver-of-notice instruments still name the full board when configured.
# - organizational_meeting_full_board — optional **True**: when an organizational meeting exists (`inc_year` only), minute it with the **full** `board_directors`
#   roster (no sole-director strip for kind `org`), and do **not** insert “Appointment of Additional Director(s)” at that meeting. Use when the board
#   intends both directors present at the post-filing organizational meeting (e.g. TeamBoost org before Q1).
# - board_meeting_chair_name — optional; used for “called to order … acting as Chair of the Board” when `board_directors` is set.
# - agm_banking_authorized_signatory — optional name in the default banking RESOLVED when the board is not sole-director (else `director_name`).
# - agm_president_report_opening_paragraph_markdown — optional `{office_locations}` / `{dev_locations}` / `{year}` paragraph replacing the default “centralized … development” opener.
# - quarterly_business_review_minutes_markdown — optional `{year}` / `{quarter}` / `{dev_locations}` template replacing default quarterly “development centers” review.
# - minute_book_compilation_preamble_markdown — optional `{display_company}` / `{first_year}` / `{last_year}` cover text for the compiled book.
# - minutes_assert_exhibits_filed — if **True**, minutes may state exhibits are **on file** / **annexed**; default **False** uses
#   **to be filed upon execution** / **designated for attachment** wording so generated text does not over-claim filing.
# - board_meeting_materials_acknowledgment_markdown — optional paragraph before business/resolutions (AGM after § IV;
#   special/quarterly after roll call) stating materials the Sole Director reviewed; must match real exhibits/files (generator does not invent them).
#
# Schedule randomization (optional; reproducible with a seed):
# - Set env `CORPORATE_MINUTES_SCHEDULE_SEED` to an int (or any string hashed to an int), or pass `--schedule-seed` to the CLI.
# - `schedule_time_jitter_minutes` (int): max ± minutes applied to nominal meeting times (rounded by `schedule_time_round_minutes`, default 5).
# - `schedule_annual_weekday_jitter` (int): max ± **weekdays** shifted from the December anchor (still clamped to December of that year).
# - `schedule_quarterly_calendar_jitter` (int): max ± **calendar** days added to each quarterly meeting date.
# - `schedule_seed_suffix` (str): optional extra salt so two registries with similar names diverge.
# - `schedule_same_day_gap_minutes` (int, default 45): minimum start-to-start gap when the **special** board meeting shares a calendar
#   day with the **annual** board meeting (written-consent corporations) so jitter cannot reverse chronology.
# - `schedule_stockholder_to_board_gap_minutes` (int, default 0): extra minutes between stockholder annual start and board AGM start
#   on the same day (`0` preserves “immediately following” / same nominal clock).
# Narrative paragraphs in the minutes are **not** randomized—only dates/times (when jitter + seed are enabled) and file mtimes
# (`_random_utime_after_meeting`) introduce variation.
STOCKHOLDER_MEETING_TIME = "1:00 PM"
BOARD_AGM_TIME = "1:00 PM"
QUARTERLY_MEETING_TIME = "1:00 PM"
SPECIAL_MEETING_TIME = "12:00 PM"
ORGANIZATIONAL_MEETING_TIME = "10:00 AM"


def organizational_meeting_date_str(co: dict, year: int) -> str | None:
    """ISO date for the post-filing organizational meeting (if configured for this company/year).

    This is intentionally **not** tied to December annual scheduling; it's meant to occur shortly after the
    SOS acceptance/filing date captured as `incorporation_filed_date_iso`.
    """
    if year != co.get("inc_year"):
        return None
    filed = _incorporation_filed_date_iso(co)
    if not filed:
        return None
    base = date.fromisoformat(filed)
    days_after = int(co.get("organizational_meeting_days_after_filing", 7) or 7)
    d = base + timedelta(days=days_after)
    # Shift off weekends (forward to Monday).
    while d.weekday() >= 5:
        d += timedelta(days=1)
    return d.strftime("%Y-%m-%d")


def _quarter_month_for(quarter: str) -> int:
    return {"Q1": 4, "Q2": 7, "Q3": 10, "Q4": 12}[quarter]


def _quarterly_exists_for_year(co: dict, year: int, quarter: str) -> bool:
    """Whether a quarterly governance meeting should exist for this year/quarter.

    For the incorporation year, quarters whose anchor month precedes the actual SOS filing month are omitted
    (you can't hold a board meeting before the corporation exists).
    """
    if year != co.get("inc_year"):
        return True
    filed = _incorporation_filed_date_iso(co)
    if not filed:
        return True
    filed_d = date.fromisoformat(filed)
    qm = _quarter_month_for(quarter)
    return qm >= filed_d.month

# Timeline place label (from get_location) → IANA timezone for file mtimes (meeting “local” business hours).
_TIMELINE_LOCATION_TZ: dict[str, str] = {
    "Palo Alto, California": "America/Los_Angeles",
    "Amstelveen, Holland": "Europe/Amsterdam",
    "Barcelona, Spain": "Europe/Madrid",
    "Sunny Isles, Florida": "America/New_York",
    "Wayne, Pennsylvania": "America/New_York",
    "Portland, Oregon": "America/Los_Angeles",
    "Knoxville, TN": "America/New_York",
    "Austin, TX": "America/Chicago",
    "Surrey, Canada": "America/Vancouver",
    "Belgrade, Serbia": "Europe/Belgrade",
    "Thessaloniki, Greece": "Europe/Athens",
    "Valencia, Spain": "Europe/Madrid",
    "Athens, Greece": "Europe/Athens",
    "Istanbul, Turkey": "Europe/Istanbul",
    "Birmingham, England": "Europe/London",
    "Hoegaarden, Belgium": "Europe/Brussels",
    "Denver, NC": "America/New_York",
    "Lantana, FL": "America/New_York",
}

def _jurisdiction(co: dict) -> str:
    """Two-letter jurisdiction code (default: DE)."""
    return str(co.get("jurisdiction") or "DE").strip().upper()


def _corporation_parenthetical(co: dict) -> str:
    """Parenthetical used on covers (rendered as plain text in .docx)."""
    j = _jurisdiction(co)
    if j == "WY":
        return "(Wyoming corporation)"
    return "(Delaware corporation)"


def _incorporation_filed_date_iso(co: dict) -> str | None:
    """Optional ISO filing/acceptance date for the certificate/articles."""
    raw = co.get("incorporation_filed_date_iso")
    if not isinstance(raw, str) or not raw.strip():
        return None
    s = raw.strip()
    # Validate shape early; templates treat this as factual.
    try:
        date.fromisoformat(s)
    except ValueError:
        return None
    return s


def _fmt_long_date(d_iso: str) -> str:
    return datetime.strptime(d_iso, "%Y-%m-%d").strftime("%B %d, %Y")


# Stock ledger JSON → (ledger entry, ledger payload) grouped by first board meeting **strictly after** `issue_date`.
_STOCK_LEDGER_ENTRIES_BY_MEETING: dict[tuple[str, str], list[tuple[dict, dict]]] | None = None


def reset_stock_ledger_meeting_index() -> None:
    """Clear cached stock-ledger / meeting index (e.g. before tests or standalone resolution export)."""
    global _STOCK_LEDGER_ENTRIES_BY_MEETING
    _STOCK_LEDGER_ENTRIES_BY_MEETING = None


def _reset_stock_ledger_meeting_index_for_tests() -> None:
    reset_stock_ledger_meeting_index()


def _norm_ledger_company_name(s: str) -> str:
    return re.sub(r"\s+", " ", s.strip().rstrip(".")).lower()


def _co_name_from_ledger_company_name(ledger_company_name: str) -> str | None:
    target = _norm_ledger_company_name(ledger_company_name)
    for key, co in company_information.items():
        if _norm_ledger_company_name(key) == target:
            return key
        md = co.get("minutes_display_name")
        if md and _norm_ledger_company_name(str(md)) == target:
            return key
    return None


_CAP_TABLE_JSON_CACHE: dict | None = None


def _load_cap_table_json() -> dict:
    global _CAP_TABLE_JSON_CACHE
    if _CAP_TABLE_JSON_CACHE is not None:
        return _CAP_TABLE_JSON_CACHE
    path = os.path.join(os.path.dirname(os.path.abspath(__file__)), "data", "cap_table.json")
    try:
        with open(path, encoding="utf-8") as f:
            _CAP_TABLE_JSON_CACHE = json.load(f)
    except (OSError, json.JSONDecodeError):
        _CAP_TABLE_JSON_CACHE = {"schema_version": "0", "companies": {}}
    return _CAP_TABLE_JSON_CACHE


def _cap_table_companies_block() -> dict:
    c = _load_cap_table_json().get("companies")
    return c if isinstance(c, dict) else {}


def _find_cap_table_json_key(co_name: str, co: dict) -> str | None:
    cap = _cap_table_companies_block()
    if not cap:
        return None
    for candidate in (co_name, str(co.get("minutes_display_name") or "").strip()):
        if not candidate:
            continue
        t = _norm_ledger_company_name(candidate)
        for k in cap.keys():
            if _norm_ledger_company_name(str(k)) == t:
                return str(k)
    return None


def cap_table_document_markdown(co_name: str, co: dict) -> str:
    """Markdown body for `*_cap_table.docx` and compiled-book appendix (from `data/cap_table.json`)."""
    display = minutes_display_name(co_name)
    root = _load_cap_table_json()
    as_of = str(root.get("as_of") or "").strip() or "(date not set in cap_table.json)"
    desc = str(root.get("description") or "").strip()
    lines: list[str] = [
        f"**Cap table summary — {display}**",
        f"**As of (file):** {as_of}",
        "",
        "This document is a **convenience summary** produced from `data/cap_table.json`. It does not replace the "
        "Corporation’s official stock ledger, executed subscription agreements, or any third-party cap-table system.",
        "",
    ]
    if desc:
        lines.append(f"_{desc}_")
        lines.append("")
    key = _find_cap_table_json_key(co_name, co)
    if not key:
        lines.append(
            "**Status:** No matching company block exists in `data/cap_table.json` for this registry entry. "
            "Add one (keyed by legal name) or align `minutes_display_name` with an existing key."
        )
        si = co.get("shares_issued")
        if isinstance(si, dict) and si:
            years = [int(y) for y in si if str(y).isdigit()]
            if years:
                ly = str(max(years))
                val = si.get(ly) or si.get(max(years))
                lines.append("")
                lines.append(
                    f"**Reference only — `company_information.shares_issued` ({ly}):** **{val}** "
                    "(not imported into the cap table file until you add an explicit summary block)."
                )
        return "\n".join(lines).strip()

    block = _cap_table_companies_block().get(key) or {}
    if not isinstance(block, dict):
        block = {}
    jur = str(block.get("jurisdiction") or _jurisdiction(co)).strip()
    lines.append(f"**Jurisdiction (summary):** {jur}")
    lines.append("")
    holders = block.get("holders")
    if isinstance(holders, list) and holders:
        lines.append("**Holders (from summary file)**")
        for h in holders:
            if not isinstance(h, dict):
                continue
            hn = str(h.get("holder") or "").strip() or "(holder name TBD)"
            sh = h.get("shares")
            st = str(h.get("status") or "").strip()
            shs = f"{sh:,}" if isinstance(sh, int) else str(sh or "").strip() or "—"
            st_part = f" — **{st}**" if st else ""
            lines.append(f"- **{hn}**: {shs} shares{st_part}")
        lines.append("")
    tot = block.get("total_issued_shares_from_ledger")
    if isinstance(tot, int):
        lines.append(f"**Total issued shares (summary field):** {tot:,}")
    elif tot is not None and str(tot).strip():
        lines.append(f"**Total issued shares (summary field):** {tot}")
    notes = str(block.get("notes") or "").strip()
    if notes:
        lines.append("")
        lines.append(f"**Notes:** {notes}")
    return "\n".join(lines).strip()


def _stock_ledger_payload_for_company(co_name: str) -> dict | None:
    for path in _stock_ledger_json_paths():
        try:
            with open(path, encoding="utf-8") as f:
                ledger = json.load(f)
        except (OSError, json.JSONDecodeError):
            continue
        legal = str(ledger.get("company_legal_name") or "").strip()
        if _co_name_from_ledger_company_name(legal) == co_name:
            return ledger
    return None


def _markdown_lines_for_consideration(cons: dict) -> list[str]:
    if not isinstance(cons, dict):
        return []
    out: list[str] = []
    typ = str(cons.get("type") or "").strip()
    if typ:
        out.append(f"- **Type:** {typ}")
    amt = cons.get("amount")
    cur = str(cons.get("currency") or "").strip()
    if amt is not None and str(amt).strip():
        out.append(f"- **Amount:** {amt} {cur}".strip())
    for label, field in (
        ("Payment method", "payment_method"),
        ("From", "bank_source"),
        ("To", "bank_destination"),
    ):
        v = str(cons.get(field) or "").strip()
        if v:
            out.append(f"- **{label}:** {v}")
    n = str(cons.get("notes") or "").strip()
    if n:
        out.append(f"- **Payment notes:** {n}")
    return out


def stock_ledger_document_markdown(co_name: str, co: dict) -> str:
    """Markdown body for `*_stock_ledger.docx` and compiled-book appendix (from `data/stock_ledgers/*.json`)."""
    display = minutes_display_name(co_name)
    ledger = _stock_ledger_payload_for_company(co_name)
    lines: list[str] = [
        f"**Stock ledger (machine-readable excerpt) — {display}**",
        "",
        "This document reflects `data/stock_ledgers/*.json` entries linked to the corporation by `company_legal_name`. "
        "The **official** stock ledger may be maintained in another system or in bound paper form; reconcile before any filing or diligence.",
        "",
    ]
    if not ledger:
        lines.append(
            "**Status:** No stock ledger JSON is linked to this company. Add `data/stock_ledgers/<slug>.json` with "
            "`company_legal_name` matching the registry (or `minutes_display_name`) to populate this appendix."
        )
        return "\n".join(lines).strip()

    legal = str(ledger.get("company_legal_name") or display).strip()
    lines.append(f"**Company (ledger file):** {legal}")
    jur = str(ledger.get("jurisdiction") or _jurisdiction(co)).strip()
    lines.append(f"**Jurisdiction (ledger):** {jur}")
    ta = ledger.get("total_authorized_shares")
    if isinstance(ta, int):
        lines.append(f"**Total authorized shares:** {ta:,}")
    elif ta is not None and str(ta).strip():
        lines.append(f"**Total authorized shares:** {ta}")
    par = ledger.get("par_value_per_share_usd")
    if isinstance(par, (int, float)):
        lines.append(f"**Par value (USD per share):** {par}")
    elif par is not None and str(par).strip():
        lines.append(f"**Par value (USD per share):** {par}")
    cur = str(ledger.get("currency") or "").strip()
    if cur:
        lines.append(f"**Currency:** {cur}")
    pol = str(ledger.get("certificate_numbering_policy") or "").strip()
    if pol:
        lines.append(f"**Certificate numbering:** {pol}")
    top_notes = str(ledger.get("notes") or "").strip()
    if top_notes:
        lines.append("")
        lines.append(f"**Ledger file notes:** {top_notes}")
    lines.append("")
    entries = ledger.get("ledger_entries")
    if not isinstance(entries, list) or not entries:
        lines.append("**Ledger entries:** *(none listed in JSON — add `ledger_entries` to mirror issuances.)*")
        return "\n".join(lines).strip()

    lines.append("**Ledger entries**")
    lines.append("")
    for i, ent in enumerate(entries, start=1):
        if not isinstance(ent, dict):
            continue
        cert = str(ent.get("certificate_number") or "").strip() or f"(entry {i})"
        sh_raw = ent.get("shares")
        sh_disp = f"{sh_raw:,}" if isinstance(sh_raw, int) else str(sh_raw or "").strip() or "—"
        sh_name = str(ent.get("shareholder") or "").strip() or "(shareholder TBD)"
        idate = str(ent.get("issue_date") or "").strip() or "(issue date TBD)"
        lines.append(f"### {cert}")
        lines.append(f"- **Shareholder:** {sh_name}")
        lines.append(f"- **Shares:** {sh_disp}")
        lines.append(f"- **Issue date:** {idate}")
        cons = ent.get("consideration") if isinstance(ent.get("consideration"), dict) else {}
        sub = _markdown_lines_for_consideration(cons)
        if sub:
            lines.append("- **Consideration:**")
            lines.extend(f"  {s}" for s in sub)
        en = str(ent.get("notes") or "").strip()
        if en:
            lines.append(f"- **Entry notes:** {en}")
        lines.append("")
    return "\n".join(lines).rstrip()


def _write_cap_table_and_stock_ledger_docx(
    safe_company_name: str,
    co_name: str,
    co: dict,
    cap_tables_dir: str,
    stock_ledgers_dir: str,
    meeting_date_iso: str,
) -> None:
    """Emit cap table `.docx` + `.csv` under ``cap_tables/`` and stock ledger `.docx` under ``stock_ledgers/`` (siblings of ``books/``)."""
    os.makedirs(cap_tables_dir, exist_ok=True)
    os.makedirs(stock_ledgers_dir, exist_ok=True)
    cap = cap_table_document_markdown(co_name, co)
    led = stock_ledger_document_markdown(co_name, co)
    cap_path = os.path.join(cap_tables_dir, f"{safe_company_name}_cap_table.docx")
    led_path = os.path.join(stock_ledgers_dir, f"{safe_company_name}_stock_ledger.docx")
    write_docx_from_minutes(cap, cap_path, meeting_date_iso, co_name)
    write_docx_from_minutes(led, led_path, meeting_date_iso, co_name)
    print(f"Writing cap table summary to {cap_path}")
    print(f"Writing stock ledger excerpt to {led_path}")
    csv_path = os.path.join(cap_tables_dir, f"{safe_company_name}_cap_table_carta_pulley.csv")
    write_cap_table_carta_pulley_csv(csv_path, co_name, co)
    print(f"Writing Carta/Pulley-style cap table CSV to {csv_path}")


# Column set aligned to common spreadsheet imports (Carta “spreadsheet import” / Pulley cap-table CSV):
# map columns in the vendor UI if headers differ slightly. `stakeholder_email` / `vesting_schedule_description` are
# intentionally blank for purchased common unless you extend the JSON schema.
CAP_TABLE_CARTA_PULLEY_CSV_FIELDNAMES = [
    "company_legal_name",
    "stakeholder_name",
    "stakeholder_email",
    "stakeholder_type",
    "security_type",
    "share_class",
    "certificate_id",
    "shares",
    "issue_date",
    "price_per_share",
    "purchase_price_total",
    "currency",
    "par_value_per_share",
    "has_vesting",
    "vesting_schedule_description",
    "notes",
]


def cap_table_carta_pulley_rows(co_name: str, co: dict) -> list[dict[str, str]]:
    """One row per ledger issuance (preferred), else one row per `cap_table.json` holder, else a single placeholder row."""
    display_co = minutes_display_name(co_name)
    ledger = _stock_ledger_payload_for_company(co_name)
    cap_key = _find_cap_table_json_key(co_name, co)
    cap_block = _cap_table_companies_block().get(cap_key) if cap_key else None
    if not isinstance(cap_block, dict):
        cap_block = None

    holder_status: dict[str, str] = {}
    if cap_block:
        for h in cap_block.get("holders") or []:
            if isinstance(h, dict) and str(h.get("holder") or "").strip():
                holder_status[_norm_ledger_company_name(str(h["holder"]).strip())] = str(
                    h.get("status") or ""
                ).strip()

    ledger_legal = str(ledger.get("company_legal_name") or "").strip() if ledger else ""
    company_cell = ledger_legal or display_co
    cur_default = str(ledger.get("currency") or "USD").strip() if ledger else "USD"
    par_raw = ledger.get("par_value_per_share_usd") if ledger else None
    if isinstance(par_raw, (int, float)):
        par_s = f"{float(par_raw):.12f}".rstrip("0").rstrip(".") or "0"
    elif par_raw is not None and str(par_raw).strip():
        par_s = str(par_raw).strip()
    else:
        par_s = ""

    def _row(
        stakeholder_name: str,
        certificate_id: str,
        shares_val: object,
        issue_date: str,
        price_per_share: str,
        purchase_total: str,
        currency: str,
        notes_parts: list[str],
    ) -> dict[str, str]:
        st = holder_status.get(_norm_ledger_company_name(stakeholder_name), "")
        np = [p for p in notes_parts if p]
        if st:
            np.append(f"cap_table_status={st}")
        sh_disp = ""
        if isinstance(shares_val, int):
            sh_disp = str(shares_val)
        elif shares_val is not None and str(shares_val).strip():
            sh_disp = str(shares_val).strip()
        return {
            "company_legal_name": company_cell,
            "stakeholder_name": stakeholder_name,
            "stakeholder_email": "",
            "stakeholder_type": "Individual",
            "security_type": "Common Stock",
            "share_class": "Common",
            "certificate_id": certificate_id,
            "shares": sh_disp,
            "issue_date": issue_date,
            "price_per_share": price_per_share,
            "purchase_price_total": purchase_total,
            "currency": currency or cur_default,
            "par_value_per_share": par_s,
            "has_vesting": "No",
            "vesting_schedule_description": "",
            "notes": " | ".join(np) if np else "",
        }

    rows: list[dict[str, str]] = []
    entries = ledger.get("ledger_entries") if ledger else None
    if ledger and isinstance(entries, list) and entries:
        for ent in entries:
            if not isinstance(ent, dict):
                continue
            sh_name = str(ent.get("shareholder") or "").strip() or "(shareholder TBD)"
            cert = str(ent.get("certificate_number") or "").strip()
            shares_raw = ent.get("shares")
            sh_int: int | None
            if isinstance(shares_raw, int):
                sh_int = shares_raw
            else:
                sh_int = None
                if shares_raw is not None and str(shares_raw).strip():
                    try:
                        sh_int = int(str(shares_raw).replace(",", "").strip(), 10)
                    except ValueError:
                        sh_int = None
            idate = str(ent.get("issue_date") or "").strip()[:10]
            cons = ent.get("consideration") if isinstance(ent.get("consideration"), dict) else {}
            amt = cons.get("amount")
            amt_f: float | None
            try:
                if amt is None or (isinstance(amt, str) and not amt.strip()):
                    amt_f = None
                else:
                    amt_f = float(amt)
            except (TypeError, ValueError):
                amt_f = None
            cur = str(cons.get("currency") or cur_default).strip()
            price = ""
            pur = ""
            if amt_f is not None and sh_int and sh_int > 0:
                pur = f"{amt_f:.10f}".rstrip("0").rstrip(".")
                price = f"{(amt_f / sh_int):.12f}".rstrip("0").rstrip(".")
            note_bits = [str(ent.get("notes") or "").strip()]
            rows.append(
                _row(
                    sh_name,
                    cert,
                    sh_int if sh_int is not None else shares_raw,
                    idate,
                    price,
                    pur,
                    cur,
                    note_bits,
                )
            )
        return rows

    if cap_block and isinstance(cap_block.get("holders"), list) and cap_block["holders"]:
        for h in cap_block["holders"]:
            if not isinstance(h, dict):
                continue
            hn = str(h.get("holder") or "").strip()
            if not hn:
                continue
            sh = h.get("shares")
            sh_int = int(sh) if isinstance(sh, int) else None
            rows.append(
                _row(
                    hn,
                    "",
                    sh_int if sh_int is not None else sh,
                    "",
                    "",
                    "",
                    cur_default,
                    ["summary-only row from cap_table.json (no matching ledger line item)"],
                )
            )
        if rows:
            return rows

    rows.append(
        _row(
            "(add stakeholders to data/cap_table.json and ledger JSON)",
            "",
            "",
            "",
            "",
            "",
            cur_default,
            [f"No ledger entries and no cap_table holders for {display_co}"],
        )
    )
    return rows


def write_cap_table_carta_pulley_csv(filepath: str, co_name: str, co: dict) -> None:
    """Write UTF-8 BOM CSV for Excel; one row per issuance (ledger) or summary holder."""
    os.makedirs(os.path.dirname(os.path.abspath(filepath)), exist_ok=True)
    rows = cap_table_carta_pulley_rows(co_name, co)
    with open(filepath, "w", encoding="utf-8-sig", newline="") as f:
        w = csv.DictWriter(
            f,
            fieldnames=CAP_TABLE_CARTA_PULLEY_CSV_FIELDNAMES,
            extrasaction="ignore",
            lineterminator="\n",
        )
        w.writeheader()
        for r in rows:
            w.writerow({k: r.get(k, "") for k in CAP_TABLE_CARTA_PULLEY_CSV_FIELDNAMES})


def write_all_companies_cap_table_carta_pulley_csv(output_root: str) -> str | None:
    """Single workbook-style CSV under `<output_root>/books/` with every registry company (ledger rows)."""
    start_cwd = os.getcwd()
    root_dir = os.path.join(start_cwd, output_root)
    books_dir = os.path.join(root_dir, "books")
    os.makedirs(books_dir, exist_ok=True)
    path = os.path.join(books_dir, "all_companies_cap_table_carta_pulley.csv")
    all_rows: list[dict[str, str]] = []
    for co_name, co in companies.items():
        all_rows.extend(cap_table_carta_pulley_rows(co_name, co))
    if not all_rows:
        return None
    with open(path, "w", encoding="utf-8-sig", newline="") as f:
        w = csv.DictWriter(
            f,
            fieldnames=CAP_TABLE_CARTA_PULLEY_CSV_FIELDNAMES,
            extrasaction="ignore",
            lineterminator="\n",
        )
        w.writeheader()
        for r in all_rows:
            w.writerow({k: r.get(k, "") for k in CAP_TABLE_CARTA_PULLEY_CSV_FIELDNAMES})
    print(f"Writing combined Carta/Pulley cap table CSV to {path}")
    return path


def _stock_ledger_json_paths() -> list[str]:
    d = os.path.join(os.path.dirname(os.path.abspath(__file__)), "data", "stock_ledgers")
    if not os.path.isdir(d):
        return []
    return sorted(glob.glob(os.path.join(d, "*.json")))


def _board_meeting_kind_from_row_title(title: str) -> str:
    low = title.lower()
    if "organizational" in low:
        return "org"
    if "special" in low:
        return "special"
    if "annual" in low and "stockholder" not in low:
        return "agm"
    if "quarterly" in low:
        return title.strip().split()[-1]
    return "unknown"


def _first_board_meeting_strictly_after(co_name: str, co: dict, d_issue: date) -> tuple[str, str] | None:
    """First scheduled board meeting date (ISO) and row kind, strictly after `d_issue` (purchase / issue date)."""
    start = int(co.get("minutes_start_year", co["inc_year"]))
    end = max(_board_series_last_calendar_year(co, start), d_issue.year)
    candidates: list[tuple[str, str]] = []
    for y in range(start, end + 1):
        for d_iso, title, _t_str, _place in _board_meeting_rows_for_year(co, y):
            try:
                md = date.fromisoformat(d_iso)
            except ValueError:
                continue
            if md <= d_issue:
                continue
            candidates.append((d_iso, _board_meeting_kind_from_row_title(title)))
    if not candidates:
        return None
    candidates.sort(key=lambda x: x[0])
    return candidates[0][0], candidates[0][1]


def _stock_ledger_resolution_blocks_for_entry(entry: dict, ledger: dict) -> list[str]:
    """Three separate board resolutions: issuance authority, consideration/payment record, books-and-records."""
    cert = str(entry.get("certificate_number") or "").strip() or "TBD"
    sh = str(entry.get("shareholder") or "").strip() or "the subscriber"
    shares_raw = entry.get("shares")
    if isinstance(shares_raw, int):
        sc = f"{shares_raw:,}"
    else:
        sc = str(shares_raw or "").strip() or "the subscribed"
    issue_raw = str(entry.get("issue_date") or "").strip()[:10]
    long_d = _fmt_long_date(issue_raw) if len(issue_raw) == 10 else "the date recorded in the stock ledger"
    cons = entry.get("consideration") if isinstance(entry.get("consideration"), dict) else {}
    pm = str(cons.get("payment_method") or "").strip() or "the method described in the Corporate records"
    src = str(cons.get("bank_source") or "").strip() or "the payor’s account"
    dst = str(cons.get("bank_destination") or "").strip() or "the Corporation’s account"
    amt = cons.get("amount")
    cur = str(cons.get("currency") or ledger.get("currency") or "USD").strip()
    if amt is not None and str(amt).strip() != "":
        pay_phrase = f"in the stated amount of **{amt} {cur}**"
    else:
        pay_phrase = (
            "for consideration the Corporation acknowledges as received (with payment particulars noted in the stock ledger)"
        )
    r1 = f"""**Issuance of Common Stock — {cert}**  
RESOLVED, that the Corporation is authorized and directed to issue **{sc}** shares of common stock to **{sh}** under certificate number **{cert}**, with an issue date of **{long_d}**, upon the terms and consideration set forth in the companion resolutions in this standalone document."""
    r2 = f"""**Acknowledgment of Consideration — {cert}**  
RESOLVED, that the Corporation acknowledges receipt of consideration for the issuance evidenced by certificate number **{cert}**: payment was received on **{long_d}** via **{pm}** from **{src}** to **{dst}**, {pay_phrase}."""
    r3 = f"""**Stock Ledger and Books and Records — {cert}**  
RESOLVED, that the **Secretary** and other appropriate officers of the Corporation are authorized and directed to record the issuance on the Corporation’s stock ledger and to reflect the issuance accurately in the Corporation’s related books and records."""
    return [r1, r2, r3]


def _ensure_stock_ledger_meeting_index() -> dict[tuple[str, str], list[tuple[dict, dict]]]:
    global _STOCK_LEDGER_ENTRIES_BY_MEETING
    if _STOCK_LEDGER_ENTRIES_BY_MEETING is not None:
        return _STOCK_LEDGER_ENTRIES_BY_MEETING
    out: dict[tuple[str, str], list[tuple[dict, dict]]] = {}
    for path in _stock_ledger_json_paths():
        try:
            with open(path, encoding="utf-8") as f:
                ledger = json.load(f)
        except (OSError, json.JSONDecodeError):
            continue
        legal = str(ledger.get("company_legal_name") or "").strip()
        co_name = _co_name_from_ledger_company_name(legal)
        if not co_name:
            continue
        co = company_information[co_name]
        for entry in ledger.get("ledger_entries") or []:
            if not isinstance(entry, dict):
                continue
            raw_date = entry.get("issue_date")
            if raw_date is None or str(raw_date).strip() == "":
                continue
            try:
                d_issue = date.fromisoformat(str(raw_date).strip()[:10])
            except ValueError:
                continue
            slot = _first_board_meeting_strictly_after(co_name, co, d_issue)
            if not slot:
                continue
            meeting_iso, _kind = slot
            out.setdefault((co_name, meeting_iso), []).append((entry, ledger))
    _STOCK_LEDGER_ENTRIES_BY_MEETING = out
    return out


def _stock_ledger_incorporating_resolution_markdown(co_name: str, meeting_date_iso: str) -> str:
    """One minutes resolution: adopt separate written resolutions (full text in standalone .docx)."""
    pairs = _ensure_stock_ledger_meeting_index().get((co_name, meeting_date_iso), [])
    if not pairs:
        return ""
    certs = ", ".join(
        f"**{str(e.get('certificate_number') or '').strip() or 'TBD'}**" for e, _ in pairs
    )
    return f"""**Written board resolutions (equity)**  
RESOLVED, that the Board hereby approves and adopts the **separate written resolutions of even date** substantially in the form filed with the Secretary as **standalone board resolutions** for certificate number(s) {certs}, which set forth the issuance, consideration, and stock ledger mechanics in full (and are **not reproduced verbatim** in these minutes)."""


def _stock_ledger_resolution_blocks_for_meeting(co_name: str, meeting_date_iso: str) -> list[str]:
    s = _stock_ledger_incorporating_resolution_markdown(co_name, meeting_date_iso).strip()
    return [s] if s else []


def _board_resolution_prefix_blocks(
    co_name: str, co: dict, year: int, meeting_date_iso: str, kind: str
) -> list[str]:
    """Stock-ledger acknowledgments (if any for this meeting date) + optional first-meeting director appointment."""
    return (
        _stock_ledger_resolution_blocks_for_meeting(co_name, meeting_date_iso)
        + _board_appointment_resolution_blocks_if_first_meeting(co_name, co, year, kind)
    )


def _jurisdiction_long_name(code: str) -> str:
    return {"DE": "Delaware", "WY": "Wyoming"}.get(code.strip().upper(), code.strip())


def _financial_year_words(co: dict) -> tuple[str, str]:
    """Return (article, noun) for AGM financial resolutions and reporting (e.g. 'the', 'calendar year')."""
    if co.get("fiscal_year_is_calendar_year", True):
        return "the", "calendar year"
    return "the", "fiscal year"


def _domestication_event(co: dict) -> dict | None:
    """Optional WY→DE domestication/continuation event metadata (used for RG)."""
    ev = co.get("domestication_event")
    if not isinstance(ev, dict):
        return None
    from_j = str(ev.get("from_jurisdiction") or "").strip().upper()
    to_j = str(ev.get("to_jurisdiction") or "").strip().upper()
    eff = str(ev.get("effective_date_iso") or "").strip()
    if not (from_j and to_j and eff):
        return None
    try:
        date.fromisoformat(eff)
    except ValueError:
        return None
    out = {"from_jurisdiction": from_j, "to_jurisdiction": to_j, "effective_date_iso": eff}
    dex = str(ev.get("documents_exhibit_label") or "").strip()
    if dex:
        out["documents_exhibit_label"] = dex
    return out


def _domestication_detailed_resolution_blocks(co: dict) -> list[str]:
    """Full domestication decision text for standalone resolutions document (not minutes)."""
    ev = _domestication_event(co)
    if not ev:
        return []
    from_ln = _jurisdiction_long_name(ev["from_jurisdiction"])
    to_ln = _jurisdiction_long_name(ev["to_jurisdiction"])
    eff_fmt = _fmt_long_date(ev["effective_date_iso"])
    ex = str(ev.get("documents_exhibit_label") or "").strip()
    pend = ""
    if ex:
        pend = (
            f"A copy of the filed domestication / continuation instruments is **annexed as {ex}**."
            if _minutes_assert_exhibits_filed(co)
            else f"Filed domestication / continuation instruments **are to be designated {ex}** for attachment **upon filing**."
        )
    blocks = [
        f"""**Prior jurisdiction ({from_ln}) — ratification through domestication**  
RESOLVED, that all corporate acts taken by or on behalf of the Corporation while domiciled in **{from_ln}** through the effectiveness of its domestication / continuation into **{to_ln}** are hereby ratified, confirmed, and approved in all respects, to the extent permitted by applicable law.""",
        f"""**Domestication / continuation ({from_ln} to {to_ln}) — approval**  
RESOLVED, that the Board hereby approves and confirms the Corporation’s domestication / continuation from **{from_ln}** to **{to_ln}**, effective **{eff_fmt}**, and ratifies and approves all acts taken to effect such domestication / continuation and to maintain the Corporation’s corporate existence and good standing in **{to_ln}**, in each case in all respects.""",
    ]
    if ex:
        blocks.append(
            f"""**Domestication / continuation — instruments and minute book**  
RESOLVED, that {pend} **FURTHER RESOLVED**, that the **Secretary** is authorized and directed to file, index, or cross-file such instruments with the Corporation’s minute book and related corporate records as counsel may advise."""
        )
    return blocks


def _domestication_resolution_blocks_if_due(co: dict, meeting_date_iso: str) -> list[str]:
    """One minutes resolution: adopt separate domestication resolutions (full text in standalone .docx)."""
    ev = _domestication_event(co)
    if not ev:
        return []
    eff = ev["effective_date_iso"]
    if meeting_date_iso < eff:
        return []
    if co.get("_domestication_motion_inserted", False):
        return []
    co["_domestication_motion_inserted"] = True
    from_ln = _jurisdiction_long_name(ev["from_jurisdiction"])
    to_ln = _jurisdiction_long_name(ev["to_jurisdiction"])
    eff_fmt = _fmt_long_date(eff)
    return [
        f"""**Written board resolutions (domestication)**  
RESOLVED, that the Board hereby approves and adopts the **separate written resolutions of even date** substantially in the form filed with the Secretary as **standalone board resolutions**, which set forth the Corporation’s domestication / continuation from **{from_ln}** to **{to_ln}**, effective **{eff_fmt}**, in full (and are **not reproduced verbatim** in these minutes)."""
    ]


def _first_board_meeting_on_or_after_iso(co: dict, eff_iso: str) -> str | None:
    """First scheduled board meeting date (ISO) on or after `eff_iso` (inclusive)."""
    start = int(co.get("minutes_start_year", co["inc_year"]))
    end = max(_board_series_last_calendar_year(co, start), int(eff_iso[:4]))
    for y in range(start, end + 1):
        for d_iso, _title, _t_str, _place in _board_meeting_rows_for_year(co, y):
            if d_iso >= eff_iso:
                return d_iso
    return None


def write_standalone_board_resolution_documents(output_root: str) -> None:
    """Write equity + domestication resolution packets under each company's ``meetings/`` folder.

    Full decision text lives here; meeting minutes use short incorporating resolutions only.
    """
    reset_stock_ledger_meeting_index()
    for _n, _co in companies.items():
        _co.pop("_domestication_motion_inserted", None)

    root_dir = os.path.abspath(output_root)
    idx = _ensure_stock_ledger_meeting_index()
    for (co_name, meet_iso), pairs in sorted(idx.items()):
        if not pairs:
            continue
        co = companies[co_name]
        safe = sanitize_company_name(co_name)
        display = str(co.get("minutes_display_name") or co_name).strip()
        body_blocks: list[str] = []
        for entry, ledger in pairs:
            body_blocks.extend(_stock_ledger_resolution_blocks_for_entry(entry, ledger))
        body = "\n\n".join(body_blocks)
        md = (
            f"**Standalone board resolutions — {display}**\n"
            f"**Related board meeting date (cycle anchor):** {meet_iso}\n"
            f"**Subject:** Equity — issuance, consideration, and stock ledger\n\n"
            f"{body}\n\n"
            f"{SIGNATURE_BLOCK_MARKER}\n\n"
            f"{board_meeting_signature_markdown(co, meet_iso, sole_director_name='Derek E. Pappas')}\n---\n"
        )
        out_dir = os.path.join(root_dir, safe, "meetings")
        os.makedirs(out_dir, exist_ok=True)
        fname = f"{safe}_{meet_iso}_equity_board_resolutions.docx"
        dest_eq = os.path.join(out_dir, fname)
        write_docx_from_minutes(md, dest_eq, meet_iso, co_name)
        print(f"Writing standalone board resolutions to {dest_eq}")

    for co_name, co in companies.items():
        ev = _domestication_event(co)
        if not ev:
            continue
        eff = ev["effective_date_iso"]
        meet_iso = _first_board_meeting_on_or_after_iso(co, eff)
        if not meet_iso:
            continue
        safe = sanitize_company_name(co_name)
        display = str(co.get("minutes_display_name") or co_name).strip()
        from_ln = _jurisdiction_long_name(ev["from_jurisdiction"])
        to_ln = _jurisdiction_long_name(ev["to_jurisdiction"])
        blocks = _domestication_detailed_resolution_blocks(co)
        if not blocks:
            continue
        body = "\n\n".join(blocks)
        md = (
            f"**Standalone board resolutions — {display}**\n"
            f"**Related board meeting date (cycle anchor):** {meet_iso}\n"
            f"**Subject:** Domestication / continuation ({from_ln} to {to_ln})\n\n"
            f"{body}\n\n"
            f"{SIGNATURE_BLOCK_MARKER}\n\n"
            f"{board_meeting_signature_markdown(co, meet_iso, sole_director_name='Derek E. Pappas')}\n---\n"
        )
        out_dir = os.path.join(root_dir, safe, "meetings")
        os.makedirs(out_dir, exist_ok=True)
        fname = f"{safe}_{meet_iso}_domestication_board_resolutions.docx"
        dest = os.path.join(out_dir, fname)
        write_docx_from_minutes(md, dest, meet_iso, co_name)
        print(f"Writing standalone board resolutions to {dest}")


def _corporation_statute_name(co: dict) -> str:
    """Full statute name for narrative references."""
    j = _jurisdiction(co)
    if j == "WY":
        return "Wyoming Business Corporation Act"
    return "Delaware General Corporation Law"


def _corp_law_section_ref(co: dict, section: str) -> str:
    """Short section citation; DE uses DGCL §X. WY maps template placeholders to W.S. 1977 WBCA sections used in generators."""
    j = _jurisdiction(co)
    if j == "DE":
        return f"DGCL §{section}"
    if j == "WY":
        # Title 17 Ch. 16 — map DGCL-shaped placeholders only (do not guess for unknown sections).
        wy_map = {"228": "17-16-704", "213": "17-16-707"}  # written consent without meeting; record date
        wy_sec = wy_map.get(section)
        if wy_sec:
            return f"W.S. 1977 § {wy_sec}"
    return f"the {_corporation_statute_name(co)}"


def _warn_if_non_de_company_has_delaware_snippets(co_name: str, co: dict) -> None:
    """Emit a runtime warning if `jurisdiction` is not DE but per-company strings still cite Delaware/DGCL."""
    if _jurisdiction(co) == "DE":
        return
    keys = (
        "stockholder_consent_bylaws_acknowledgment",
        "stockholder_consent_bylaws_mechanics_suffix",
        "stockholders_quorum_collective_sentence",
        "board_notice_waiver_bylaws_ref",
    )
    bad: list[str] = []
    for k in keys:
        v = co.get(k)
        if not isinstance(v, str):
            continue
        low = v.lower()
        if "dgcl" in low or "delaware general corporation law" in low:
            bad.append(k)
    if bad:
        warnings.warn(
            f"{co_name}: jurisdiction is {_jurisdiction(co)} but company_information fields still reference "
            f"Delaware/DGCL: {', '.join(bad)}. Update those strings or set jurisdiction to DE.",
            UserWarning,
            stacklevel=2,
        )


def _dgcl_section_bold_md(co: dict, section: str) -> str:
    """Bold markdown DGCL section cite for Delaware corporations; otherwise statute name only."""
    if _jurisdiction(co) == "DE":
        return f"**{_corp_law_section_ref(co, section)}**"
    return f"the {_corporation_statute_name(co)}"


def reliance_standard(co: dict) -> str:
    j = _jurisdiction(co)
    if j == "DE":
        ref = _dgcl_section_bold_md(co, "141(e)")
        return (
            "In taking the actions reflected in these minutes, the Sole Director relied in good faith on information, opinions, reports, and "
            "statements—including financial and operational materials prepared for this meeting and presentations from officers of the "
            "Corporation—as to matters the Sole Director reasonably believed were within such persons’ professional or expert competence, "
            f"as contemplated by {ref}.\n"
        )
    if j == "WY":
        return (
            "In taking the actions reflected in these minutes, the Sole Director relied in good faith on information, opinions, reports, and "
            "statements—including financial and operational materials prepared for this meeting and presentations from officers of the "
            "Corporation—as to matters the Sole Director reasonably believed were within such persons’ professional or expert competence, "
            "as contemplated by **W.S. 1977 § 17-16-830** (Wyoming Business Corporation Act; standards for directors and reliance on "
            "information from officers and others reasonably believed reliable in their areas of competence).\n"
        )
    return (
        "In taking the actions reflected in these minutes, the Sole Director relied in good faith on information, opinions, reports, and "
        "statements—including financial and operational materials prepared for this meeting and presentations from officers of the "
        "Corporation—as to matters the Sole Director reasonably believed were within such persons’ professional or expert competence, "
        f"as contemplated by the {_corporation_statute_name(co)}.\n"
    )


def _minutes_boilerplate_variant_index(co_name: str, date_iso: str, meeting_kind: str, modulo: int) -> int:
    """Stable per-meeting index for rotating equivalent boilerplate phrasing (deterministic; not cryptographic)."""
    if modulo <= 0:
        return 0
    digest = hashlib.md5(f"{co_name}|{date_iso}|{meeting_kind}".encode("utf-8")).digest()
    return int.from_bytes(digest[:4], "big") % modulo


def _de_141e_reliance_variant_paragraphs(co: dict) -> list[str]:
    """Semantically similar §141(e) reliance wordings (Delaware); index rotates by meeting. Citations use `_corp_law_section_ref` (short DGCL form)."""
    ref = _dgcl_section_bold_md(co, "141(e)")
    return [
        (
            "In taking the actions reflected in these minutes, the Sole Director relied in good faith on information, opinions, reports, and "
            "statements—including financial and operational materials prepared for this meeting and presentations from officers of the "
            "Corporation—as to matters the Sole Director reasonably believed were within such persons’ professional or expert competence, "
            f"as contemplated by {ref}.\n"
        ),
        (
            "For the actions described herein, the Sole Director relied in good faith on materials and oral presentations furnished for the meeting—"
            "including financial and operating summaries from officers of the Corporation—and on other information, reports, and statements "
            f"presented as to matters within the presenters’ professional or expert competence, within the meaning of {ref}.\n"
        ),
        (
            "The Sole Director stated that, in approving the matters minuted here, he relied in good faith on officer-prepared financial and operational "
            "materials and on other information, opinions, reports, and statements reasonably believed reliable on subjects within the presenters’ "
            f"expert or professional competence, consistent with {ref}.\n"
        ),
        (
            "The Sole Director recorded reliance, in good faith, on the financial and operational materials circulated or reviewed for this meeting and "
            "on statements from officers of the Corporation on matters reasonably treated as within such persons’ professional or expert competence, "
            f"as permitted by {ref}.\n"
        ),
        (
            "After reviewing the materials on file for the meeting, the Sole Director relied in good faith on information, opinions, reports, and "
            "statements—including presentations from officers of the Corporation—on matters reasonably believed to fall within such persons’ "
            f"professional or expert competence, under {ref}.\n"
        ),
    ]


def _wy_director_reliance_variant_paragraphs() -> list[str]:
    """W.S. §17-16-830-style reliance; index rotates by meeting."""
    return [
        (
            "In taking the actions reflected in these minutes, the Sole Director relied in good faith on information, opinions, reports, and "
            "statements—including financial and operational materials prepared for this meeting and presentations from officers of the "
            "Corporation—as to matters the Sole Director reasonably believed were within such persons’ professional or expert competence, "
            "as contemplated by **W.S. 1977 § 17-16-830** (Wyoming Business Corporation Act; standards for directors and reliance on "
            "information from officers and others reasonably believed reliable in their areas of competence).\n"
        ),
        (
            "The Sole Director stated that he relied in good faith on officer-prepared financial and operational materials and on other information "
            "reasonably believed reliable on matters within the presenters’ areas of competence, consistent with **W.S. 1977 § 17-16-830** of the "
            "Wyoming Business Corporation Act.\n"
        ),
        (
            "For the resolutions adopted here, the Sole Director relied in good faith on materials furnished for the meeting and on oral and written "
            "presentations from officers, in each case on subjects treated as within such persons’ professional or expert competence, as contemplated "
            "by **W.S. 1977 § 17-16-830** (Wyoming Business Corporation Act).\n"
        ),
    ]


def board_director_reliance_paragraph(
    co: dict,
    co_name: str | None = None,
    meeting_date_iso: str | None = None,
    meeting_kind: str = "",
) -> str:
    """Director reliance paragraph after Treasurer’s Report (141(e)-style for DE, W.S. §17-16-830 for WY, etc.).

    Override with **`board_meeting_reliance_markdown`** (markdown), or `""` to omit.

    When **`meeting_date_iso`** and registry **`co_name`** are supplied, Delaware and Wyoming defaults rotate among a small
    set of equivalent phrasings so minutes are not byte-identical meeting-to-meeting. Omit those args to preserve the
    legacy single default paragraph (e.g. for tests).
    """
    raw = co.get("board_meeting_reliance_markdown")
    if raw is not None:
        s = str(raw).strip()
        return f"{s}\n" if s else ""
    if not meeting_date_iso or not co_name:
        return reliance_standard(co)
    j = _jurisdiction(co)
    if j == "DE":
        variants = _de_141e_reliance_variant_paragraphs(co)
        i = _minutes_boilerplate_variant_index(co_name, meeting_date_iso, meeting_kind or "board", len(variants))
        return variants[i]
    if j == "WY":
        variants = _wy_director_reliance_variant_paragraphs()
        i = _minutes_boilerplate_variant_index(co_name, meeting_date_iso, meeting_kind or "board", len(variants))
        return variants[i]
    return reliance_standard(co)


def _normalized_board_directors(co: dict) -> list[dict[str, str]]:
    """If set, `board_directors` is a list of dicts with `name` and optional `title` (default **Director**)."""
    raw = co.get("board_directors")
    if not raw or not isinstance(raw, list):
        return []
    out: list[dict[str, str]] = []
    for item in raw:
        if isinstance(item, str) and item.strip():
            out.append({"name": item.strip(), "title": "Director"})
        elif isinstance(item, dict):
            nm = str(item.get("name") or "").strip()
            if not nm:
                continue
            title = str(item.get("title") or "Director").strip() or "Director"
            out.append({"name": nm, "title": title})
    return out


def _board_remote_141i_clause(co: dict) -> str:
    """Explicit §141(i) tail for Delaware board remote-presence lines (audit: cite statute, not only bylaws)."""
    if _jurisdiction(co) != "DE" or not co.get("virtual_ok", True):
        return ""
    ref = _dgcl_section_bold_md(co, "141(i)")
    return f", as permitted under {ref} and the Corporation’s bylaws"


def _append_141i_to_remote_lines(co: dict, remote_lines: list[str]) -> list[str]:
    tail = _board_remote_141i_clause(co)
    if not tail:
        return remote_lines
    out: list[str] = []
    for ln in remote_lines:
        s = ln.rstrip()
        if s.endswith("."):
            s = s[:-1]
        out.append(s + tail + ".")
    return out


def _quorum_notice_remote_variant_lines(co: dict, director_name: str) -> tuple[list[str], list[str], list[str]]:
    """Returns (quorum_lines, notice_lines, remote_lines) for sole-director board meetings."""
    statute = _corporation_statute_name(co)
    quorum_lines = [
        (
            f"The Sole Director being present, a quorum was present, and the meeting was duly constituted to transact business "
            f"in accordance with the {statute}."
        ),
        (
            f"With **{director_name}** present as the Sole Director, the Board had a quorum and lawfully convened to transact business under the {statute}."
        ),
        (
            f"The Sole Director’s presence satisfied the quorum requirement, and the meeting proceeded as duly constituted under the {statute}."
        ),
        (
            f"Quorum being established by the attendance of the Sole Director, the meeting was duly organized for business under the {statute}."
        ),
    ]
    notice_lines = [
        "The Sole Director confirmed that notice of the meeting was duly given or waived.",
        "The Sole Director confirmed that **notice** had been duly given or **validly waived** for this meeting.",
        "Notice for the meeting had been provided or waived as required, which the Sole Director confirmed for the record.",
        "The Sole Director acknowledged on the record that notice requirements were satisfied by proper notice or waiver.",
    ]
    remote_lines = [
        (
            "The Sole Director participated via communications equipment by means of which all persons participating in the meeting could hear each other, "
            "and such participation constituted presence in person at the meeting."
        ),
        (
            "The Sole Director attended using remote communications by which each participant could hear the others, and treated that participation as "
            "**presence in person** where permitted by applicable law and the Corporation’s bylaws."
        ),
        (
            "Remote participation was used for the Sole Director’s attendance; the means employed allowed contemporaneous hearing among participants and "
            "were treated as satisfying any applicable **in-person** presence requirement."
        ),
        (
            "The Sole Director joined the meeting by approved digital means, with audio contemporaneous among participants, and such attendance was "
            "recorded as **present in person** at the meeting for quorum purposes."
        ),
    ]
    remote_lines = _append_141i_to_remote_lines(co, remote_lines)
    return quorum_lines, notice_lines, remote_lines


def _quorum_notice_remote_variant_lines_multi(co: dict, director_names: list[str]) -> tuple[list[str], list[str], list[str]]:
    """Quorum / notice / remote variants when the full board (two or more directors) is present."""
    statute = _corporation_statute_name(co)
    joined = " and ".join(f"**{n}**" for n in director_names)
    nlist = ", ".join(director_names)
    quorum_lines = [
        (
            f"With {joined} present—constituting the **entire** membership of the Board of Directors—a quorum was established and the meeting was "
            f"duly constituted to transact business in accordance with the {statute}. The Corporation’s governance practice is that **all directors attend each board meeting**."
        ),
        (
            f"All directors ({nlist}) were present, the Board was fully constituted, and quorum requirements under the {statute} and the Corporation’s bylaws were satisfied. "
            "Each director is expected to attend every board meeting, and that expectation was met for this meeting."
        ),
        (
            f"Attendance by {joined} satisfied quorum as the full Board; the meeting lawfully convened under the {statute}. "
            "No director was absent."
        ),
    ]
    notice_lines = [
        "The directors confirmed that notice of the meeting was duly given or waived.",
        "The directors confirmed that **notice** had been duly given or **validly waived** for this meeting.",
        "Notice for the meeting had been provided or waived as required, which the directors confirmed for the record.",
        "The directors acknowledged on the record that notice requirements were satisfied by proper notice or waiver.",
    ]
    remote_lines = [
        (
            "The directors participated via communications equipment by means of which all persons participating in the meeting could hear each other, "
            "and such participation constituted presence in person at the meeting."
        ),
        (
            "The directors attended using remote communications by which each participant could hear the others, and the Board treated that participation as "
            "**presence in person** where permitted by applicable law and the Corporation’s bylaws."
        ),
        (
            "Remote participation was used; the means employed allowed contemporaneous hearing among participants and "
            "were treated as satisfying any applicable **in-person** presence requirement for each director."
        ),
        (
            "The directors joined the meeting by approved digital means, with audio contemporaneous among participants, and such attendance was "
            "recorded as **present in person** at the meeting for quorum purposes."
        ),
    ]
    remote_lines = _append_141i_to_remote_lines(co, remote_lines)
    return quorum_lines, notice_lines, remote_lines


def board_roll_quorum_markdown_sole_director(
    co: dict,
    co_name: str,
    date_iso: str,
    meeting_kind: str,
    director_name: str = "Derek E. Pappas",
) -> str:
    """Roll call + quorum + notice + remote block; layout varies by `board_roll_quorum_layout` on the company dict."""
    q_lines, n_lines, r_lines = _quorum_notice_remote_variant_lines(co, director_name)
    salt = {"agm": 0, "special": 17, "quarterly": 31}.get(meeting_kind, 0)
    vi = _minutes_boilerplate_variant_index(co_name, date_iso, f"{meeting_kind}|roll", len(q_lines))
    vj = _minutes_boilerplate_variant_index(co_name, date_iso, f"{meeting_kind}|notice|{salt}", len(n_lines))
    vk = _minutes_boilerplate_variant_index(co_name, date_iso, f"{meeting_kind}|remote|{salt}", len(r_lines))
    quorum = q_lines[vi]
    notice = n_lines[vj]
    remote = board_remote_presence_paragraph(
        co, co_name=co_name, meeting_date_iso=date_iso, meeting_kind=f"{meeting_kind}|remotepick", remote_variant_index=vk
    ).rstrip("\n")

    layout = int(co.get("board_roll_quorum_layout", 0)) % 4
    # 0: quorum, notice, remote (default). 1: notice, quorum, remote. 2: shorter fused quorum+notice then remote. 3: single prose paragraph (RG-style).
    if layout == 1:
        mid = f"{notice}\n{quorum}"
    elif layout == 2:
        mid = f"{quorum} {notice}"
    elif layout == 3:
        absent_bit = (
            f"There were **no** additional directors and **no** absences. {quorum} {notice}"
            if meeting_kind == "agm"
            else f"{quorum} {notice}"
        )
        mid = absent_bit
    else:
        mid = f"{quorum}\n{notice}"

    if meeting_kind == "agm":
        if layout == 3:
            body = f"""**III. Roll Call and Quorum**
**Director Present:**  
{director_name} (Sole Director)

{mid}
{remote}

"""
        else:
            body = f"""**III. Roll Call and Quorum**
**Director Present:**  
{director_name} (Sole Director)

**Director Absent:**  
None

{mid}
{remote}

"""
        return body

    # special / quarterly: colon style headers, tighter lines
    if layout == 3:
        body = f"""**II. Roll Call and Quorum:**
**Director Present:** {director_name} (Sole Director)

{mid}
{remote}

"""
    else:
        body = f"""**II. Roll Call and Quorum:**
**Director Present:** {director_name} (Sole Director)  
**Director Absent:** None  

{mid}
{remote}

"""
    return body


def board_roll_quorum_markdown_multi_director(
    co: dict,
    co_name: str,
    date_iso: str,
    meeting_kind: str,
    directors: list[dict[str, str]],
) -> str:
    """Roll call + quorum when `board_directors` lists the full board (each expected at every meeting)."""
    names = [d["name"] for d in directors]
    q_lines, n_lines, r_lines = _quorum_notice_remote_variant_lines_multi(co, names)
    salt = {"agm": 0, "special": 17, "quarterly": 31}.get(meeting_kind, 0)
    vi = _minutes_boilerplate_variant_index(co_name, date_iso, f"{meeting_kind}|roll|multi", len(q_lines))
    vj = _minutes_boilerplate_variant_index(co_name, date_iso, f"{meeting_kind}|notice|multi|{salt}", len(n_lines))
    vk = _minutes_boilerplate_variant_index(co_name, date_iso, f"{meeting_kind}|remote|multi|{salt}", len(r_lines))
    quorum = q_lines[vi]
    notice = n_lines[vj]
    remote = board_remote_presence_paragraph(
        co, co_name=co_name, meeting_date_iso=date_iso, meeting_kind=f"{meeting_kind}|remotepick|multi", remote_variant_index=vk
    ).rstrip("\n")

    layout = int(co.get("board_roll_quorum_layout", 0)) % 4
    if layout == 1:
        mid = f"{notice}\n{quorum}"
    elif layout == 2:
        mid = f"{quorum} {notice}"
    elif layout == 3:
        mid = f"{quorum} {notice}"
    else:
        mid = f"{quorum}\n{notice}"

    present_lines = "\n".join(f"{d['name']} ({d.get('title') or 'Director'})" for d in directors)

    if meeting_kind == "agm":
        body = f"""**III. Roll Call and Quorum**
**Directors Present:**  
{present_lines}

**Directors Absent:**  
None

{mid}
{remote}

"""
        return body

    body = f"""**II. Roll Call and Quorum:**
**Directors Present:**  
{present_lines}  
**Directors Absent:** None  

{mid}
{remote}

"""
    return body


def board_roll_quorum_markdown(
    co: dict,
    co_name: str,
    date_iso: str,
    meeting_kind: str,
    *,
    director_name: str = "Derek E. Pappas",
) -> str:
    """Roll call + quorum; uses `board_directors` when set, otherwise sole-director template."""
    bds = _normalized_board_directors(co)
    if len(bds) >= 2:
        return board_roll_quorum_markdown_multi_director(co, co_name, date_iso, meeting_kind, bds)
    return board_roll_quorum_markdown_sole_director(co, co_name, date_iso, meeting_kind, director_name=director_name)


def board_remote_presence_paragraph(
    co: dict,
    *,
    co_name: str | None = None,
    meeting_date_iso: str | None = None,
    meeting_kind: str = "",
    remote_variant_index: int | None = None,
) -> str:
    """Sole-director remote-presence sentence in board minutes (AGM / special / quarterly).

    Default text is identical across companies when `virtual_ok` is true—fine legally, but repetitive for readers.
    Set **`board_meeting_remote_presence_markdown`** to a company-specific sentence (markdown), or `""` to omit this block.

    When **`co_name`**, **`meeting_date_iso`**, and **`meeting_kind`** are provided (and no override string), one of several
    equivalent remote-presence formulations is selected deterministically. Callers may pass **`remote_variant_index`** to
    align with the roll-call block’s pick without re-hashing.
    """
    raw = co.get("board_meeting_remote_presence_markdown")
    if raw is not None:
        s = str(raw).strip()
        return f"{s}\n" if s else ""
    if not co.get("virtual_ok", True):
        return ""
    bds = _normalized_board_directors(co)
    if len(bds) >= 2:
        names = [d["name"] for d in bds]
        _, _, r_lines = _quorum_notice_remote_variant_lines_multi(co, names)
    else:
        _, _, r_lines = _quorum_notice_remote_variant_lines(co, "Derek E. Pappas")
    if co_name and meeting_date_iso:
        if remote_variant_index is None:
            remote_variant_index = _minutes_boilerplate_variant_index(
                co_name, meeting_date_iso, f"{meeting_kind or 'board'}|remote", len(r_lines)
            )
        line = r_lines[remote_variant_index % len(r_lines)]
        return f"{line}\n"
    return f"{r_lines[0]}\n"


def board_meeting_materials_acknowledgment_block(co: dict) -> str:
    """Optional paragraph: materials the Sole Director reviewed before the meeting (e.g. exhibit index).

    Set **`board_meeting_materials_acknowledgment_markdown`** to counsel-approved text that matches **real**
    records (PDFs, emails, annexed exhibits). Omit the key or use a blank string when not used. The generator does
    not invent filenames, versions, or dates.
    """
    raw = co.get("board_meeting_materials_acknowledgment_markdown")
    if not raw:
        return ""
    s = str(raw).strip()
    if not s:
        return ""
    return f"{s}\n\n"


# Company information (canonical registry for minute generation)
#
# Prior annual board minutes (AGM § IV): for every company here **except DATA RECORD SCIENCE, INC.**, `minutes_start_year`
# equals `inc_year`—the first AGM minutes this program generates are the corporation’s first annual board cycle after incorporation.
# **DATA RECORD SCIENCE** alone predates the series (`inc_year` 2006, `minutes_start_year` 2022); its first generated AGM uses
# “compilation series begins …” wording so the minutes do not read as denying earlier corporate life.
company_information = {
    "Hippo, Inc": {
        "minutes_display_name": "Hippo, Inc.",
        "address": "30 N Gould St Ste 21106, Sheridan, WY 82801",
        "par": "$.0001",
        "inc_year": 2022,
        "incorporation_filed_date_iso": "2022-04-07",
        "incorporation_jurisdiction": "DE",
        "minutes_start_year": 2022,
        "director_election_standard": "plurality",
        "shares_issued": {2022: "8,000,000", 2023: "8,160,000", 2024: "8,160,000", 2025: "8,160,000", 2026: "8,160,000"},
        "annual_day_offset": 0,
        "meeting_stagger_day": 0,
        "stockholder_meeting": "written_consent",
        # Annex executed consent PDF to board book; minutes cross-reference (AGM). Set to None until filed.
        "sole_stockholder_consent_exhibit_label": "Exhibit A",
        "use_timeline_place": True,
        "virtual_ok": True,
        # Engineering / contractor geography for Hippo only (not used for other registry companies).
        "development_centers_line": "Serbia; Bosnia and Herzegovina; Tunisia",
        "board_directors": [
            {"name": "Derek E. Pappas", "title": "Director"},
            {"name": "Marija Cejovic", "title": "Director"},
        ],
        "board_meeting_chair_name": "Derek E. Pappas",
        "board_sole_director_first_chronological_meeting": True,
        "organizational_meeting_full_board": True,
        "board_meeting_reliance_markdown": (
            "In taking the actions reflected in these minutes, the directors relied in good faith on information, opinions, reports, and statements—including "
            "financial and operational materials prepared for this meeting and presentations from officers of the Corporation—as to matters the directors "
            "reasonably believed were within such persons’ professional or expert competence, as contemplated by **DGCL §141(e)**.\n"
        ),
        "board_roll_quorum_layout": 1,
        "minute_book_compilation_preamble_markdown": (
            "**Compiled board minutes — {display_company}**\n\n"
            "Single volume covering calendar years **{first_year}** through **{last_year}**. "
            "The Corporation may maintain additional minutes or instruments outside this span.\n\n"
            "Where these minutes reference exhibits, signed counterparts or labeled annexes may be bound with this book or filed separately.\n\n"
            "---"
        ),
        "organizational_bylaws_exhibit_label": "Exhibit A",
        "organizational_bylaws_document_description": "Amended and Restated Bylaws of Hippo, Inc.",
        "organizational_officers_elected": [
            {
                "name": "Derek E. Pappas",
                "titles": ["President", "Secretary", "Treasurer", "Chief Executive Officer"],
            },
            {"name": "Marija Cejovic", "title": "Chief Operating Officer"},
        ],
        # Signature formatting (intentionally varied per company).
        "signature_block_style": "executed_by",
        "signature_block_include_date": True,
        "signature_block_date_format": "iso",
        "signature_block_date_label": "Date:",
        "signature_block_include_title_in_label": True,
        "signature_block_spacing_lines": 1,
        "signature_block_print_signing_lines": True,
        "primary_banking_institution": "JPMorgan Chase Bank, N.A.",
        "agm_president_report_product_line": (
            "The President’s report summarized the Corporation’s consumer and data platforms for **{year}**, including the **social shopping network** marketed as "
            "**Hippo Shopping** on the Apple App Store and Google Play, the **PriceStarz** **browser-based comparison-shopping extension**, **large-scale web crawling**, "
            "**templates and extraction workflows** producing **structured data records**, and a **data-processing pipeline** for deals and product-offer records that "
            "**groups offers for the same product** and, where possible, **associates deal-offer records with the correct product group**—together with merchant integrations, "
            "public APIs, analytics, and back-office tooling supporting those systems."
        ),
        "agm_discussion_items_line": (
            "The directors discussed the Corporation’s **{next_year}** roadmap across **Hippo Shopping**, **PriceStarz**, **crawl coverage and extraction quality**, "
            "**offer-grouping and deal-linking accuracy** in the processing pipeline, merchant and marketplace reliability, and **mobile and extension release** cadence."
        ),
        "special_meeting_purpose": (
            "Pre-annual review of Hippo Shopping, PriceStarz, web crawling, structured-data extraction, offer-grouping pipeline, marketplace, and mobile operations"
        ),
        "special_meeting_ratification_resolution_markdown": (
            "**Ratification of Shopping, Extension, Crawling, Extraction, and Pipeline Operations**  \n"
            "RESOLVED, that operational and engineering decisions affecting the Corporation’s **social shopping application**, **PriceStarz extension**, **web crawling**, "
            "**structured-data extraction**, **deals and product-offer processing pipeline**, consumer marketplace surfaces, data ingestion, and **native mobile applications** "
            "during {year} are hereby ratified, confirmed, and approved in all respects."
        ),
        "treasurer_contingent_obligations_clause": (
            "including notes payable and similar obligations that remain contingent on a future liquidity event, "
            "the timing of which has not yet been determined."
        ),
        "quarterly_default_ratification_resolution": (
            "RESOLVED, that **Hippo Shopping**, **PriceStarz**, **web crawling**, **structured-data extraction**, **deals and product-offer pipeline**, marketplace, "
            "data-pipeline, **mobile application**, and supporting cloud infrastructure work completed during the quarter—and related intellectual property—is hereby "
            "ratified, confirmed, and approved as assets of the Corporation."
        ),
        "agm_ip_affirmation_sentence": (
            "All software, data models, algorithms, and related intellectual property developed **for the Corporation** during the year "
            "under applicable **contractor and consultant arrangements** (including services performed by personnel **in Serbia; Bosnia and Herzegovina; Tunisia**) "
            "were reaffirmed as **properly titled to and the exclusive property of the Corporation** under those arrangements and applicable law, "
            "without implying ownership of **counterparties’ equipment or premises**."
        ),
        # Cite filed Amended and Restated Bylaws (`bylaws_text/Hippo, Inc. - Bylaws.docx.pdf.txt`).
        "stockholder_consent_bylaws_acknowledgment": (
            "The undersigned acknowledges that this consent is intended to comply with **Article III, Section 13** "
            "of the Corporation’s **Amended and Restated Bylaws** (Action Without Meeting), as well as **Section 228** of the DGCL."
        ),
        "stockholder_consent_bylaws_mechanics_suffix": (
            "Under **Article III, Section 13(b)** of those Bylaws, written consents must be delivered to the Corporation "
            "within **sixty (60) days** of the earliest dated consent; delivery to the registered office, if used, must be "
            "by hand or by certified or registered mail, return receipt requested, as set forth in the Bylaws."
        ),
        "board_notice_waiver_bylaws_ref": (
            "the Corporation’s **Amended and Restated Bylaws**, including **Article IV, Section 21** "
            "(meetings of the Board of Directors; notice; and waiver of notice)"
        ),
    },
    "Ritual Growth, Inc.": {
        "address": "30 N Gould St Ste 27616, Sheridan, WY 82801",
        "par": "$.0001",
        "inc_year": 2022,
        "incorporation_filed_date_iso": "2022-06-10",
        "incorporation_jurisdiction": "DE",
        "prior_incorporation": {
            "jurisdiction": "WY",
            "filed_date_iso": "2021-09-14",
        },
        "domestication_event": {
            "from_jurisdiction": "WY",
            "to_jurisdiction": "DE",
            "effective_date_iso": "2022-06-10",
            "documents_exhibit_label": "Exhibit A",
        },
        "minutes_start_year": 2022,
        "shares_issued": {2022: "4,000,000", 2023: "4,000,000", 2024: "4,000,000", 2025: "4,000,000", 2026: "4,000,000"},
        "annual_day_offset": 1,
        "meeting_stagger_day": 1,
        "stockholder_meeting": "written_consent",
        "sole_stockholder_consent_exhibit_label": "Exhibit A",
        "use_timeline_place": True,
        "virtual_ok": True,
        # `bylaws_text/Bylaws of Ritual Growth, Inc.pdf.txt` — Roman articles II / III / VIII.
        "stockholder_consent_bylaws_acknowledgment": (
            "The undersigned acknowledges that this consent is intended to comply with **Article II, Section 9** "
            "of the **By-Laws of Ritual Growth, Inc.** (Action Without Meeting), as well as **Section 228** of the DGCL."
        ),
        "stockholder_consent_bylaws_mechanics_suffix": (
            "Under **Article II, Section 9** of those By-Laws, delivery to the Corporation’s registered office shall be "
            "by hand or by certified or registered mail, return receipt requested, where that delivery method is used; "
            "**prompt notice** of action taken without a meeting is required when consent is less than unanimous among "
            "applicable stockholders, as described in the By-Laws."
        ),
        "board_notice_waiver_bylaws_ref": (
            "the **By-Laws of Ritual Growth, Inc.**, including **Article III, Section 8** (notice and place of meetings of "
            "the Board of Directors) and **Article VIII, Section 4** (waiver of notice)"
        ),
        "agm_president_report_product_line": (
            "The directors summarized continued **server-side and web application** engineering for **{year}**, including API services, deployment automation, "
            "and hosting integrations. The President’s report noted that **core product-development efforts were restarted in calendar years 2023, 2025, and 2026**, "
            "each time replanning architecture, backlog, and release sequencing for the Corporation’s primary web stack while preserving continuity for long-running services."
        ),
        "agm_president_report_infrastructure_line": (
            "The Corporation continued to operate hardware and cloud infrastructure using hosting providers including **DigitalOcean** and **Hetzner**."
        ),
        "agm_president_report_operating_exhibit_label": "Exhibit B",
        "development_centers_line": "Poland; Romania; Portugal",
        "primary_banking_institution": "Bank of America, N.A.",
        "agm_discussion_items_line": (
            "The directors discussed **{next_year}** delivery priorities for the Corporation’s **server-side and web stack**, including recovery planning after "
            "**development restarts in 2023, 2025, and 2026**, partner commitments, hosting footprint, and release scheduling."
        ),
        "special_meeting_purpose": (
            "Pre-annual review of server-side and web application delivery, hosting footprint, partner commitments, and post-restart engineering stability"
        ),
        "special_meeting_ratification_resolution_markdown": (
            "**Ratification of Server, Web, and Hosting Operations**  \n"
            "RESOLVED, that engineering, hosting, and go-to-market decisions affecting the Corporation’s **server-side services**, **web applications**, and **API** "
            "delivery during {year} are hereby ratified, confirmed, and approved in all respects."
        ),
        "treasurer_contingent_obligations_clause": (
            "including convertible instruments and vendor payment terms that remain contingent on a future liquidity event, "
            "the timing of which has not yet been determined."
        ),
        "quarterly_default_ratification_resolution": (
            "RESOLVED, that all **server-side**, **web**, **API**, and supporting infrastructure work completed during the quarter—and related intellectual property—is hereby "
            "ratified, confirmed, and approved as assets of the Corporation."
        ),
        "agm_ip_affirmation_sentence": (
            "All application code, APIs, and related intellectual property developed **for the Corporation** during the year under applicable "
            "**contractor and consultant arrangements** were reaffirmed as **properly titled to and the exclusive property of the Corporation**, "
            "without implying ownership of **counterparties’ equipment or premises**."
        ),
        "board_directors": [
            {"name": "Derek E. Pappas", "title": "Director"},
            {"name": "Marija Cejovic", "title": "Director"},
        ],
        "board_meeting_chair_name": "Derek E. Pappas",
        "board_sole_director_first_chronological_meeting": True,
        "organizational_meeting_full_board": True,
        "board_meeting_reliance_markdown": (
            "In taking the actions reflected in these minutes, the directors relied in good faith on information, opinions, reports, and statements—including "
            "financial and operational materials prepared for this meeting and presentations from officers of the Corporation—as to matters the directors "
            "reasonably believed were within such persons’ professional or expert competence, as contemplated by **DGCL §141(e)**.\n"
        ),
        "board_roll_quorum_layout": 3,
        "organizational_bylaws_exhibit_label": "Exhibit B",
        "organizational_bylaws_document_description": "By-Laws of Ritual Growth, Inc.",
        "organizational_officers_elected": [
            {
                "name": "Derek E. Pappas",
                "titles": ["President", "Secretary", "Treasurer", "Chief Executive Officer"],
            },
            {"name": "Marija Cejovic", "title": "Chief Operating Officer"},
        ],
        "minute_book_compilation_preamble_markdown": (
            "**{display_company} — board minute compilation**\n\n"
            "Years **{first_year}**–**{last_year}** (inclusive). This compilation is not represented as exhaustive of all corporate acts.\n\n"
            "Exhibits are referenced only where the underlying minutes call for them; physical execution copies may be cross-filed.\n\n"
            "---"
        ),
        "signature_block_style": "signature",
        "signature_block_include_date": True,
        "signature_block_date_format": "long",
        "signature_block_date_label": "Dated:",
        "signature_block_spacing_lines": 1,
        "signature_block_print_signing_lines": True,
    },
    "DATA RECORD SCIENCE, INC.": {
        "address": "30 N Gould St Ste 24165, Sheridan, WY 82801",
        "par": "$0.001",
        # Originally incorporated in Delaware in 2006 (as Yoterra, Inc.), later renamed to Data Record Science, Inc.
        # Sole registry company with minutes_start_year > inc_year: compiled board minutes begin in 2022 (see § IV wording in generate_agm).
        "inc_year": 2006,
        "minutes_start_year": 2022,
        "director_election_standard": "plurality",
        "shares_issued": {2022: "5,346,132", 2023: "5,346,132", 2024: "5,346,132", 2025: "5,346,132", 2026: "5,346,132"},
        "annual_day_offset": 2,
        "meeting_stagger_day": 2,
        "stockholder_meeting": "annual_meeting_stockholders",
        # Annual stockholder meeting §222: for a single voting stockholder, waiver-only path is usually simplest; use
        # "notice_focus" if counsel prefers formal notice. Set annual_stockholder_notice_exhibit_label when you annex PDFs.
        "annual_stockholder_notice_record": "waiver_focus",
        "annual_stockholder_notice_exhibit_label": "Exhibit A",
        "use_timeline_place": True,
        "virtual_ok": True,
        "stockholders_roll_call": [
            {"name": "Derek E. Pappas", "presence": "present in person"},
        ],
        "annual_stockholder_director_election_votes": [
            {"name": "Derek E. Pappas", "shares": "5,346,132", "vote": "FOR"},
        ],
        # Majority voting stockholder only: quorum satisfied by his presence alone (no other voters present).
        "stockholders_quorum_collective_sentence": (
            "**Derek E. Pappas** holds a **majority** of the outstanding shares of the Corporation entitled to vote at the meeting; "
            "he was the **only stockholder present** (in person or by proxy) entitled to vote at the meeting, and his presence "
            "**alone satisfied** the quorum requirement under the DGCL and the Corporation’s bylaws."
        ),
        # Patent holding company: no product-engineering geography list (quarterly minutes use a custom review paragraph).
        "development_centers_line": "",
        "board_roll_quorum_layout": 0,
        "agm_president_report_opening_paragraph_markdown": (
            "The President reported that the Corporation is operated principally as a **Delaware patent holding company**, "
            "holding patents and related rights without conducting commercial software engineering as an operating business. "
            "Corporate administration, outside counsel coordination, and portfolio maintenance were overseen from **{office_locations}**; "
            "the Corporation did not maintain third-party product-development centers comparable to a commercial SaaS operator during the year."
        ),
        "agm_president_report_product_line": (
            "The Sole Director summarized patent annuity and prosecution updates, portfolio housekeeping, and counsel reporting for the year. "
        ),
        "agm_discussion_items_line": (
            "The Sole Director discussed portfolio strategy for {next_year}, including claim scope reviews, continuation strategy, "
            "and coordination with outside patent counsel."
        ),
        "minute_book_compilation_preamble_markdown": (
            "**Minute book — {display_company}**\n\n"
            "Board and stockholder materials generated for **{first_year}** through **{last_year}**. "
            "This volume is limited to that span; the Corporation’s other records may include additional instruments.\n\n"
            "References to exhibits in these minutes describe records maintained for the Corporation; physical counterparts may be filed separately.\n\n"
            "---"
        ),
        "signature_block_style": "none",
        "signature_block_include_date": True,
        "signature_block_date_format": "iso",
        "signature_block_date_label": "Date:",
        "signature_block_name_prefix": "Executed by:",
        "signature_block_spacing_lines": 1,
        "signature_block_print_signing_lines": True,
        "quarterly_business_review_minutes_markdown": (
            "The Sole Director reviewed quarterly **franchise tax**, **registered agent**, and **minute-book** compliance, and confirmed that "
            "the Corporation’s **patent portfolio** records remained current as reported by counsel. The Corporation did not operate product "
            "engineering centers during **{quarter} {year}**; intellectual-property assets continued to be held at the parent level."
        ),
        "primary_banking_institution": "Wells Fargo Bank, N.A.",
        "special_meeting_purpose": "Pre-annual review of corporate records, patent portfolio status, and registered-agent compliance",
        "special_meeting_ratification_resolution_markdown": (
            "**Ratification of Corporate and Patent-Portfolio Administration**  \n"
            "RESOLVED, that corporate housekeeping, outside-counsel coordination, and patent-portfolio maintenance decisions affecting the Corporation during {year} "
            "are hereby ratified, confirmed, and approved in all respects."
        ),
        "treasurer_contingent_obligations_clause": (
            "including legacy acquisition-related holdbacks and similar obligations that remain contingent on a future liquidity event, "
            "the timing of which has not yet been determined."
        ),
        "quarterly_default_ratification_resolution": (
            "RESOLVED, that all corporate, patent-portfolio, and minute-book administration actions taken during the quarter—and related intellectual property records—"
            "are hereby ratified, confirmed, and approved as reflected in the Corporation’s books."
        ),
        "agm_ip_affirmation_sentence": (
            "Patents, patent applications, and related intellectual property held for investment were reaffirmed as assets of the Corporation."
        ),
    },
    "TeamBoost.ai, Inc.": {
        "address": "30 N Gould St Ste 23049, Sheridan, WY 82801",
        "par": "$.0001",
        # Filed January 30, 2023 (Delaware).
        "inc_year": 2023,
        "incorporation_filed_date_iso": "2023-01-13",
        "incorporation_jurisdiction": "DE",
        "minutes_start_year": 2023,
        "shares_issued": {2023: "10,000,000", 2024: "10,000,000", 2025: "10,000,000", 2026: "10,000,000"},
        "annual_day_offset": 3,
        "meeting_stagger_day": 4,
        "stockholder_meeting": "written_consent",
        "sole_stockholder_consent_exhibit_label": "Exhibit A",
        "use_timeline_place": True,
        "virtual_ok": True,
        # Same article/section map as Ritual (`bylaws_text/Bylaws of TeamBoost.ai, Inc.pdf.txt`).
        "stockholder_consent_bylaws_acknowledgment": (
            "The undersigned acknowledges that this consent is intended to comply with **Article II, Section 9** "
            "of the **By-Laws of TeamBoost.ai, Inc.** (Action Without Meeting), as well as **Section 228** of the DGCL."
        ),
        "stockholder_consent_bylaws_mechanics_suffix": (
            "Under **Article II, Section 9** of those By-Laws, delivery to the Corporation’s registered office shall be "
            "by hand or by certified or registered mail, return receipt requested, where that delivery method is used; "
            "**prompt notice** of action taken without a meeting is required when consent is less than unanimous among "
            "applicable stockholders, as described in the By-Laws."
        ),
        "board_notice_waiver_bylaws_ref": (
            "the **By-Laws of TeamBoost.ai, Inc.**, including **Article III, Section 8** (notice and place of meetings of "
            "the Board of Directors) and **Article VIII, Section 4** (waiver of notice)"
        ),
        "agm_president_report_product_line": {
            "2023": (
                "The directors summarized product development for **{year}** with primary emphasis on **native mobile applications** and the Corporation’s "
                "**AI-assisted company-management** direction—supported by **APIs, AI bots, background services, and owned or leased servers**—while **customer-facing "
                "web application** engineering moved from planning into initial implementation ahead of broader **2024** web releases."
            ),
            "default": (
                "The directors summarized continued work on an **AI platform for company management**, including the **customer-facing web application** "
                "(scaling from **2024** alongside **native mobile clients** in production from **2023**), **AI bots and AI tooling**, **server-side services**, and "
                "supporting cloud infrastructure for those products."
            ),
        },
        "agm_president_report_infrastructure_line": (
            "The Corporation continued to operate hardware and cloud infrastructure using hosting providers including **DigitalOcean** and **Hetzner**."
        ),
        "agm_president_report_operating_exhibit_label": "Exhibit B",
        "development_centers_line": "Brazil; Nigeria; Portugal",
        "quarterly_meeting_time": "2:00 PM",
        "primary_banking_institution": "First Citizens Bank & Trust Company",
        "agm_discussion_items_line": (
            "The directors discussed **{next_year}** priorities for the Corporation’s **AI-led company-management system**, including **web and mobile** delivery, "
            "**AI bot** reliability, **server and toolchains** scaling, model governance, and customer onboarding."
        ),
        "special_meeting_purpose": (
            "Pre-annual review of AI company-management platform delivery—web application, native mobile clients, AI bots, servers, and AI tooling"
        ),
        "special_meeting_ratification_resolution_markdown": (
            "**Ratification of AI, Web, Mobile, Bot, and Server Operations**  \n"
            "RESOLVED, that engineering and product decisions affecting the Corporation’s **AI-assisted company-management** offerings—including **web applications**, "
            "**native mobile applications**, **AI bots**, **servers**, **AI tools**, and **APIs**—during {year} are hereby ratified, confirmed, and approved in all respects."
        ),
        "treasurer_contingent_obligations_clause": (
            "including SAFEs, convertible notes, and similar instruments that remain contingent on a future liquidity event, "
            "the timing of which has not yet been determined."
        ),
        "quarterly_default_ratification_resolution": (
            "RESOLVED, that all **web**, **native mobile**, **AI bot**, **server**, **AI tooling**, **API**, and supporting cloud infrastructure work completed during the "
            "quarter—and related intellectual property—is hereby ratified, confirmed, and approved as assets of the Corporation."
        ),
        "agm_ip_affirmation_sentence": (
            "All web and mobile clients, **AI bots**, **AI tools**, server-side services, APIs, and related intellectual property developed **for the Corporation** during the year "
            "under applicable **contractor and consultant arrangements** were reaffirmed as **properly titled to and the exclusive property of the Corporation**, "
            "without implying ownership of **counterparties’ equipment or premises**."
        ),
        "board_directors": [
            {"name": "Derek E. Pappas", "title": "Director"},
            {"name": "Marija Cejovic", "title": "Director"},
        ],
        "board_meeting_chair_name": "Derek E. Pappas",
        "board_sole_director_first_chronological_meeting": True,
        "organizational_meeting_full_board": True,
        "board_meeting_reliance_markdown": (
            "In taking the actions reflected in these minutes, the directors relied in good faith on information, opinions, reports, and statements—including "
            "financial and operational materials prepared for this meeting and presentations from officers of the Corporation—as to matters the directors "
            "reasonably believed were within such persons’ professional or expert competence, as contemplated by **DGCL §141(e)**.\n"
        ),
        "board_roll_quorum_layout": 2,
        "organizational_bylaws_exhibit_label": "Exhibit A",
        "organizational_bylaws_document_description": "By-Laws of TeamBoost.ai, Inc.",
        "organizational_officers_elected": [
            {
                "name": "Derek E. Pappas",
                "titles": ["President", "Secretary", "Treasurer", "Chief Executive Officer"],
            },
            {"name": "Marija Cejovic", "title": "Chief Operating Officer"},
        ],
        "minute_book_compilation_preamble_markdown": (
            "**{display_company} — compiled minutes**\n\n"
            "Board governance meetings for **{first_year}** through **{last_year}**.\n\n"
            "Exhibit references in the minutes are to instruments on the corporate record; execution copies may be bound or filed elsewhere.\n\n"
            "---"
        ),
        "signature_block_style": "executed_by",
        "signature_block_include_date": False,
        "signature_block_include_title_in_label": False,
        "signature_block_label_template": "**Signed:**",
        "signature_block_spacing_lines": 1,
        "signature_block_print_signing_lines": True,
    },
    "SurveyTeams, Inc.": {
        "minutes_display_name": "SurveyTeams, Inc.",
        # Lease address (also used as principal address in minutes unless you provide a separate mailing/principal line).
        "address": "30 N Gould St #58611, Sheridan, WY 82801, USA",
        "par": "$.0001",
        # Inc/year not provided; defaulting minutes generation to start in 2026.
        "inc_year": 2026,
        "minutes_start_year": 2026,
        "director_election_standard": "plurality",
        # Two-member board: both directors are expected at every board meeting; minutes / quorum reflect full-board attendance.
        "board_directors": [
            {"name": "Derek E. Pappas", "title": "Director"},
            {"name": "Mohamed Mohamed", "title": "Director"},
        ],
        "board_meeting_chair_name": "Derek E. Pappas",
        "board_meeting_reliance_markdown": (
            "In taking the actions reflected in these minutes, the directors relied in good faith on information, opinions, reports, and statements—including "
            "financial and operational materials prepared for this meeting and presentations from officers of the Corporation—as to matters the directors "
            "reasonably believed were within such persons’ professional or expert competence, as contemplated by **DGCL §141(e)**.\n"
        ),
        # Share data per intake:
        # - 10,000,000 authorized
        # - 8,000,000 issued/outstanding (50% / 50%: 4,000,000 Derek; 4,000,000 Mohamed)
        "shares_authorized": "10,000,000",
        "shares_issued": {2026: "8,000,000"},
        "stockholder_shares": {
            "Derek E. Pappas": "4,000,000",
            "Mohamed Mohamed": "4,000,000",
        },
        "annual_day_offset": 4,
        "meeting_stagger_day": 4,
        # Two named stockholders: generate formal annual stockholder meeting minutes + notice/waiver instruments.
        "stockholder_meeting": "annual_meeting_stockholders",
        "annual_stockholder_notice_record": "waiver_focus",
        "annual_stockholder_notice_exhibit_label": "Exhibit A",
        "use_timeline_place": True,
        "virtual_ok": True,
        "stockholders_roll_call": [
            {"name": "Mohamed Mohamed", "presence": "present in person"},
            {"name": "Derek E. Pappas", "presence": "present in person"},
        ],
        "annual_stockholder_director_election_votes": [
            {"name": "Mohamed Mohamed", "shares": "4,000,000", "vote": "FOR"},
            {"name": "Derek E. Pappas", "shares": "4,000,000", "vote": "FOR"},
        ],
        "stockholders_quorum_collective_sentence": (
            "Collectively, the stockholders present at the meeting held a **majority** of the outstanding shares of the Corporation "
            "entitled to vote at the meeting, and their presence satisfied the quorum requirement under the DGCL and the Corporation’s bylaws."
        ),
        # Optional metadata captured from intake (not currently used by templates).
        "irs_ein": "41-3602747",
        "irs_legal_name": "SURVEYTEAMS INC",
        "officers": {"CEO": "Mohamed Mohamed", "CTO": "Derek E. Pappas"},
        "agm_banking_authorized_signatory": "Mohamed Mohamed",
        "agm_president_report_opening_paragraph_markdown": (
            "**Mohamed Mohamed**, President, reported on the Corporation’s operational and engineering activities for the calendar year, "
            "including centralized management of globally distributed development and the use of operational office location(s) "
            "during the calendar year, with operations conducted from {office_locations} and engineering support sourced through "
            "**contractor and consultant arrangements** with personnel **in {dev_locations}**, "
            "while confirming that management, oversight, and decision-making remained centralized and continuously recorded "
            "through the Corporation’s official records. "
        ),
        "agm_president_report_product_line": (
            "The directors summarized continued engineering of **SurveyTeams’ server-backed web application** for designing, distributing, and analyzing "
            "**hierarchical organizational surveys**, including authenticated multi-tenant access, survey versioning, **server-side** analytics and exports, "
            "and supporting APIs and operational tooling for research teams."
        ),
        "development_centers_line": "Egypt; United Arab Emirates; United States (distributed)",
        "primary_banking_institution": "Mercury Bank",
        "agm_discussion_items_line": (
            "The directors discussed **{next_year}** roadmap items for **hierarchical organizational surveys**, including **web and server** reliability, "
            "respondent privacy controls, sample-weighting methodology, multilingual instrument delivery, and enterprise deployment workflows."
        ),
        "special_meeting_purpose": (
            "Pre-annual review of SurveyTeams **server and web application** operations, hierarchical survey workflows, and research compliance"
        ),
        "special_meeting_ratification_resolution_markdown": (
            "**Ratification of Survey Server, Web Application, and Research Operations**  \n"
            "RESOLVED, that engineering and operations decisions affecting the Corporation’s **survey servers**, **web application**, hierarchical survey tooling, "
            "sampling workflows, and research compliance posture during {year} are hereby ratified, confirmed, and approved in all respects."
        ),
        "treasurer_contingent_obligations_clause": (
            "including founder advances, deferred vendor invoices, and similar items that remain contingent on a future liquidity event, "
            "the timing of which has not yet been determined."
        ),
        "quarterly_default_ratification_resolution": (
            "RESOLVED, that all **survey server**, **web application**, survey-tooling, data-collection, and supporting infrastructure work completed during the quarter—and "
            "related intellectual property—is hereby ratified, confirmed, and approved as assets of the Corporation."
        ),
        "agm_ip_affirmation_sentence": (
            "All survey instruments, weighting libraries, **web and server application code**, APIs, and related intellectual property developed **for the Corporation** during the year "
            "under applicable **contractor and consultant arrangements** were reaffirmed as **properly titled to and the exclusive property of the Corporation**, "
            "without implying ownership of **counterparties’ equipment or premises**."
        ),
        "minute_book_compilation_preamble_markdown": (
            "**SurveyTeams board minutes — compiled**\n\n"
            "**{display_company}** · years **{first_year}**–**{last_year}**.\n\n"
            "This file collects generated minutes for the stated years only; other corporate instruments may exist.\n\n"
            "---"
        ),
        "signature_block_style": "executed_by",
        "signature_block_include_date": True,
        "signature_block_date_format": "long",
        "signature_block_date_label": "Date:",
        "signature_block_include_title_in_label": True,
        "signature_block_spacing_lines": 2,
        "signature_block_print_signing_lines": True,
    },
    "Loki Sports Enterprises, Inc.": {
        "minutes_display_name": "Loki Sports Enterprises, Inc.",
        "address": "30 N Gould St Ste 24709, Sheridan, WY 82801",
        "jurisdiction": "WY",
        "board_meeting_chair_name": "Derek E. Pappas",
        "par": "$.0001",
        # WY domestic profit corporation filing (2023).
        "inc_year": 2023,
        "minutes_start_year": 2023,
        "shares_issued": {
            2023: "10,000,000",
            2024: "10,000,000",
            2025: "10,000,000",
            2026: "10,000,000",
        },
        "annual_day_offset": 5,
        "meeting_stagger_day": 5,
        # Sole stockholder: use written consent pack (DGCL-style templates; adjust if you later add WY-specific variants).
        "stockholder_meeting": "written_consent",
        "sole_stockholder_consent_exhibit_label": "Exhibit A",
        "use_timeline_place": True,
        "virtual_ok": True,
        # Optional metadata captured from intake (not currently used by templates).
        "irs_ein": "93-2976555",
        "irs_legal_name": "LOKI SPORTS ENTERPRISES INC",
        "wy_sos_filing_id": "2023-001316332",
        "dba": "DEREK EDWIN PAPPAS",
        "mailing_address": "1317 Edgewater Dr Num 1961, Orlando, FL 32804",
        "board_roll_quorum_layout": 2,
        "development_centers_line": "United States; Pakistan",
        "agm_president_report_product_line": (
            "The Sole Director summarized sports-media and venue-integration work during the year, including coordination with domestic (United States) "
            "operations and **hockey stick manufacturing** run through **Pakistan**-based production partners, with inventory and quality oversight centralized through the Corporation. "
        ),
        "minute_book_compilation_preamble_markdown": (
            "**Wyoming corporation — compiled board minutes**\n\n"
            "**{display_company}** ({first_year}–{last_year}).\n\n"
            "Wyoming-law references in these minutes follow the Wyoming Business Corporation Act. Exhibit cross-references are to records on file for the Corporation.\n\n"
            "---"
        ),
        "signature_block_style": "executed_by",
        "signature_block_include_date": True,
        "signature_block_date_format": "iso",
        "signature_block_date_label": "Date:",
        "signature_block_include_title_in_label": False,
        "signature_block_label_template": "**Executed and agreed:**",
        "signature_block_spacing_lines": 1,
        "signature_block_print_signing_lines": True,
        "primary_banking_institution": "Truist Bank",
        "agm_discussion_items_line": (
            "The Sole Director discussed fan engagement, venue partnerships, and media integrations for {next_year}, including tournament-season logistics, "
            "rights-holder coordination, and in-venue mobile experiences."
        ),
        "special_meeting_purpose": "Pre-annual review of sports media products and venue-facing integrations",
        "special_meeting_ratification_resolution_markdown": (
            "**Ratification of Sports Media and Venue Integration Work**  \n"
            "RESOLVED, that engineering and commercial decisions affecting the Corporation’s sports media stack and venue-facing integrations during {year} "
            "are hereby ratified, confirmed, and approved in all respects."
        ),
        "treasurer_contingent_obligations_clause": (
            "including equipment financing and season-cycle receivables that remain contingent on a future liquidity event, "
            "the timing of which has not yet been determined."
        ),
        "quarterly_default_ratification_resolution": (
            "RESOLVED, that all sports-media, venue-integration, and supporting infrastructure work completed during the quarter—and related intellectual property—is hereby "
            "ratified, confirmed, and approved as assets of the Corporation."
        ),
        "agm_ip_affirmation_sentence": (
            "All sports-media software, venue-integration tooling, product designs, manufacturing specifications for hockey sticks produced through the Corporation’s "
            "supply chain, and related intellectual property developed **for the Corporation** during the year were reaffirmed as **properly titled to and the exclusive property of the Corporation**; "
            "the Board noted that **production partners’ facilities and equipment are not owned by the Corporation**."
        ),
    },
}

# Backwards-compatible alias (existing generator code expects `companies`).
companies = company_information


def _co_registry_key_for(co: dict) -> str | None:
    """Resolve the dict key in `companies` for this company config (identity match)."""
    for k, v in companies.items():
        if v is co:
            return k
    return None


def schedule_seed_from_environment() -> int | None:
    """Integer seed from `CORPORATE_MINUTES_SCHEDULE_SEED`, or None if unset (no schedule randomization)."""
    raw = os.environ.get("CORPORATE_MINUTES_SCHEDULE_SEED", "").strip()
    if not raw:
        return None
    try:
        return int(raw, 0)
    except ValueError:
        return int(hashlib.sha256(raw.encode("utf-8")).hexdigest()[:12], 16)


def _schedule_rng_for_co(co: dict, year: int, salt: str) -> random.Random | None:
    base = schedule_seed_from_environment()
    if base is None:
        return None
    key = _co_registry_key_for(co)
    if not key:
        return None
    suf = str(co.get("schedule_seed_suffix", ""))
    digest = hashlib.blake2b(f"{base}|{key}|{year}|{salt}|{suf}".encode("utf-8"), digest_size=8).digest()
    return random.Random(int.from_bytes(digest, "big"))


def _add_signed_weekdays(d: date, n: int) -> date:
    """Move `n` weekdays (Mon–Fri) forward (n>0) or backward (n<0), skipping weekends."""
    if n == 0:
        return d
    step = 1 if n > 0 else -1
    out = d
    remaining = abs(n)
    while remaining:
        out += timedelta(days=step)
        if out.weekday() < 5:
            remaining -= 1
    return out


def _clamp_date_to_december(d: date, year: int) -> date:
    lo, hi = date(year, 12, 1), date(year, 12, 31)
    if d < lo:
        return lo
    if d > hi:
        return hi
    return d


def _clock_to_minutes_since_midnight(clock: str) -> int:
    t = datetime.strptime(clock.strip(), "%I:%M %p")
    return t.hour * 60 + t.minute


def _minutes_since_midnight_to_ampm(ms: int) -> str:
    ms = max(8 * 60, min(18 * 60 - 1, ms))
    hh, mm = divmod(ms, 60)
    disp = hh % 12
    if disp == 0:
        disp = 12
    return f"{disp}:{mm:02d} {'PM' if hh >= 12 else 'AM'}"


def _raw_scheduled_meeting_time(co: dict, year: int, salt: str, nominal_clock: str) -> str:
    """Nominal time like `1:00 PM`, optionally jittered when seed + `schedule_time_jitter_minutes` are set (no same-day ordering)."""
    max_j = int(co.get("schedule_time_jitter_minutes", 0) or 0)
    rng = _schedule_rng_for_co(co, year, salt)
    if rng is None or max_j <= 0:
        return nominal_clock
    base = _clock_to_minutes_since_midnight(nominal_clock)
    delta = rng.randint(-max_j, max_j)
    step = int(co.get("schedule_time_round_minutes", 5) or 5)
    if step > 0:
        delta = (delta // step) * step
    return _minutes_since_midnight_to_ampm(base + delta)


def scheduled_meeting_time(co: dict, year: int, salt: str, nominal_clock: str) -> str:
    """Raw jittered clock (no cross-meeting ordering). Prefer `scheduled_*` wrappers, which enforce same-day sequencing."""
    return _raw_scheduled_meeting_time(co, year, salt, nominal_clock)


def _apply_ordered_start_minutes(raw_minutes: list[int], gap_minutes: int) -> list[int]:
    """Enforce ordered start times; if gap_minutes > 0, each start is at least gap_minutes after the previous start."""
    out: list[int] = []
    prev_start: int | None = None
    for m in raw_minutes:
        if prev_start is not None:
            if gap_minutes <= 0:
                if m < prev_start:
                    m = prev_start
            else:
                lo = prev_start + gap_minutes
                if m < lo:
                    m = lo
        m = max(8 * 60, min(18 * 60 - 1, m))
        out.append(m)
        prev_start = m
    return out


def _meeting_clocks_for_year(co: dict, year: int) -> dict[str, str]:
    """All jittered clocks for `year`, with same-calendar-day meetings ordered and non-overlapping (sole-director realistic)."""
    ann = annual_meeting_date_str(co, year)
    sp = board_special_meeting_date_str(co, year)
    raw_special = _raw_scheduled_meeting_time(co, year, "special_board", SPECIAL_MEETING_TIME)
    raw_stock = _raw_scheduled_meeting_time(co, year, "stockholder", STOCKHOLDER_MEETING_TIME)
    raw_board = _raw_scheduled_meeting_time(co, year, "board_agm", BOARD_AGM_TIME)
    raw_q: dict[str, str] = {
        q: _raw_scheduled_meeting_time(co, year, f"quarterly_{q}", QUARTERLY_MEETING_TIME)
        for q in ("Q1", "Q2", "Q3", "Q4")
    }
    gap_written = int(co.get("schedule_same_day_gap_minutes", 45) or 45)
    gap_stock_board = int(co.get("schedule_stockholder_to_board_gap_minutes", 0) or 0)

    result: dict[str, str] = {
        "special_board": raw_special,
        "stockholder": raw_stock,
        "board_agm": raw_board,
        **{f"quarterly_{q}": raw_q[q] for q in raw_q},
    }

    by_date: dict[str, list[tuple[int, str, str]]] = defaultdict(list)
    # tuple: sort order, result_key, raw_clock_str

    if co.get("stockholder_meeting") == "annual_meeting_stockholders":
        by_date[sp].append((0, "special_board", raw_special))
        by_date[ann].append((1, "stockholder", raw_stock))
        by_date[ann].append((2, "board_agm", raw_board))
    else:
        by_date[ann].append((0, "special_board", raw_special))
        by_date[ann].append((2, "board_agm", raw_board))

    for qi, q in enumerate(("Q1", "Q2", "Q3", "Q4")):
        qd = quarterly_meeting_date_str(co, year, q)
        by_date[qd].append((10 + qi, f"quarterly_{q}", raw_q[q]))

    for d_iso in sorted(by_date.keys()):
        events = sorted(by_date[d_iso], key=lambda t: t[0])
        keys = {e[1] for e in events}
        if len(events) <= 1:
            continue
        if keys == {"stockholder", "board_agm"}:
            gap_use = gap_stock_board
        elif "special_board" in keys and "board_agm" in keys and len(keys) == 2:
            gap_use = gap_written
        else:
            gap_use = gap_written if len(keys) > 1 else 0
        raw_m = [_clock_to_minutes_since_midnight(e[2]) for e in events]
        adj_m = _apply_ordered_start_minutes(raw_m, gap_use)
        for (_, key, _), ms in zip(events, adj_m, strict=True):
            result[key] = _minutes_since_midnight_to_ampm(ms)

    return result


def scheduled_stockholder_meeting_time(co: dict, year: int) -> str:
    return _meeting_clocks_for_year(co, year)["stockholder"]


def scheduled_board_agm_time(co: dict, year: int) -> str:
    return _meeting_clocks_for_year(co, year)["board_agm"]


def scheduled_special_meeting_time(co: dict, year: int) -> str:
    return _meeting_clocks_for_year(co, year)["special_board"]


def scheduled_quarterly_meeting_time(co: dict, year: int, quarter: str) -> str:
    # Allow per-company nominal time override (still jittered/ordered if schedule seed is enabled).
    nominal = str(co.get("quarterly_meeting_time") or QUARTERLY_MEETING_TIME).strip() or QUARTERLY_MEETING_TIME
    # Recompute clocks with the nominal override by temporarily shadowing the constant for this calculation.
    # (We keep the seed salt stable; only the baseline clock changes.)
    clocks = _meeting_clocks_for_year(co, year)
    if nominal == QUARTERLY_MEETING_TIME:
        return clocks[f"quarterly_{quarter}"]
    return _raw_scheduled_meeting_time(co, year, f"quarterly_{quarter}", nominal)


# Accomplishments (President’s report summary + operating addendum detail) — `audit_reports/all_corp_accomplishments_2021-2025.json`.
# Top-level keys in JSON: "Hippo", "TB", "RG", … mapped from `companies` dict keys below.
_ACCOMPLISHMENTS_JSON_KEY: dict[str, str] = {
    "Hippo, Inc": "Hippo",
    "TeamBoost.ai, Inc.": "TB",
    "Ritual Growth, Inc.": "RG",
    "Loki Sports Enterprises, Inc.": "LOKI",
}
_ACCOMPLISHMENTS_PATH = os.path.join(
    os.path.dirname(os.path.abspath(__file__)),
    "audit_reports",
    "all_corp_accomplishments_2021-2025.json",
)
_accomplishments_cache: dict | None = None

_MARIJA_VACATIONS_JSON = os.path.join(
    os.path.dirname(os.path.abspath(__file__)),
    "vacations",
    "marija_cejovic.json",
)
_marija_vacation_ranges_cache: list[tuple[date, date]] | None = None


def _load_marija_vacation_ranges() -> list[tuple[date, date]]:
    """Inclusive date ranges from `vacations/marija_cejovic.json` (empty if missing)."""
    global _marija_vacation_ranges_cache
    if _marija_vacation_ranges_cache is not None:
        return _marija_vacation_ranges_cache
    out: list[tuple[date, date]] = []
    if os.path.isfile(_MARIJA_VACATIONS_JSON):
        try:
            with open(_MARIJA_VACATIONS_JSON, encoding="utf-8") as f:
                payload = json.load(f)
            for row in payload.get("vacations") or []:
                if not isinstance(row, dict):
                    continue
                a = str(row.get("start_date") or "").strip()
                b = str(row.get("end_date") or "").strip()
                if not a or not b:
                    continue
                out.append((date.fromisoformat(a), date.fromisoformat(b)))
        except (OSError, ValueError, json.JSONDecodeError):
            out = []
    _marija_vacation_ranges_cache = out
    return _marija_vacation_ranges_cache


def _marija_vacation_blocked_iso_in_month(year: int, month: int) -> set[str]:
    """ISO dates in `year`-`month` that fall on Marija’s vacation (for schedule shifting)."""
    blocked: set[str] = set()
    for a, b in _load_marija_vacation_ranges():
        cur = a
        while cur <= b:
            if cur.year == year and cur.month == month:
                blocked.add(cur.strftime("%Y-%m-%d"))
            cur += timedelta(days=1)
    return blocked


def _board_series_last_calendar_year(co: dict, ref_year: int) -> int:
    sk = co.get("shares_issued") or {}
    years: list[int] = []
    for k in sk:
        if isinstance(k, int):
            years.append(k)
        elif isinstance(k, str) and k.isdigit():
            years.append(int(k, 10))
    mx = max(years) if years else ref_year
    return max(ref_year, mx)


def _quarterly_meeting_date_impl(co: dict, year: int, quarter: str) -> date:
    """Quarterly governance date as `date` (monthly anchor + stagger + jitter + annual/special collision avoidance)."""
    stagger = co.get("meeting_stagger_day", 0)
    if quarter == "Q1":
        y, month, base_day = year, 4, 1
    elif quarter == "Q2":
        y, month, base_day = year, 7, 1
    elif quarter == "Q3":
        y, month, base_day = year, 10, 1
    elif quarter == "Q4":
        y, month, base_day = year, 12, 1
    else:
        raise ValueError(f"Unknown quarter: {quarter}")
    day = base_day + stagger
    d = date(y, month, day)
    qj = int(co.get("schedule_quarterly_calendar_jitter", 0) or 0)
    if qj > 0:
        rng = _schedule_rng_for_co(co, year, f"quarterly_cal_{quarter}")
        if rng is not None:
            cand = d + timedelta(days=rng.randint(-qj, qj))
            if cand.year == y and cand.month == month:
                d = cand
    blocked = {annual_meeting_date_str(co, year), board_special_meeting_date_str(co, year)}
    return _shift_date_within_month_avoiding(y, month, d, blocked)


def _board_meeting_chronological_rows(co: dict, co_name: str) -> list[tuple[str, int, str, int]]:
    """Sorted board meetings: (date_iso, type_order, kind, year). `kind` is special|agm|Q1|…|Q4."""
    if not co.get("board_sole_director_first_chronological_meeting"):
        return []
    start = co.get("minutes_start_year", co["inc_year"])
    last = _board_series_last_calendar_year(co, start)
    rows: list[tuple[str, int, str, int]] = []
    for y in range(start, last + 1):
        org = organizational_meeting_date_str(co, y)
        if org:
            rows.append((org, -1, "org", y))
        sp = board_special_meeting_date_str(co, y)
        ann = annual_meeting_date_str(co, y)
        rows.append((sp, 0, "special", y))
        rows.append((ann, 1, "agm", y))
        for q in ("Q1", "Q2", "Q3", "Q4"):
            qd = _quarterly_meeting_date_impl(co, y, q).strftime("%Y-%m-%d")
            qi = {"Q1": 2, "Q2": 3, "Q3": 4, "Q4": 5}[q]
            rows.append((qd, qi, q, y))
    rows.sort(key=lambda r: (r[0], r[1]))
    return rows


def _board_meeting_chronological_index(co_name: str, co: dict, year: int, kind: str) -> int:
    """Index of this meeting in the corporation’s full board schedule (0 = first); raises if not found."""
    rows = _board_meeting_chronological_rows(co, co_name)
    for i, r in enumerate(rows):
        if r[2] == kind and r[3] == year:
            return i
    raise ValueError(f"No chronological board row for {co_name!r} {year} {kind!r}")


def _effective_co_for_board_meeting(co: dict, co_name: str, year: int, kind: str) -> dict:
    """Shallow copy with `board_directors` cleared for the single first chronological meeting when configured."""
    if not co.get("board_sole_director_first_chronological_meeting"):
        return co
    if kind == "org" and co.get("organizational_meeting_full_board"):
        return co
    if _board_meeting_chronological_index(co_name, co, year, kind) > 0:
        return co
    out = dict(co)
    out["board_directors"] = []
    # Rely on standard sole-director / variant reliance text, not the multi-director counsel block.
    out.pop("board_meeting_reliance_markdown", None)
    return out


def _accomplishments_root() -> dict:
    global _accomplishments_cache
    if _accomplishments_cache is None:
        if os.path.isfile(_ACCOMPLISHMENTS_PATH):
            with open(_ACCOMPLISHMENTS_PATH, encoding="utf-8") as f:
                _accomplishments_cache = json.load(f)
        else:
            _accomplishments_cache = {}
    return _accomplishments_cache


def accomplishments_for_year(co_name: str, year: int) -> tuple[str | None, list[str]]:
    """Return (summary text or None, cleaned detail bullets) from accomplishments JSON for this company and calendar year."""
    jkey = _ACCOMPLISHMENTS_JSON_KEY.get(co_name)
    if not jkey:
        return None, []
    block = _accomplishments_root().get(jkey)
    if not isinstance(block, dict):
        return None, []
    yblock = block.get(str(year))
    if not isinstance(yblock, dict):
        return None, []
    summary_raw = yblock.get("summary")
    summary = summary_raw.strip() if isinstance(summary_raw, str) and summary_raw.strip() else None
    raw_list = yblock.get("annual_report", [])
    details: list[str] = []
    if isinstance(raw_list, list):
        for x in raw_list:
            if isinstance(x, str) and x.strip():
                details.append(x.strip())
    return summary, details


def _minutes_assert_exhibits_filed(co: dict) -> bool:
    """When True, minutes may say instruments are already on file / annexed. Default False avoids over-claiming."""
    return bool(co.get("minutes_assert_exhibits_filed"))


def _minute_book_exhibit_suffix(
    co: dict,
    exhibit_label: str | None,
    *,
    filed_noun: str,
    pending_noun: str,
) -> str:
    """Appendix clause for stockholder minutes: annexed vs to-be-filed."""
    if not exhibit_label:
        return ""
    if _minutes_assert_exhibits_filed(co):
        return f" A copy of {filed_noun} is **annexed to these minutes as {exhibit_label}**."
    return (
        f" {pending_noun} **are to be filed** with these minutes **as {exhibit_label}** following execution "
        "(or delivery, as applicable)."
    )


def agm_operating_addendum_markdown(
    co_name: str, year: int, exhibit_label: str, detail_lines: list[str]
) -> str:
    """Markdown body for the AGM operating addendum (same text as standalone .docx)."""
    if not detail_lines:
        return ""
    co = companies[co_name]
    display = minutes_display_name(co_name)
    bullets = "\n".join(f"- {line}" for line in detail_lines)
    if _minutes_assert_exhibits_filed(co):
        purpose_attach = (
            f"This addendum supplements the **Minutes of the Annual Meeting of the Board of Directors** of the Corporation for **{year}** "
            f"and is annexed to those minutes as **{exhibit_label}**."
        )
    else:
        purpose_attach = (
            f"This addendum supplements the **Minutes of the Annual Meeting of the Board of Directors** of the Corporation for **{year}** "
            f"and is **designated** as **{exhibit_label}** for attachment to those minutes **upon filing** with the minute book."
        )
    return f"""
**Operating addendum ({exhibit_label})**
**{display}**
{_corporation_parenthetical(co)}

**Purpose**
{purpose_attach}

**Detailed accomplishments ({year})**

{bullets}

**Other materials**
To the extent referenced in the minutes, this addendum may also include or incorporate by reference supplemental technical materials furnished for the meeting (including technical specifications, roadmaps, KPIs, and architecture diagrams).

---
""".strip() + "\n"


def _generate_agm_operating_addendum_docx(
    company_name_year: str,
    co_name: str,
    year: int,
    exhibit_label: str,
    detail_lines: list[str],
) -> None:
    """Write `{company_name_year}_agm_operating_addendum.docx` (Exhibit referenced in AGM minutes)."""
    if not detail_lines:
        return
    co = companies[co_name]
    mdate = annual_meeting_date_str(co, year)
    content = agm_operating_addendum_markdown(co_name, year, exhibit_label, detail_lines)
    path = meeting_filename(co_name, mdate, "agm_operating_addendum", ext="docx")
    print(f"Writing AGM operating addendum to {path}")
    write_docx_from_minutes(content, path, mdate, co_name)


def office_locations_for_year(ranges, year):
    year_start = date(year, 1, 1)
    year_end = date(year, 12, 31)

    locations = []

    for start, end, location in ranges:
        start_d = date.fromisoformat(start)
        end_d = date.fromisoformat(end)

        # overlap test
        if start_d <= year_end and end_d >= year_start:
            if location not in locations:
                locations.append(location)

    return locations


def office_locations_for_year_after_incorporation(ranges, year: int, inc_year: int) -> list:
    """Like `office_locations_for_year`, but drops timeline segments that end before `inc_year` begins (calendar year)."""
    year_start = date(year, 1, 1)
    year_end = date(year, 12, 31)
    corp_start = date(inc_year, 1, 1)
    locations: list[str] = []
    for start, end, location in ranges:
        start_d = date.fromisoformat(start)
        end_d = date.fromisoformat(end)
        if end_d < corp_start:
            continue
        if start_d <= year_end and end_d >= year_start:
            if location not in locations:
                locations.append(location)
    return locations


def normalize_locations(locations):
    normalized = []
    for loc in locations:
        if loc == "Denver, NC":
            normalized.append("Denver, North Carolina, USA")
        elif loc == "Lantana, FL":
            normalized.append("Lantana, Florida, USA")
        elif loc == "Wayne, Pennsylvania":
            normalized.append("Wayne, Pennsylvania, USA")
        else:
            normalized.append(loc)
    return "; ".join(normalized)

def development_locations():
    """Legacy default contractor/consultant geography list (used only if `development_centers_line` is omitted)."""
    return "Serbia; Bosnia and Herzegovina; Tunisia"


def development_centers_line_for_company(co: dict) -> str:
    """Semicolon-separated regions (contractor/consultant personnel) for minutes (per-company)."""
    if "development_centers_line" in co and isinstance(co["development_centers_line"], str):
        s = co["development_centers_line"].strip()
        if s:
            return s
        return ""
    return development_locations()


def _agm_discussion_items_line(co: dict, year: int) -> str:
    """AGM §VI discussion paragraph; supports optional `{next_year}` in company override."""
    raw = co.get("agm_discussion_items_line")
    if isinstance(raw, str) and raw.strip():
        return raw.strip().format(year=year, next_year=year + 1)
    return (
        f"The Sole Director discussed the Corporation’s transition plan for {year + 1}, including security audits, "
        "penetration testing, and commercialization readiness."
    )


def _special_meeting_primary_resolution_block(co: dict, year: int) -> str:
    """First resolution in the annual-cycle special board meeting (ratification theme differs by company)."""
    custom = co.get("special_meeting_ratification_resolution_markdown")
    if isinstance(custom, str) and custom.strip():
        return custom.strip().format(year=year, next_year=year + 1)
    return f"""**Ratification of International Operations**  
RESOLVED, that all operational and management decisions made during the Corporation’s international operations cycle for the year {year} are hereby ratified, confirmed, and approved in all respects."""


def _treasurer_report_minutes_paragraph(co: dict, issued: str) -> str:
    """Treasurer’s Report body in AGM minutes; optional full override or contingent-obligations clause only."""
    full = co.get("treasurer_report_minutes_paragraph")
    if isinstance(full, str) and full.strip():
        return full.strip().format(issued=issued, par=co["par"])
    clause = co.get(
        "treasurer_contingent_obligations_clause",
        "including notes payable, are contingent and payable upon the occurrence of a future liquidity event, the timing of which has not yet been determined.",
    )
    ack = (
        "The directors acknowledged the status of such obligations and confirmed continued oversight of these matters. "
        if len(_normalized_board_directors(co)) >= 2
        else "The Sole Director acknowledged the status of such obligations and confirmed continued oversight of these matters. "
    )
    return (
        "The Treasurer reported that the Corporation remains solvent and that certain outstanding obligations, "
        f"{clause.strip()} "
        f"{ack}"
        f"Franchise taxes and registered agent fees are paid and current. The Corporation has {issued} shares of common stock "
        f"issued and outstanding at a par value of {co['par']} per share."
    )

# 2. HELPER LOGIC
def get_location(date_str):
    target = datetime.strptime(date_str, "%Y-%m-%d")
    for start, end, loc in locations_timeline:
        s_dt = datetime.strptime(start, "%Y-%m-%d")
        e_dt = datetime.strptime(end, "%Y-%m-%d")
        if s_dt <= target <= e_dt:
            return loc
    raise Exception("Location not found for date: " + date_str)


def _zoneinfo_for_meeting(co: dict, meeting_date_iso: str) -> ZoneInfo:
    """Timezone for the operational meeting place (timeline row or registered-office fallback)."""
    if co.get("use_timeline_place", True):
        loc = get_location(meeting_date_iso)
        iana = _TIMELINE_LOCATION_TZ.get(loc)
        if not iana:
            raise ValueError(f"No IANA timezone mapped for timeline location: {loc!r}")
        return ZoneInfo(iana)
    # Principal address (Wyoming) — Mountain Time
    return ZoneInfo("America/Denver")


def _random_utime_after_meeting(filepath: str, meeting_date_iso: str, co: dict) -> None:
    """Set atime/mtime to a random weekday time 1–3 calendar days after the meeting, in the meeting timezone (9:00–16:59 local)."""
    meeting_d = datetime.strptime(meeting_date_iso, "%Y-%m-%d").date()
    tz = _zoneinfo_for_meeting(co, meeting_date_iso)
    weekday_candidates = []
    for delta in (1, 2, 3):
        d = meeting_d + timedelta(days=delta)
        if d.weekday() < 5:
            weekday_candidates.append(d)
    if not weekday_candidates:
        d = meeting_d + timedelta(days=4)
        while d.weekday() >= 5:
            d += timedelta(days=1)
        weekday_candidates = [d]
    chosen = random.choice(weekday_candidates)
    minute_of_day = random.randint(9 * 60, 17 * 60 - 1)
    h, m = divmod(minute_of_day, 60)
    s = random.randint(0, 59)
    local_dt = datetime(
        chosen.year, chosen.month, chosen.day, h, m, s, tzinfo=tz
    )
    ts = local_dt.timestamp()
    os.utime(filepath, (ts, ts))


def _annual_series_date(year: int, offset_days: int) -> date:
    """Pick a December weekday for the annual meeting series: first Monday on/after Dec 8, then consecutive weekdays."""
    ref = date(year, 12, 8)
    days_to_monday = (0 - ref.weekday()) % 7
    anchor_monday = ref + timedelta(days=days_to_monday)
    return anchor_monday + timedelta(days=offset_days)


def annual_meeting_date_str(co, year):
    """ISO date (YYYY-MM-DD) for this corporation’s annual meetings (stockholders / board / special) for the year."""
    offset_days = co.get("annual_day_offset", 0)
    d = _annual_series_date(year, offset_days)
    jw = int(co.get("schedule_annual_weekday_jitter", 0) or 0)
    if jw > 0:
        rng = _schedule_rng_for_co(co, year, "annual_weekday")
        if rng is not None:
            d = _clamp_date_to_december(_add_signed_weekdays(d, rng.randint(-jw, jw)), year)
    return d.strftime("%Y-%m-%d")


def stockholder_annual_record_date_str(co, year: int) -> str:
    """ISO date for stockholders entitled to vote at the annual meeting (record date), prior to meeting date.

    Fixed offset (calendar days) for template continuity; adjust per bylaws if a business-day or different window applies.
    """
    meeting = datetime.strptime(annual_meeting_date_str(co, year), "%Y-%m-%d").date()
    return (meeting - timedelta(days=10)).strftime("%Y-%m-%d")


def board_special_meeting_date_str(co: dict, year: int) -> str:
    """ISO date for the annual-cycle special board meeting.

    For corporations that hold an annual meeting of stockholders, this is the **record date** so the
    board may adopt the record-date resolution **on** that date (avoiding a record date that precedes
    the board action that fixes it under Delaware-style sequencing used in these templates).
    """
    if co.get("stockholder_meeting") == "annual_meeting_stockholders":
        return stockholder_annual_record_date_str(co, year)
    return annual_meeting_date_str(co, year)


def quarterly_meeting_date_str(co, year, quarter):
    """Governance check date per quarter; bumped by meeting_stagger_day so the same director isn’t quadruple-booked on one calendar day across corporations."""
    if not _quarterly_exists_for_year(co, year, quarter):
        # Not applicable before incorporation in the incorporation year.
        # Callers that enumerate quarterlies should skip generation when this returns None-like.
        # We return an ISO string anyway to avoid widening types; it should not be used because callers are gated.
        return annual_meeting_date_str(co, year)
    d = _quarterly_meeting_date_impl(co, year, quarter)
    y, month = d.year, d.month

    # In the incorporation year, if the quarterly anchor month matches the filing month, don't schedule before filing.
    filed = _incorporation_filed_date_iso(co) if year == co.get("inc_year") else None
    if filed:
        filed_d = date.fromisoformat(filed)
        if d.month == filed_d.month and d < filed_d:
            d = filed_d + timedelta(days=1)
            while d.weekday() >= 5:
                d += timedelta(days=1)
    co_name = _co_registry_key_for(co)
    if (
        co_name
        and co.get("board_sole_director_first_chronological_meeting")
        and _board_meeting_chronological_index(co_name, co, year, quarter) > 0
    ):
        blocked = {annual_meeting_date_str(co, year), board_special_meeting_date_str(co, year)}
        vac = _marija_vacation_blocked_iso_in_month(y, month)
        d = _shift_date_within_month_avoiding(y, month, d, blocked | vac)
    return d.strftime("%Y-%m-%d")


def _shift_date_within_month_avoiding(year: int, month: int, d: date, blocked_iso: set[str]) -> date:
    """If `d` is the annual or special-board ISO date, nudge within the same calendar month (± days)."""
    if d.strftime("%Y-%m-%d") not in blocked_iso:
        return d
    cand = d
    for _ in range(35):
        cand += timedelta(days=1)
        if cand.month != month or cand.year != year:
            break
        if cand.strftime("%Y-%m-%d") not in blocked_iso:
            return cand
    cand = d
    for _ in range(35):
        cand -= timedelta(days=1)
        if cand.month != month or cand.year != year:
            break
        if cand.strftime("%Y-%m-%d") not in blocked_iso:
            return cand
    for day in range(1, calendar.monthrange(year, month)[1] + 1):
        cand = date(year, month, day)
        if cand.strftime("%Y-%m-%d") not in blocked_iso:
            return cand
    return d


def meeting_place_line(co, date_str):
    """Where line for minutes: timeline place + remote participation, or principal address only."""
    if co.get("use_timeline_place", True):
        loc = get_location(date_str)
        if co.get("virtual_ok", True):
            return f"{loc}, via digital communication"
        return loc
    if co.get("virtual_ok", True):
        return f"{co['address']}, via digital communication"
    return co["address"]


def principal_address_note_markdown(co: dict) -> str:
    """
    Footnote after **Principal Address:** (or after **Location** in special/quarterly) clarifying filing/notice address vs **Place** / domicile.
    Override with `minutes_principal_address_note` (markdown); use empty string "" to suppress. When the key is absent, a default note is
    emitted if the principal address appears to be Wyoming (e.g. Sheridan registered-office style) so readers are not left guessing.
    """
    raw = co.get("minutes_principal_address_note")
    if raw is not None:
        s = str(raw).strip()
        return f"{s}\n" if s else ""
    addr = (co.get("address") or "").strip()
    if not addr:
        return ""
    if "Sheridan" not in addr and ", WY" not in addr and " WY " not in addr:
        return ""
    statute = _corporation_statute_name(co)
    if _jurisdiction(co) == "DE":
        return (
            "**Address note:** The **principal address** above is the Corporation’s designated notice and filing address on the corporate records. "
            f"The Corporation is governed by the **{statute}**; statutory references in these minutes follow that domicile. **Place** (and, where applicable, "
            "remote participation) reflects where the meeting was conducted for the stated date and may differ from the principal address.\n"
        )
    return (
        "**Address note:** The **principal address** above is the Corporation’s designated notice and filing address on the corporate records. "
        f"The Corporation is governed by the **{statute}**. **Place** identifies where the meeting was conducted for the stated date (including by approved remote means).\n"
    )


def minutes_display_name(co_name: str) -> str:
    co = companies[co_name]
    return co.get("minutes_display_name", co_name)


def _annual_stockholder_notice_section_iv(co: dict) -> str:
    """Section IV for annual stockholder meeting minutes: waiver-first, notice-first, or combined (counsel to choose)."""
    mode = co.get("annual_stockholder_notice_record", "combined")
    ex = co.get("annual_stockholder_notice_exhibit_label")
    assert_filed = _minutes_assert_exhibits_filed(co)

    if mode == "waiver_focus":
        suffix = _minute_book_exhibit_suffix(
            co,
            ex,
            filed_noun="the executed waivers of notice",
            pending_noun="Executed waivers of notice",
        )
        if assert_filed:
            return f"""**IV. Notice; Waiver**
The Chairperson confirmed that **written waivers of notice** of this annual meeting, executed by stockholders entitled to cast the votes required by applicable law and the Corporation’s bylaws, are **on file** with the records of the Corporation, and that the meeting was held in reliance on such waivers in accordance with the {_corporation_statute_name(co)} and the bylaws.{suffix}"""
        return f"""**IV. Notice; Waiver**
The Chairperson confirmed that this annual meeting was held in reliance on **written waivers of notice** for stockholders entitled to cast the votes required by applicable law and the Corporation’s bylaws, consistent with the {_corporation_statute_name(co)} and the bylaws.{suffix}"""

    if mode == "notice_focus":
        if ex:
            if _minutes_assert_exhibits_filed(co):
                suffix = (
                    f" A copy of this notice (or related delivery documentation) is **annexed to these minutes as {ex}**."
                )
            else:
                suffix = (
                    f" A copy of this notice (or related delivery documentation) **is to be filed** with these minutes **as {ex}** "
                    "following delivery."
                )
        else:
            suffix = ""
        return f"""**IV. Notice**
The Chairperson confirmed that **notice** of this annual meeting was given to each stockholder entitled to vote, **not less than** the minimum time period required by the {_corporation_statute_name(co)} and the Corporation’s bylaws, and that such notice stated the date, time, and principal place (if any) of the meeting and the **means of remote communication**, if any, for participating in the meeting.{suffix}"""

    if assert_filed:
        annex_w = f" Waivers, if any, may be **annexed to these minutes as {ex}**." if ex else ""
        return f"""**IV. Notice**
The Chairperson confirmed that **notice** of this annual meeting was given to each stockholder entitled to vote, **not less than** the minimum time period required by the {_corporation_statute_name(co)} and the Corporation’s bylaws, and that such notice stated the date, time, and principal place (if any) of the meeting and the **means of remote communication**, if any, for participating in the meeting. The Chairperson further confirmed that, to the extent notice was waived, **written waivers of notice** executed by stockholders entitled to cast sufficient votes to satisfy the requirements of applicable law and the bylaws are **on file** with the records of the Corporation.{annex_w}"""

    waiver_suffix = _minute_book_exhibit_suffix(
        co,
        ex,
        filed_noun="the executed waivers of notice",
        pending_noun="Executed waivers of notice",
    )
    return f"""**IV. Notice**
The Chairperson confirmed that **notice** of this annual meeting was given to each stockholder entitled to vote, **not less than** the minimum time period required by the {_corporation_statute_name(co)} and the Corporation’s bylaws, and that such notice stated the date, time, and principal place (if any) of the meeting and the **means of remote communication**, if any, for participating in the meeting. The Chairperson further confirmed that, to the extent notice was waived, the Corporation **is proceeding** on the basis of **written waivers of notice** from stockholders entitled to cast sufficient votes to satisfy the requirements of applicable law and the bylaws.{waiver_suffix}"""


def format_stockholders_roll_call_block(co: dict) -> str:
    roll = co.get("stockholders_roll_call")
    absent = co.get("stockholders_absent_line", "None.")
    if not roll:
        return f"""**Stockholders Present:**  
The stockholders of the Corporation holding a majority of the outstanding shares of the Corporation entitled to vote at the meeting.

**Stockholders Absent:**  
{absent}"""
    lines = "\n".join(f"- **{r['name']}**, {r['presence']}" for r in roll)
    collective = co.get(
        "stockholders_quorum_collective_sentence",
        "Together with other record holders voting in person or by valid proxy as tallied for this meeting, these stockholders "
        "constituted holders of a majority of the outstanding shares entitled to vote at the meeting.",
    )
    return f"""**Stockholders Present:**  
{lines}

{collective}

**Stockholders Absent:**  
{absent}"""


def annual_stockholder_director_vote_tabulation_markdown(co: dict) -> str:
    """Optional per-holder vote lines after §VI for `annual_meeting_stockholders` (remediation: defensible tabulation)."""
    raw = co.get("annual_stockholder_director_election_votes")
    if not isinstance(raw, list) or not raw:
        return ""
    lines: list[str] = []
    for item in raw:
        if not isinstance(item, dict):
            continue
        nm = str(item.get("name") or "").strip()
        if not nm:
            continue
        vote = str(item.get("vote") or "FOR").strip().upper()
        shares = str(item.get("shares") or "").strip()
        sh_part = f"holding **{shares}** shares, " if shares else ""
        lines.append(f"- **{nm}**, {sh_part}cast **{vote}**.")
    if not lines:
        return ""
    return (
        "\n**VI-A. Vote Tabulation (Election of Directors)**\n\n"
        "The Chairperson caused the following votes to be recorded for the election resolution adopted in Section VI:\n\n"
        + "\n".join(lines)
        + "\n"
    )


def _stockholder_waiver_signature_blocks(co: dict, execution_date: str) -> str:
    """Signature area for stockholder waiver form (plain text; no lines; no dates)."""
    roll = co.get("stockholders_roll_call")
    if not roll:
        return """**Executed by (stockholder):**

"""
    parts = []
    for r in roll:
        parts.append(
            f"""**Executed by (stockholder):** {r["name"]}"""
        )
    # Extra whitespace between individual executions for readability.
    return "\n\n\n".join(parts)


def stockholder_waiver_of_notice_annual_meeting_markdown(
    company_name_year: str, year: int, co_name: str
) -> str:
    """Markdown for standalone waiver of notice (annual stockholders)."""
    co = companies[co_name]
    date_iso = annual_meeting_date_str(co, year)
    record_date = stockholder_annual_record_date_str(co, year)
    place = meeting_place_line(co, date_iso)
    as_meeting = datetime.strptime(date_iso, "%Y-%m-%d").strftime("%B %d, %Y")
    display_company = minutes_display_name(co_name)
    sig_blocks = _stockholder_waiver_signature_blocks(co, as_meeting)
    t_stock = scheduled_stockholder_meeting_time(co, year)
    return f"""
**Waiver of Notice of Annual Meeting of Stockholders**
**{display_company}**
{_corporation_parenthetical(co)}

The undersigned record stockholder(s) of **{display_company}** (the “Corporation”) entitled to vote at the annual meeting described below, intending to be legally bound, **waive notice** of that meeting and of any postponement or adjournment thereof to the extent permitted by the {_corporation_statute_name(co)}, the Corporation’s **certificate of incorporation**, and **bylaws**.

**Meeting**  
**Date:** {date_iso} ({as_meeting})  
**Time:** {t_stock}  
**Place / remote means:** {place}

**Record date ({_corp_law_section_ref(co, "213")}):** {record_date}  
(Stockholders of record as of this date are entitled to notice of, and to vote at, the meeting, subject to the certificate of incorporation and bylaws.)

**Business**  
The meeting may include election of directors and any other annual business proper under the charter and bylaws.

{sig_blocks}
---
"""


def generate_stockholder_waiver_of_notice_annual_meeting(
    company_name_year: str, year: int, co_name: str
) -> None:
    """Standalone waiver of notice for the annual stockholder meeting (file alongside minutes; sign and annex as Exhibit A if used)."""
    co = companies[co_name]
    date_iso = annual_meeting_date_str(co, year)
    content = stockholder_waiver_of_notice_annual_meeting_markdown(company_name_year, year, co_name)
    path = meeting_filename(co_name, date_iso, "waiver_of_notice_annual_stockholder_meeting", ext="docx")
    print(f"Writing Waiver of Notice (annual stockholders) to {path}")
    write_docx_from_minutes(content, path, date_iso, co_name)


def notice_of_annual_stockholder_meeting_markdown(
    company_name_year: str, year: int, co_name: str
) -> str:
    """Markdown for §222-style notice of annual stockholder meeting."""
    co = companies[co_name]
    date_iso = annual_meeting_date_str(co, year)
    record_date = stockholder_annual_record_date_str(co, year)
    place = meeting_place_line(co, date_iso)
    as_meeting = datetime.strptime(date_iso, "%Y-%m-%d").strftime("%B %d, %Y")
    notice_date_iso = board_special_meeting_date_str(co, year)
    as_notice = datetime.strptime(notice_date_iso, "%Y-%m-%d").strftime("%B %d, %Y")
    display_company = minutes_display_name(co_name)
    principal = co["address"]
    officer = co.get("notice_signatory_line", "Derek E. Pappas, President")
    t_stock = scheduled_stockholder_meeting_time(co, year)
    return f"""
**Notice of Annual Meeting of Stockholders**
**{display_company}**
{_corporation_parenthetical(co)}

To the stockholders of the Corporation:

Notice is hereby given that an **annual meeting of stockholders** of **{display_company}** (the “Corporation”) will be held:

**Date:** {as_meeting} ({date_iso})  
**Time:** {t_stock}  
**Place / means of participation:** {place}

**Record date:** The **record date** for determining stockholders entitled to notice of and to vote at the meeting (or any adjournment or postponement) is **{record_date}**, as fixed by the Board of Directors in accordance with the bylaws and the {_corporation_statute_name(co)}.

**Purpose**  
To elect directors and to transact such other business as may properly come before the meeting and any adjournment or postponement, in accordance with the certificate of incorporation and bylaws.

Stockholders may attend and vote in person or, where permitted by the bylaws and applicable law, by remote communication in the manner described in this notice or in supplemental instructions provided by the Corporation.

This notice is given pursuant to the {_corporation_statute_name(co)} and the Corporation’s bylaws. Notice may be delivered by hand, United States mail, or electronic transmission in the manner and within the time frames permitted by applicable law and the bylaws.

**Principal executive office:** {principal}

By order of the Board of Directors,

{officer}

**Date of this notice:** {as_notice} ({notice_date_iso})
---
"""


def generate_notice_of_annual_stockholder_meeting(
    company_name_year: str, year: int, co_name: str
) -> None:
    """Standalone §222-style notice of annual stockholder meeting (deliver or file as required; may annex to minute book)."""
    co = companies[co_name]
    date_iso = annual_meeting_date_str(co, year)
    content = notice_of_annual_stockholder_meeting_markdown(company_name_year, year, co_name)
    path = meeting_filename(co_name, date_iso, "notice_of_annual_stockholder_meeting", ext="docx")
    print(f"Writing Notice of Annual Stockholder Meeting to {path}")
    write_docx_from_minutes(content, path, date_iso, co_name)


def _board_meeting_rows_for_year(co: dict, year: int) -> list[tuple[str, str, str, str]]:
    """(date_iso, meeting_title, time_str, place_line) sorted chronologically; matches minuted board meetings for the year."""
    org = organizational_meeting_date_str(co, year)
    annual = annual_meeting_date_str(co, year)
    special = board_special_meeting_date_str(co, year)
    t_special = scheduled_special_meeting_time(co, year)
    t_board = scheduled_board_agm_time(co, year)
    # (date_iso, minutes_since_midnight, tie_seq, title, time_str, place)
    rows: list[tuple[str, int, int, str, str, str]] = []
    if org:
        rows.append(
            (
                org,
                _clock_to_minutes_since_midnight(ORGANIZATIONAL_MEETING_TIME),
                -1,
                "Organizational Meeting of the Board of Directors",
                ORGANIZATIONAL_MEETING_TIME,
                meeting_place_line(co, org),
            )
        )
    # Special board meeting (often on record date for stockholder-annual corps) precedes the December annual board block.
    rows.append(
        (
            special,
            _clock_to_minutes_since_midnight(t_special),
            0,
            "Special Meeting of the Board of Directors",
            t_special,
            meeting_place_line(co, special),
        )
    )
    rows.append(
        (
            annual,
            _clock_to_minutes_since_midnight(t_board),
            1,
            "Annual Meeting of the Board of Directors",
            t_board,
            meeting_place_line(co, annual),
        )
    )
    for seq, quarter in enumerate(("Q1", "Q2", "Q3", "Q4"), start=2):
        if not _quarterly_exists_for_year(co, year, quarter):
            continue
        qd = quarterly_meeting_date_str(co, year, quarter)
        t_q = scheduled_quarterly_meeting_time(co, year, quarter)
        rows.append(
            (
                qd,
                _clock_to_minutes_since_midnight(t_q),
                seq,
                f"Quarterly Governance Meeting – {year} {quarter}",
                t_q,
                meeting_place_line(co, qd),
            )
        )
    rows.sort(key=lambda r: (r[0], r[1], r[2]))
    return [(r[0], r[3], r[4], r[5]) for r in rows]


def board_waiver_of_notice_markdown(company_name_year: str, year: int, co_name: str) -> str:
    """Markdown for waiver of notice (board meetings listed for the year); sole director or full board."""
    co = companies[co_name]
    director_name = "Derek E. Pappas"
    doc_date_iso = annual_meeting_date_str(co, year)
    rows = _board_meeting_rows_for_year(co, year)
    bullet_lines = "\n".join(
        f"- **{title}** — **Date:** {d_iso}; **Time:** {t_str}; **Place / remote means:** {place}"
        for d_iso, title, t_str, place in rows
    )
    bylaws_ref = co.get("board_notice_waiver_bylaws_ref") or "**the Corporation\u2019s bylaws**"
    bds = _normalized_board_directors(co)
    if len(bds) >= 2:
        names_join = " and ".join(f"**{d['name']}**" for d in bds)
        intro = (
            f"The undersigned, {names_join}, constituting **all** of the directors of **{co_name}** (the “Corporation”), intending to be legally bound, "
            f"**each waives all notice** of the time, place, and purposes of each meeting of the Board of Directors of the Corporation listed below, "
            f"and of any postponement or adjournment of any such meeting, to the extent permitted by the **{_corporation_statute_name(co)}**, "
            f"the Corporation’s **certificate of incorporation**, and {bylaws_ref}. This waiver is given to supplement the minutes of the Corporation, "
            "which state that notice of each such meeting was duly given **or waived**."
        )
        sig = board_meeting_signature_markdown(co, doc_date_iso, sole_director_name=director_name)
    else:
        intro = (
            f"The undersigned, **{director_name}**, Sole Director of **{co_name}** (the “Corporation”), intending to be legally bound, **waives all notice** "
            f"of the time, place, and purposes of each meeting of the Board of Directors of the Corporation listed below, and of any postponement or adjournment "
            f"of any such meeting, to the extent permitted by the **{_corporation_statute_name(co)}**, the Corporation’s **certificate of incorporation**, and {bylaws_ref}. "
            "This waiver is given to supplement the minutes of the Corporation, which state that notice of each such meeting was duly given **or waived**."
        )
        sig = signature_block(co, director_name, doc_date_iso, title="Sole Director")
    return f"""
**Waiver of Notice of Meetings of the Board of Directors**
**{co_name}**
{_corporation_parenthetical(co)}

**Calendar year {year}**

{intro}

**Meetings covered**

{bullet_lines}

{sig}
---
"""


def generate_board_waiver_of_notice(
    company_name_year: str, year: int, co_name: str
) -> None:
    """Optional standalone waiver: sole director waives notice of board meetings minuted for this year (file with minute book).

    Supports the board minutes line that notice of each meeting was duly given **or waived** (Delaware DGCL + typical bylaws).
    """
    co = companies[co_name]
    doc_date_iso = annual_meeting_date_str(co, year)
    content = board_waiver_of_notice_markdown(company_name_year, year, co_name)
    path = meeting_filename(co_name, doc_date_iso, "waiver_of_notice_board_meetings", ext="docx")
    print(f"Writing Waiver of Notice (board meetings) to {path}")
    write_docx_from_minutes(content, path, doc_date_iso, co_name)


SIGNATURE_BLOCK_MARKER = "<<<SIGNATURE_BLOCK>>>"


def _format_signature_date(co: dict, date_iso: str) -> str:
    """Format signature date from an ISO date string.

    Supported formats:
    - "iso" (default): YYYY-MM-DD
    - "long": Month DD, YYYY
    """
    fmt = str(co.get("signature_block_date_format") or "iso").strip().lower()
    if fmt == "long":
        return datetime.strptime(date_iso, "%Y-%m-%d").strftime("%B %d, %Y")
    return date_iso


def wet_signing_lines_markdown(co: dict, name: str, title: str, date_iso: str, *, include_filled_date: bool = True) -> str:
    """Blank signature rule plus **Name:** / **Title:** / **Date:** when `signature_block_print_signing_lines` is set."""
    if not co.get("signature_block_print_signing_lines"):
        return ""
    rule = str(co.get("signature_block_signing_rule_line") or "_______________________________").strip()
    date_label = str(co.get("signature_block_date_label") or "Date:").strip()
    if include_filled_date and bool(co.get("signature_block_include_date", True)):
        date_val = _format_signature_date(co, date_iso)
        date_line = f"**{date_label}** {date_val}"
    else:
        date_line = f"**{date_label}** {rule}"
    return f"{rule}\n**Name:** {name}\n**Title:** {title}\n{date_line}\n"


def signature_block(co: dict, name: str, date: str, *, title: str = "Sole Director") -> str:
    """Signature block; optional wet-ink lines via `signature_block_print_signing_lines`.

    Uses a marker so the .docx writer can keep the block together to avoid splitting across pages.
    """
    style = (co.get("signature_block_style") or "executed_by").strip().lower()
    include_date = bool(co.get("signature_block_include_date", True))
    include_title_in_label = bool(co.get("signature_block_include_title_in_label", True))
    label_template = co.get("signature_block_label_template")
    name_prefix = str(co.get("signature_block_name_prefix") or "").strip()
    date_label = str(co.get("signature_block_date_label") or "Date:").strip()
    rendered_date = _format_signature_date(co, date)
    spacing_lines = int(co.get("signature_block_spacing_lines", 1))
    print_lines = bool(co.get("signature_block_print_signing_lines", False))

    if style == "signature":
        header = "**Signature:**"
    elif style == "none":
        header = ""
    else:
        if isinstance(label_template, str) and label_template.strip():
            header = label_template.strip()
        elif include_title_in_label and title:
            header = f"**Executed by ({title}):**"
        else:
            header = "**Executed by:**"

    lines: list[str] = [SIGNATURE_BLOCK_MARKER]
    if header:
        lines.append(header)
    if print_lines:
        rule = str(co.get("signature_block_signing_rule_line") or "_______________________________").strip()
        lines.append(rule)
        lines.append(f"**Name:** {name}")
        lines.append(f"**Title:** {title}")
        if include_date:
            lines.append(f"**{date_label}** {rendered_date}")
        else:
            lines.append(f"**{date_label}** {rule}")
    else:
        if name_prefix:
            lines.append(f"{name_prefix} {name}".rstrip())
        else:
            lines.append(name)
        if include_date:
            lines.append(f"**{date_label}** {rendered_date}")
    lines.extend([""] * max(spacing_lines, 0))
    return "\n".join(lines)


def board_meeting_signature_markdown(co: dict, date_iso: str, *, sole_director_name: str = "Derek E. Pappas") -> str:
    """Closing signatures for board minutes: dual signatures when `board_directors` lists the full board."""
    bds = _normalized_board_directors(co)
    if len(bds) >= 2:
        parts: list[str] = []
        for d in bds:
            title = (d.get("title") or "Director").strip() or "Director"
            parts.append(signature_block(co, d["name"], date_iso, title=title))
        return "\n".join(parts)
    return signature_block(co, sole_director_name, date_iso, title="Sole Director")


def _adopted_resolutions_section(co: dict, section_heading: str, resolution_parts: list[str]) -> str:
    """Resolutions section: sole director vs full board wording based on `board_directors`."""
    cleaned = [p.strip() for p in resolution_parts if p and str(p).strip()]
    multi = len(_normalized_board_directors(co)) >= 2
    if not cleaned:
        return f"{section_heading}\nNone. No resolutions were presented for adoption.\n\n"
    if len(cleaned) == 1:
        intro = (
            "Upon consideration, the Board adopted the following resolution:\n\n"
            if multi
            else "Upon consideration, the Sole Director adopted the following resolution:\n\n"
        )
    else:
        intro = (
            "Upon consideration, the Board adopted the following resolutions:\n\n"
            if multi
            else "Upon consideration, the Sole Director adopted the following resolutions:\n\n"
        )
    body = "\n\n".join(cleaned) + "\n\n"
    return f"{section_heading}\n{intro}{body}"


def _board_appointment_resolution_blocks_if_first_meeting(
    co_name: str, co: dict, year: int, kind: str
) -> list[str]:
    """Optional appointment block(s) for the first (chronological) board meeting when configured.

    Used to formally minute that a second director is appointed/elected after the initial sole-director meeting.
    """
    if not co.get("board_sole_director_first_chronological_meeting"):
        return []
    if kind == "org" and co.get("organizational_meeting_full_board"):
        return []
    try:
        if _board_meeting_chronological_index(co_name, co, year, kind) != 0:
            return []
    except ValueError:
        return []

    bds = _normalized_board_directors(co)
    # Only add the appointment motion when the company is configured with a multi-director board.
    if len(bds) < 2:
        return []

    # Convention here: appoint any non-Derek directors listed in `board_directors`.
    appointees = [d for d in bds if d.get("name") and d["name"] != "Derek E. Pappas"]
    if not appointees:
        return []

    names = ", ".join(d["name"] for d in appointees)
    return [
        f"""**Appointment of Additional Director(s)**  
RESOLVED, that the size of the Board of Directors of the Corporation is fixed at **{len(bds)}** director(s); and  
FURTHER RESOLVED, that **{names}** is hereby appointed and elected to serve as a director of the Corporation, to hold office until a successor is duly elected and qualified or until earlier resignation or removal, effective immediately following this meeting."""
    ]


def _agm_president_report_product_line_resolved(co: dict, year: int, multi: bool) -> str:
    """`agm_president_report_product_line`: str (optional `{year}` / `{next_year}`) or dict with year keys and optional `default`."""
    raw = co.get("agm_president_report_product_line")
    if isinstance(raw, dict):
        chunk = raw.get(str(year)) or raw.get("default")
        if isinstance(chunk, str) and chunk.strip():
            return chunk.strip().format(year=year, next_year=year + 1)
    elif isinstance(raw, str) and raw.strip():
        return raw.strip().format(year=year, next_year=year + 1)
    return (
        "The directors summarized continued development of the Corporation’s software and service offerings. "
        if multi
        else "The Sole Director summarized continued development of the Corporation’s software and service offerings. "
    )


def _agm_president_report_body(co: dict, office_locations: str, dev_locations: str, co_name: str, year: int) -> str:
    """President’s Report narrative: operational baseline, optional product/hosting lines, accomplishments summary, addendum."""
    multi = len(_normalized_board_directors(co)) >= 2
    period = _period_phrase(co)
    opener = co.get("agm_president_report_opening_paragraph_markdown")
    if isinstance(opener, str) and opener.strip():
        base = (
            opener.format(
                office_locations=office_locations,
                dev_locations=dev_locations or "locations noted in the President’s report",
                year=year,
            ).strip()
            + " "
        )
    else:
        if dev_locations.strip():
            geo_ops = (
                f"with operations conducted from {office_locations} and engineering support sourced through "
                f"**contractor and consultant arrangements** with personnel **in {dev_locations.strip()}**"
            )
        else:
            geo_ops = (
                f"with operations conducted from {office_locations} and engineering support sourced through "
                "**contractor and consultant arrangements** as described in the President’s report"
            )
        if multi:
            base = (
                f"The directors reported on the Corporation’s operational and engineering activities for {period}, "
                "including centralized management of globally distributed development and the use of operational office location(s) "
                f"during {period}, {geo_ops}, "
                "while confirming that management, oversight, and decision-making remained centralized and continuously recorded "
                "through the Corporation’s official records. "
            )
        else:
            base = (
                f"The Sole Director reported on the Corporation’s operational and engineering activities for {period}, "
                "including centralized management of globally distributed development and the use of operational office location(s) "
                f"during {period}, {geo_ops}, "
                "while confirming that management, oversight, and decision-making remained centralized and continuously recorded "
                "through the Corporation’s official records. "
            )
    product = _agm_president_report_product_line_resolved(co, year, multi)
    if not product.endswith(" "):
        product = product.rstrip() + " "
    infra_raw = co.get("agm_president_report_infrastructure_line") or ""
    if isinstance(infra_raw, str) and infra_raw.strip():
        ir = infra_raw.strip()
        infra = " " + (ir.format(year=year, next_year=year + 1) if "{" in ir else ir)
    else:
        infra = ""

    summary, detail_items = accomplishments_for_year(co_name, year)
    summary_sentence = ""
    if summary:
        art, noun = _financial_year_words(co)
        summary_sentence = (
            f" The President’s report included the following **accomplishments summary** for {art} {noun} {year}: "
            f"{summary}"
        )

    lab = co.get("agm_president_report_operating_exhibit_label")
    if detail_items and not lab:
        lab = "Exhibit B"

    attach = (
        f"annexed as **{lab}**"
        if _minutes_assert_exhibits_filed(co)
        else f"designated as **{lab}** for attachment to these minutes upon filing"
    )
    exhibit = ""
    if detail_items and lab:
        exhibit = (
            f" A written addendum {attach} sets forth **detailed accomplishments** "
            "for the period covered by the President’s report, together with supplemental technical materials furnished "
            "for the meeting (including technical specifications, roadmaps, KPIs, and architecture diagrams) where applicable."
        )
    elif lab:
        exhibit = (
            f" A written addendum {attach} sets forth additional detail furnished for the meeting, "
            "including technical specifications, roadmaps, KPIs, and architecture diagrams."
        )

    ip_default = (
        " All software, algorithms, and intellectual property developed **for the Corporation** during the year "
        "under applicable **contractor, consultant, and employment arrangements**, regardless of where services were performed, "
        "were reaffirmed as **properly titled to and the exclusive property of the Corporation** under those arrangements and applicable law, "
        "without implying ownership of **counterparties’ equipment, facilities, or premises**."
    )
    ip_custom = co.get("agm_ip_affirmation_sentence")
    if isinstance(ip_custom, str) and ip_custom.strip():
        ip_close = " " + ip_custom.strip()
    else:
        ip_close = ip_default
    return base + product + infra + summary_sentence + exhibit + ip_close


def _period_phrase(co: dict) -> str:
    """e.g. 'the calendar year' / 'the fiscal year' for narrative paragraphs."""
    art, noun = _financial_year_words(co)
    return f"{art} {noun}"


def _agm_banking_resolution_markdown(co: dict, director_name: str) -> str:
    bank = co.get("primary_banking_institution", "JPMorgan Chase Bank, N.A.")
    raw = co.get("agm_banking_authorized_signatories")
    names: list[str] = []
    if isinstance(raw, list):
        names = [str(x).strip() for x in raw if str(x).strip()]
    if not names:
        one = (
            str(co.get("agm_banking_authorized_signatory") or "").strip()
            or str((co.get("officers") or {}).get("CEO") or "").strip()
            or director_name
        )
        names = [one]
    if len(names) == 1:
        return f"""**Banking Authorization**  
RESOLVED, that **{names[0]}** is authorized to open, maintain, and manage one or more corporate bank accounts in the name of the Corporation at {bank}, and any successor institution, and to act as the **sole authorized signatory** with full authority to execute all related documents."""
    bolded = ", ".join(f"**{n}**" for n in names[:-1]) + f", and **{names[-1]}**"
    return f"""**Banking Authorization**  
RESOLVED, that each of {bolded} is authorized to open, maintain, and manage one or more corporate bank accounts in the name of the Corporation at {bank}, and any successor institution, and to act as an **authorized signatory** with full authority to execute routine banking documents, in each case subject to any dual-signature or other controls the Board or the financial institution may require."""


def _agm_resolutions_block(co: dict, director_name: str, year: int, *, extra_blocks: list[str] | None = None) -> str:
    if "annual_board_resolution_blocks" in co:
        parts = list(co["annual_board_resolution_blocks"])
    else:
        period = _financial_year_words(co)[1]
        parts = [
            f"""**Approval of Financial Reports**  
RESOLVED, that the financial statements for {period} {year} are hereby approved.""",
            f"""**Approval of {year + 1} Budget**  
RESOLVED, that the operating, engineering, and marketing budget for {period} {year + 1} is hereby approved.""",
            _agm_banking_resolution_markdown(co, director_name),
        ]
    if extra_blocks:
        parts = list(extra_blocks) + parts
    return _adopted_resolutions_section(co, "**VII. Resolutions**", parts)


def _special_resolutions_block(
    co_name: str,
    co: dict,
    year: int,
    record_date_resolution: str,
    *,
    extra_blocks: list[str] | None = None,
) -> str:
    parts: list[str] = []
    if extra_blocks:
        parts.extend(list(extra_blocks))
    parts.extend(_domestication_resolution_blocks_if_due(co, board_special_meeting_date_str(co, year)))
    parts.append(_special_meeting_primary_resolution_block(co, year))
    extra = record_date_resolution.strip()
    if extra:
        parts.append(extra)
    return _adopted_resolutions_section(co, "**III. Resolutions:**", parts)


def _quarterly_resolutions_block(
    co_name: str,
    render_co: dict,
    base_co: dict,
    year: int,
    quarter: str,
    *,
    extra_blocks: list[str] | None = None,
) -> str:
    """Quarterly resolutions.

    `render_co` drives the sole-director vs multi-director wording; `base_co` is used for one-time governance actions
    that may take effect immediately after the meeting (e.g. appointing an additional director).
    """
    if "quarterly_resolution_blocks" in render_co:
        parts = list(render_co["quarterly_resolution_blocks"])
    else:
        qdef = render_co.get(
            "quarterly_default_ratification_resolution",
            "RESOLVED, that all operational, infrastructure, and intellectual property assets created during the quarter are hereby ratified, confirmed, and approved as assets of the Corporation.",
        )
        parts = [qdef]
    prefix = list(extra_blocks or [])
    parts = (
        _domestication_resolution_blocks_if_due(base_co, quarterly_meeting_date_str(base_co, year, quarter))
        + _board_appointment_resolution_blocks_if_first_meeting(co_name, base_co, year, quarter)
        + prefix
        + parts
    )
    n = len([p for p in parts if p and str(p).strip()])
    if n == 0:
        heading = "**IV. Resolutions**"
    elif n == 1:
        heading = "**IV. Resolution:**"
    else:
        heading = "**IV. Resolutions**"
    return _adopted_resolutions_section(render_co, heading, parts)


# 3. OUTPUT HELPERS

def write_docx_from_minutes(
    content: str,
    filepath: str,
    meeting_date_iso: str | None = None,
    co_name: str | None = None,
    *,
    minute_book_page_breaks: bool = False,
):
    doc = Document()
    pending_keep_with_next = False
    for raw_line in content.splitlines():
        line = raw_line.rstrip()
        if minute_book_page_breaks and line.strip() == MEETING_BOOK_PAGE_BREAK_MARKER:
            doc.add_page_break()
            continue
        if line.strip() == SIGNATURE_BLOCK_MARKER:
            pending_keep_with_next = True
            continue
        if not line.strip():
            if pending_keep_with_next:
                pending_keep_with_next = False
            doc.add_paragraph()
            continue
        # Render inline markdown-style bold using **...** into bold runs
        p = doc.add_paragraph()
        if pending_keep_with_next:
            # Try to keep the signature block from splitting across pages.
            # (Word honors these flags when it can.)
            p.paragraph_format.keep_with_next = True
            p.paragraph_format.keep_together = True
        i = 0
        while i < len(line):
            if line[i:i+2] == "**":
                j = line.find("**", i + 2)
                if j != -1:
                    bold_text = line[i+2:j]
                    run = p.add_run(bold_text)
                    run.bold = True
                    i = j + 2
                else:
                    # No closing ** found; write the remaining text as-is
                    p.add_run(line[i:])
                    break
            else:
                next_bold = line.find("**", i)
                if next_bold == -1:
                    p.add_run(line[i:])
                    break
                else:
                    p.add_run(line[i:next_bold])
                    i = next_bold
    doc.save(filepath)
    if meeting_date_iso and co_name:
        _random_utime_after_meeting(filepath, meeting_date_iso, companies[co_name])

# 3. GENERATORS
def generate_agm(co_name, year):
    co = companies[co_name]
    eco = _effective_co_for_board_meeting(co, co_name, year, "agm")
    date = annual_meeting_date_str(co, year)
    place = meeting_place_line(co, date)
    issued = co["shares_issued"].get(year, "4,000,000")
    display_company = minutes_display_name(co_name)
    t_stock = scheduled_stockholder_meeting_time(co, year)
    t_board = scheduled_board_agm_time(co, year)

    # select locations for the given year (clip pre-incorporation timeline unless opted out)
    if co.get("agm_locations_respect_incorporation_year", True):
        locations = office_locations_for_year_after_incorporation(
            locations_timeline, year, co["inc_year"]
        )
    else:
        locations = office_locations_for_year(locations_timeline, year)

    # normalize and format for template insertion
    office_locations = normalize_locations(locations)

    dev_locations = development_centers_line_for_company(co)

    director_name = "Derek E. Pappas"
    chair_name = str(eco.get("board_meeting_chair_name") or director_name).strip() or director_name
    full_board = _normalized_board_directors(eco)
    multi_director_board = len(full_board) >= 2
    inc_year = co["inc_year"]
    minutes_start_year = co.get("minutes_start_year", inc_year)
    if year > minutes_start_year:
        prior_date = annual_meeting_date_str(co, year - 1)
        approver = "the Board" if multi_director_board else "the Sole Director"
        prior_minutes_section = f"""**IV. Approval of Prior Minutes**
The minutes of the prior Annual Meeting of the Board of Directors held on {prior_date} were reviewed and approved by {approver}."""
    elif year == minutes_start_year and minutes_start_year > inc_year:
        # Only when the minute book series starts after incorporation (currently: DATA RECORD SCIENCE). All other registry
        # companies have minutes_start_year == inc_year, so their first generated AGM uses the “first after incorporation” branch below.
        prior_minutes_section = f"""**IV. Approval of Prior Minutes**
Board minutes included in **this** compiled minute book series begin with calendar year **{minutes_start_year}**. The Corporation was incorporated in **{inc_year}**. No annual board minutes from prior calendar years **within this compilation series** were presented for approval."""
    else:
        inc_filed = _incorporation_filed_date_iso(co)
        inc_phrase = f"on {_fmt_long_date(inc_filed)}" if inc_filed else f"in {inc_year}"
        prior_minutes_section = f"""**IV. Approval of Prior Minutes**
This was the first Annual Meeting of the Board of Directors following incorporation of the Corporation {inc_phrase}; no prior annual meeting of the Board was held and no prior annual board minutes were presented for approval."""

    if co.get("stockholder_meeting") == "annual_meeting_stockholders":
        call_intro = f"Immediately following the Annual Meeting of Stockholders of the Corporation held on {date} commencing at {t_stock}, "
    else:
        call_intro = ""

    roll_quorum_block = board_roll_quorum_markdown(eco, co_name, date, "agm", director_name=director_name)
    reliance_141e_line = board_director_reliance_paragraph(eco, co_name, date, "agm")

    consent_cross_ref = ""
    if co.get("stockholder_meeting") == "written_consent":
        as_of_fmt = datetime.strptime(date, "%Y-%m-%d").strftime("%B %d, %Y")
        lab = co.get("sole_stockholder_consent_exhibit_label")
        law228 = _corp_law_section_ref(co, "228")
        noter = "The Board noted" if multi_director_board else "The Sole Director noted"
        if _minutes_assert_exhibits_filed(co):
            annex = f" (annexed as **{lab}**)" if lab else ""
            consent_line = (
                f"{noter} that the **Written Consent of Sole Stockholder** dated {as_of_fmt}, adopting stockholder resolutions "
                f"under **{law228}** for the year {year}, is **on file** with the minutes of the stockholders of the Corporation{annex}."
            )
        else:
            des = f", **to be designated {lab}** upon filing" if lab else ""
            consent_line = (
                f"{noter} that the **Written Consent of Sole Stockholder** dated {as_of_fmt}, adopting stockholder resolutions "
                f"under **{law228}** for the year {year}, **will be filed** with the minutes of the stockholders of the Corporation **upon execution**{des}."
            )
        consent_cross_ref = f"""
**Stockholder Written Consent ({year})**  
{consent_line}

"""

    addr_note = principal_address_note_markdown(co)
    materials_ack_block = board_meeting_materials_acknowledgment_block(co)

    if multi_director_board:
        others = [d["name"] for d in full_board if d["name"] != chair_name]
        others_txt = " and ".join(others) if others else ""
        co_tail = (
            f"by **{chair_name}**, acting as Chair of the Board, with **{others_txt}** also present as a director—"
            "all directors constituting the full membership of the Board were in attendance."
            if others_txt
            else f"by **{chair_name}**, acting as Chair of the Board, with all directors present."
        )
        call_to_order_body = (
            f"{call_intro}The Annual Meeting of the Board of Directors of {display_company} (the “Corporation”) "
            f"was called to order at {t_board} on {date} {co_tail}"
        )
    else:
        call_to_order_body = (
            f"{call_intro}The Annual Meeting of the Board of Directors of {display_company} (the “Corporation”) "
            f"was called to order at {t_board} on {date} by {director_name}, acting as Sole Director, President, and Treasurer of the Corporation."
        )

    return f"""
**Minutes of the Annual Meeting of the Board of Directors**
{_corporation_parenthetical(co)}

**I. Meeting Information**
**Company Name:** {display_company}
**Principal Address:** {co['address']}
{addr_note}**Date:** {date}
**Time:** {t_board}
**Place:** {place}
**Type of Meeting:** Annual Meeting of the Board of Directors

**II. Call to Order**
{call_to_order_body}

{roll_quorum_block}{prior_minutes_section}
{materials_ack_block}**V. Reports of Officers**

**President’s Report:**  
{_agm_president_report_body(eco, office_locations, dev_locations, co_name, year)}

**Treasurer’s Report:**  
{_treasurer_report_minutes_paragraph(eco, issued)}

{reliance_141e_line}

**VI. Discussion Items**
{_agm_discussion_items_line(eco, year)}

{_agm_resolutions_block(eco, director_name, year, extra_blocks=_board_resolution_prefix_blocks(co_name, co, year, date, "agm"))}
{consent_cross_ref}
**VIII. Adjournment**
There being no further business to come before the Board, the meeting was adjourned.

{board_meeting_signature_markdown(eco, date, sole_director_name=director_name)}
---
"""


def generate_special(co_name, year):
    co = companies[co_name]
    eco = _effective_co_for_board_meeting(co, co_name, year, "special")
    annual_date = annual_meeting_date_str(co, year)
    date = board_special_meeting_date_str(co, year)
    place = meeting_place_line(co, date)
    t_stock = scheduled_stockholder_meeting_time(co, year)
    t_special = scheduled_special_meeting_time(co, year)

    director_name = "Derek E. Pappas"
    chair_name = str(eco.get("board_meeting_chair_name") or director_name).strip() or director_name
    full_board = _normalized_board_directors(eco)
    multi_director_board = len(full_board) >= 2
    display_company = minutes_display_name(co_name)

    roll_quorum_block = board_roll_quorum_markdown(eco, co_name, date, "special", director_name=director_name)
    reliance_141e_line = board_director_reliance_paragraph(eco, co_name, date, "special")

    record_date_resolution = ""
    if co.get("stockholder_meeting") == "annual_meeting_stockholders":
        rd = stockholder_annual_record_date_str(co, year)
        record_date_resolution = f"""
**Record Date for Annual Meeting of Stockholders ({_corp_law_section_ref(co, "213")})**  
RESOLVED, that **{rd}** is hereby fixed as the record date for determining the stockholders entitled to notice of and to vote at the Annual Meeting of Stockholders of the Corporation to be held on **{annual_date}** commencing at **{t_stock}**, in accordance with the Corporation’s bylaws and the {_corporation_statute_name(co)}.

"""

    addr_note = principal_address_note_markdown(co)
    materials_ack_block = board_meeting_materials_acknowledgment_block(co)

    if multi_director_board:
        others = [d["name"] for d in full_board if d["name"] != chair_name]
        others_txt = " and ".join(others) if others else ""
        special_call = (
            f"The Special Meeting of the Board of Directors of {display_company} (the “Corporation”) was called to order at {t_special} on {date} "
            f"by **{chair_name}**, acting as Chair of the Board, with **{others_txt}** also present as a director."
            if others_txt
            else f"The Special Meeting of the Board of Directors of {display_company} (the “Corporation”) was called to order at {t_special} on {date} "
            f"by **{chair_name}**, acting as Chair of the Board, with all directors present."
        )
    else:
        special_call = (
            f"The Special Meeting of the Board of Directors of {display_company} (the “Corporation”) was called to order at {t_special} on {date} "
            f"by {director_name}, acting as Sole Director of the Corporation."
        )

    return f"""
**Minutes of the Special Meeting of the Board of Directors - {year}**
**{display_company}**
(Board of Directors – {_jurisdiction(co)} corporation)

**Meeting Details**
**Date of Meeting:** {date}
**Time of Meeting:** {t_special}
**Location of Meeting:** {place}
{addr_note}**Purpose:** {co.get("special_meeting_purpose", "Pre-annual board review of operations")}

**I. Call to Order:**
{special_call}

{roll_quorum_block}{materials_ack_block}{_special_resolutions_block(
co_name,
eco,
year,
record_date_resolution,
extra_blocks=_board_resolution_prefix_blocks(co_name, co, year, date, "special"),
)}
{reliance_141e_line}

**IV. Adjournment:**
There being no further business to come before the Board, the meeting was adjourned.

{board_meeting_signature_markdown(eco, date, sole_director_name=director_name)}
---"""


def _organizational_bylaws_adoption_resolution(co: dict) -> str:
    lab = str(co.get("organizational_bylaws_exhibit_label") or "").strip()
    desc = str(co.get("organizational_bylaws_document_description") or "").strip()
    if lab and desc:
        if _minutes_assert_exhibits_filed(co):
            core = f"described as **{desc}**, a copy of which is **annexed to these minutes as {lab}**"
        else:
            core = (
                f"described as **{desc}**, **to be designated {lab}** for attachment to these minutes **upon filing**"
            )
    elif lab:
        if _minutes_assert_exhibits_filed(co):
            core = f"**annexed to these minutes as {lab}**"
        else:
            core = f"**to be designated {lab}** for attachment to these minutes **upon filing**"
    else:
        core = "presented to the meeting"
    return f"""**Adoption of Bylaws**  
RESOLVED, that the bylaws {core} are hereby adopted as the bylaws of the Corporation, and the Secretary is authorized and directed to place a copy in the Corporation’s minute book."""


def _organizational_officer_election_resolution(co: dict) -> str:
    roster = co.get("organizational_officers_elected")
    if isinstance(roster, list) and roster:
        lines: list[str] = []
        for item in roster:
            if isinstance(item, dict):
                nm = str(item.get("name") or "").strip()
                t_one = str(item.get("title") or "").strip()
                titles_raw = item.get("titles")
                titles: list[str] = []
                if isinstance(titles_raw, list):
                    titles = [str(t).strip() for t in titles_raw if str(t).strip()]
                elif t_one:
                    titles = [t_one]
                if nm and titles:
                    if len(titles) == 1:
                        lines.append(f"**{nm}** as **{titles[0]}**")
                    else:
                        ox = ", ".join(f"**{t}**" for t in titles[:-1])
                        lines.append(f"**{nm}** as {ox}, and **{titles[-1]}**")
                elif nm:
                    lines.append(f"**{nm}**")
            else:
                s = str(item).strip()
                if s:
                    lines.append(s)
        roster_txt = "; ".join(lines)
        return f"""**Election of Officers**  
RESOLVED, that the following individuals are hereby elected and appointed to serve as officers of the Corporation, to serve at the pleasure of the Board and in accordance with the bylaws: {roster_txt}."""
    return """**Election of Officers**  
RESOLVED, that the individuals currently acting as officers of the Corporation are hereby elected and appointed to continue serving in their respective offices, to serve at the pleasure of the Board and in accordance with the bylaws."""


def generate_quarterly(co_name, year, quarter):
    co = companies[co_name]
    eco = _effective_co_for_board_meeting(co, co_name, year, quarter)
    date = quarterly_meeting_date_str(co, year, quarter)
    place = meeting_place_line(co, date)

    dev_locations = development_centers_line_for_company(co)
    contractor_regions = dev_locations.strip() or "regions covered by the Corporation’s contractor and consultant arrangements"
    t_quarter = scheduled_quarterly_meeting_time(co, year, quarter)

    director_name = "Derek E. Pappas"
    chair_name = str(eco.get("board_meeting_chair_name") or director_name).strip() or director_name
    full_board = _normalized_board_directors(eco)
    multi_director_board = len(full_board) >= 2
    display_company = minutes_display_name(co_name)

    roll_quorum_block = board_roll_quorum_markdown(eco, co_name, date, "quarterly", director_name=director_name)
    reliance_141e_line = board_director_reliance_paragraph(eco, co_name, date, f"quarterly-{quarter}")
    addr_note = principal_address_note_markdown(co)
    materials_ack_block = board_meeting_materials_acknowledgment_block(co)

    custom_review = eco.get("quarterly_business_review_minutes_markdown")
    if isinstance(custom_review, str) and custom_review.strip():
        business_review = custom_review.strip().format(
            year=year, quarter=quarter, dev_locations=dev_locations or "N/A"
        )
    elif multi_director_board:
        business_review = (
            "The directors reviewed quarterly infrastructure stability and confirmed that **software and related intellectual property** "
            f"developed **for the Corporation** during the quarter under its **contractor and consultant arrangements**, including services "
            f"performed by personnel **in {contractor_regions}**, **is properly titled to and is the exclusive property of the Corporation** "
            "under those arrangements and applicable law. The Board acknowledged that **consulting and contracting firms, their personnel, "
            "and their equipment and premises are not owned by the Corporation**."
        )
    else:
        business_review = (
            "The Sole Director reviewed quarterly infrastructure stability and confirmed that **software and related intellectual property** "
            f"developed **for the Corporation** during the quarter under its **contractor and consultant arrangements**, including services "
            f"performed by personnel **in {contractor_regions}**, **is properly titled to and is the exclusive property of the Corporation** "
            "under those arrangements and applicable law. The Sole Director acknowledged that **consulting and contracting firms, their personnel, "
            "and their equipment and premises are not owned by the Corporation**."
        )

    if multi_director_board:
        others = [d["name"] for d in full_board if d["name"] != chair_name]
        others_txt = " and ".join(others) if others else ""
        quarterly_call = (
            f"The Quarterly Governance Meeting of the Board of Directors of {display_company} (the “Corporation”) was called to order at {t_quarter} on {date} "
            f"by **{chair_name}**, acting as Chair of the Board, with **{others_txt}** also present as a director."
            if others_txt
            else f"The Quarterly Governance Meeting of the Board of Directors of {display_company} (the “Corporation”) was called to order at {t_quarter} on {date} "
            f"by **{chair_name}**, acting as Chair of the Board, with all directors present."
        )
    else:
        quarterly_call = (
            f"The Quarterly Governance Meeting of the Board of Directors of {display_company} (the “Corporation”) was called to order at {t_quarter} on {date} "
            f"by {director_name}, acting as Sole Director of the Corporation."
        )

    return f"""
**Minutes of the Quarterly Governance Meeting - {year} {quarter}**
**{display_company}**
(Board of Directors – {_jurisdiction(co)} corporation)

**Meeting Details**
**Date of Meeting:** {date}
**Time of Meeting:** {t_quarter}
**Location of Meeting:** {place}
{addr_note}
**I. Call to Order:**
{quarterly_call}

{roll_quorum_block}{materials_ack_block}**III. Business Review:**
{business_review}

{reliance_141e_line}

{_quarterly_resolutions_block(
        co_name,
        eco,
        co,
        year,
        quarter,
        extra_blocks=_stock_ledger_resolution_blocks_for_meeting(co_name, date),
    )}

**V. Adjournment:**
There being no further business to come before the Board, the meeting was adjourned.

{board_meeting_signature_markdown(eco, date, sole_director_name=director_name)}
---"""


def generate_organizational(co_name: str, year: int) -> str:
    """Post-filing organizational meeting (only when `incorporation_filed_date_iso` is set for `inc_year`)."""
    co = companies[co_name]
    org = organizational_meeting_date_str(co, year)
    if not org:
        return ""
    # Organizational meeting should use the same effective-board logic (first meeting may be sole-director).
    eco = _effective_co_for_board_meeting(co, co_name, year, "org")
    place = meeting_place_line(co, org)
    t_org = ORGANIZATIONAL_MEETING_TIME

    director_name = "Derek E. Pappas"
    chair_name = str(eco.get("board_meeting_chair_name") or director_name).strip() or director_name
    full_board = _normalized_board_directors(eco)
    multi_director_board = len(full_board) >= 2
    display_company = minutes_display_name(co_name)

    roll_quorum_block = board_roll_quorum_markdown(eco, co_name, org, "organizational", director_name=director_name)
    reliance_141e_line = board_director_reliance_paragraph(eco, co_name, org, "organizational")
    addr_note = principal_address_note_markdown(co)

    if multi_director_board:
        others = [d["name"] for d in full_board if d["name"] != chair_name]
        others_txt = " and ".join(others) if others else ""
        call = (
            f"The Organizational Meeting of the Board of Directors of {display_company} (the “Corporation”) was called to order at {t_org} on {org} "
            f"by **{chair_name}**, acting as Chair of the Board, with **{others_txt}** also present as a director."
            if others_txt
            else f"The Organizational Meeting of the Board of Directors of {display_company} (the “Corporation”) was called to order at {t_org} on {org} "
            f"by **{chair_name}**, acting as Chair of the Board, with all directors present."
        )
    else:
        call = (
            f"The Organizational Meeting of the Board of Directors of {display_company} (the “Corporation”) was called to order at {t_org} on {org} "
            f"by {director_name}, acting as Sole Director of the Corporation."
        )

    # Organizational resolutions: formation housekeeping + any one-time domestication/appointment blocks.
    blocks: list[str] = []
    blocks.extend(_domestication_resolution_blocks_if_due(eco, org))
    blocks.extend(_board_appointment_resolution_blocks_if_first_meeting(co_name, co, year, "org"))
    blocks.extend(_stock_ledger_resolution_blocks_for_meeting(co_name, org))
    blocks.extend(
        [
            """**Ratification of Formation Actions**  
RESOLVED, that all actions taken to organize and form the Corporation, including filing the certificate/articles of incorporation and taking related organizational steps, are hereby ratified, confirmed, and approved in all respects.""",
            _organizational_bylaws_adoption_resolution(co),
            _organizational_officer_election_resolution(co),
        ]
    )

    resolutions = _adopted_resolutions_section(eco, "**III. Resolutions:**", blocks)

    return f"""
**Minutes of the Organizational Meeting of the Board of Directors**
**{display_company}**
(Board of Directors – {_jurisdiction(co)} corporation)

**Meeting Details**
**Date of Meeting:** {org}
**Time of Meeting:** {t_org}
**Location of Meeting:** {place}
{addr_note}
**I. Call to Order:**
{call}

{roll_quorum_block}{resolutions}
{reliance_141e_line}

**IV. Adjournment:**
There being no further business to come before the Board, the meeting was adjourned.

{board_meeting_signature_markdown(eco, org, sole_director_name=director_name)}
---""".lstrip()

def generate_quarterly_summary(company_name_year, year, quarter, co_name):
    """Generate a summary of the quarterly meeting for reporting purposes."""
    co = companies[co_name]
    qdate = quarterly_meeting_date_str(co, year, quarter)
    quarterly_docx = meeting_filename(
        co_name,
        qdate,
        "quarterly",
        quarter=quarter,
        ext="docx",
    )
    quarterly_content = generate_quarterly(co_name, year, quarter)
    print(f"Writing Quarterly meeting minutes to {quarterly_docx}")
    write_docx_from_minutes(quarterly_content, quarterly_docx, qdate, co_name)
    print(f"Generating file for {co_name} {year} {quarter}")


def generate_annual_meeting_stockholders(co_name, year):
    """Formal annual meeting minutes for corporations with more than one stockholder (e.g. DATA RECORD SCIENCE, INC.)."""
    co = companies[co_name]
    date = annual_meeting_date_str(co, year)
    record_date = stockholder_annual_record_date_str(co, year)
    special_board_date = board_special_meeting_date_str(co, year)
    place = meeting_place_line(co, date)
    issued = co["shares_issued"].get(year, co["shares_issued"].get(2025))
    display_company = minutes_display_name(co_name)
    t_stock = scheduled_stockholder_meeting_time(co, year)

    chair = "Derek E. Pappas"
    election_standard = co.get("director_election_standard", "plurality")
    election_sentence = (
        "Directors were elected by a plurality of the votes cast by the shares present in person or by proxy and entitled to vote."
        if election_standard == "plurality"
        else "Directors were elected by the requisite vote under the Corporation’s bylaws and applicable law."
    )

    roll_block = format_stockholders_roll_call_block(co)
    vote_tab = annual_stockholder_director_vote_tabulation_markdown(co)
    record_date_source = (
        f"The Chairperson confirmed that **{record_date}** is the **record date** for determining the stockholders entitled to vote at this meeting, "
        f"having been fixed by the Board of Directors pursuant to resolutions adopted at the **Special Meeting of the Board of Directors** held on "
        f"**{special_board_date}** (as reflected in the minutes of such meeting), in accordance with the Corporation’s bylaws and the "
        f"{_corporation_statute_name(co)}. The Chairperson further confirmed that the record date is **not less than** the minimum interval before "
        "this annual meeting required by applicable law and the Corporation’s bylaws."
    )

    addr_note = principal_address_note_markdown(co)
    report_period = _financial_year_words(co)[1]

    return f"""
**Minutes of the Annual Meeting of Stockholders**
**{display_company}**
{_corporation_parenthetical(co)}

**I. Meeting Information**
**Company Name:** {display_company}
**Principal Address:** {co['address']}
{addr_note}**Date:** {date}
**Time:** {t_stock}
**Place:** {place}
**Record Date (stockholders entitled to vote; {_corp_law_section_ref(co, "213")}):** {record_date}
**Type of Meeting:** Annual Meeting of Stockholders

**II. Call to Order and Organization**
The Annual Meeting of Stockholders of {display_company} (the “Corporation”) was called to order commencing at {t_stock} on {date}. Pursuant to the Corporation’s bylaws, {chair}, acting as President of the Corporation, served as Chairperson of the meeting, and the Secretary (or a person designated by the Chairperson) served as Secretary of the meeting.

**III. Roll Call and Quorum**
{record_date_source}

The following stockholders were present in person or by remote participation (as permitted under the Corporation’s bylaws and applicable law) and were entitled to vote at the meeting:

{roll_block}

The Chairperson declared that a quorum of stockholders was present and that the meeting was duly constituted to transact business.
The Chairperson further confirmed that any stockholders participating remotely were able to hear and be heard contemporaneously and that the Corporation had reasonable means to verify that each such person was a stockholder or proxyholder entitled to vote at the meeting.

**Stock Ledger / Voting List; Proxies (bylaws and applicable law)**
The Chairperson confirmed that an **alphabetized list of the names of the stockholders of record** entitled to vote at this meeting (or a certified extract of the stock ledger) as of the record date was produced and made available for inspection by stockholders at the meeting in accordance with the Corporation’s bylaws, and that **proxies** and votes received in accordance with the bylaws were accepted for shares entitled to vote in accordance with such proxies.

{_annual_stockholder_notice_section_iv(co)}

**V. Reports**
The Chairperson presented and summarized the Corporation’s **operational and financial highlights** for the {report_period}. Stockholders had a reasonable opportunity to **ask questions** regarding the Chairperson’s report.

**VI. Election of Directors**
The following resolution was presented and adopted by the stockholders by the requisite vote under the Corporation’s bylaws and applicable law:

**Election of Director**  
RESOLVED, that {chair} is hereby elected as a director of the Corporation, to serve until the next annual meeting of stockholders and until such director’s successor is duly elected and qualified.

{election_sentence}
{vote_tab}
**VII. Shares Outstanding**
The Chairperson noted for the record that the Corporation had {issued} shares of common stock issued and outstanding at a par value of {co['par']} per share as of the date of the meeting.

**VIII. Adjournment**
There being no further business properly brought before the meeting, the meeting was adjourned.

**Executed by (Chairperson of the Meeting):**
{wet_signing_lines_markdown(co, chair, "President and Chairperson of the Meeting", date)}
---
"""


def majority_stockholder_written_consent_ratification_markdown(
    company_name_year: str, year: int, co_name: str
) -> str:
    """Markdown for majority stockholder written consent (ratification; action by written consent)."""
    co = companies[co_name]
    board_date = annual_meeting_date_str(co, year)
    as_of = datetime.strptime(board_date, "%Y-%m-%d").strftime("%B %d, %Y")
    ensuing_period = _financial_year_words(co)[1]
    mechanics = f"""**Written Consent Mechanics**
This Written Consent is intended to be delivered to the Corporation and to become effective in accordance with the {_corporation_statute_name(co)} and the Corporation’s bylaws, including any timing requirements applicable to the delivery of consents bearing dated signatures. The Corporation is authorized and directed to file this Written Consent with the minutes of the proceedings of the stockholders of the Corporation and to give any prompt notice required by applicable law, the certificate of incorporation, and the bylaws."""
    return f"""
**{co_name}.**
**Written Consent of Majority Stockholders**
**Action by Written Consent of Stockholders**
({_corp_law_section_ref(co, "228")})

The undersigned, being the stockholders of {co_name}, a {_jurisdiction(co)} corporation (the "Corporation"), holding not less than the minimum number of votes that would be necessary to authorize the following action at a meeting at which all shares entitled to vote thereon were present and voted, hereby adopt the following resolutions by written consent pursuant to the {_corporation_statute_name(co)}, effective as of the date set forth below.

**Ratification of Annual Board Meeting**
RESOLVED, that all actions taken and resolutions adopted by the Board of Directors of the Corporation at the Annual Meeting of the Board of Directors held on {as_of} (or as otherwise recorded in the minutes of such meeting), including the approval of financial statements, the budget for the ensuing {ensuing_period}, officer actions, and banking authorizations, are hereby ratified, confirmed, and approved in all respects.

**Notice to Non-Consenting Stockholders**
RESOLVED, that the Corporation is authorized and directed to provide prompt notice of the taking of the foregoing corporate action by written consent, to the extent required by applicable law, the Corporation’s certificate of incorporation, and the Corporation’s bylaws.

{mechanics}

**Effective Date**
This Written Consent shall be effective as of {as_of}, and shall be filed with the minutes of the proceedings of the stockholders of the Corporation.

**Executed by (stockholders):**


---
"""


def generate_majority_stockholder_written_consent_ratification(company_name_year: str, year: int, co_name: str):
    """Majority stockholder written consent ratifying same-year annual board actions, for multi-stockholder corporations."""
    co = companies[co_name]
    board_date = annual_meeting_date_str(co, year)
    content = majority_stockholder_written_consent_ratification_markdown(company_name_year, year, co_name)
    output = meeting_filename(
        co_name,
        board_date,
        "majority_stockholders_written_consent_ratification_of_annual_board_actions",
        ext="docx",
    )
    print(f"Writing Majority Stockholders Written Consent (ratification) to {output}")
    write_docx_from_minutes(content, output, board_date, co_name)

def sanitize_company_name(name):
    # remove the comma from company name for file naming
    safe_name = name.replace(", ", "_").replace(" ", "_")
    
    # remove a dot at the end
    if safe_name.endswith("."):
        safe_name = safe_name[:-1]
    
    # remove any remaining dots
    

    # remove "Inc" or "inc" from the end
    if safe_name.lower().endswith("inc"):
        safe_name = safe_name[:-3]
    
    # lowercase the company name for file naming
    safe_name = safe_name.lower()
    
    # remove trailing underscores
    safe_name = safe_name.rstrip('_')
    
    # replace . with underscore
    safe_name = safe_name.replace('.', '_')
    
    return safe_name


def meetings_dir(company_root: str) -> str:
    """``<company_root>/meetings/`` — single folder per company for all meeting-minute (and related) ``.docx`` files."""
    return os.path.join(os.path.abspath(company_root), "meetings")


def _slugify_filename_part(s: str) -> str:
    s = s.strip().lower()
    s = re.sub(r"[^a-z0-9]+", "_", s)
    s = re.sub(r"_+", "_", s)
    return s.strip("_")


def meeting_filename(
    co_name: str,
    meeting_date_iso: str,
    meeting_name: str,
    *,
    quarter: str | None = None,
    ext: str = "docx",
) -> str:
    """
    Canonical file naming format for sortable output:

      <company name (no .inc.)>_<year>_<month>_<day>_(<q_{1,2,3,4}>)*_<meeting name>.<ext>

    This ensures simple lexicographic sorting matches chronological order within
    a company folder.
    """
    dt = datetime.strptime(meeting_date_iso, "%Y-%m-%d")
    safe = sanitize_company_name(co_name)
    y = f"{dt.year:04d}"
    m = f"{dt.month:02d}"
    d = f"{dt.day:02d}"

    q_part = ""
    if quarter:
        q = quarter.strip().upper()
        m_q = re.fullmatch(r"Q([1-4])", q)
        if m_q:
            q_part = f"_q_{m_q.group(1)}"
        else:
            q_part = "_" + _slugify_filename_part(quarter)

    meeting_part = _slugify_filename_part(meeting_name)
    return f"{safe}_{y}_{m}_{d}{q_part}_{meeting_part}.{ext.lstrip('.')}"



def generate_annual(company_name_year: str, co_name: str, year: int):
    co = companies[co_name]
    mdate = annual_meeting_date_str(co, year)
    agm_docx = meeting_filename(co_name, mdate, "agm", ext="docx")
    agm_content = generate_agm(co_name, year)
    print(f"Writing AGM minutes to {agm_docx}")
    write_docx_from_minutes(agm_content, agm_docx, mdate, co_name)

    _summary, detail_items = accomplishments_for_year(co_name, year)
    exhibit = co.get("agm_president_report_operating_exhibit_label")
    if detail_items and not exhibit:
        exhibit = "Exhibit B"
    addendum_path = meeting_filename(co_name, mdate, "agm_operating_addendum", ext="docx")
    if detail_items and exhibit:
        _generate_agm_operating_addendum_docx(company_name_year, co_name, year, exhibit, detail_items)
    elif os.path.isfile(addendum_path):
        os.remove(addendum_path)


def generate_special_meeting(company_name_year: str, co_name: str, year: int):
    """Generate special meeting minutes."""
    co = companies[co_name]
    mdate = board_special_meeting_date_str(co, year)
    special_docx = meeting_filename(co_name, mdate, "yearly_special_meeting", ext="docx")
    special_content = generate_special(co_name, year)
    print(f"Writing Special meeting minutes to {special_docx}")
    write_docx_from_minutes(special_content, special_docx, mdate, co_name)


def sole_stockholder_written_consent_markdown(co_name: str, year: int) -> str:
    """Markdown for sole stockholder written consent (action by written consent)."""
    co = companies[co_name]
    display_name = minutes_display_name(co_name)
    date = annual_meeting_date_str(co, year)
    as_of = datetime.strptime(date, "%Y-%m-%d").strftime("%B %d, %Y")
    shareholder_term = "Sole Stockholder"
    voting_shares = co.get(
        "voting_shares_description",
        "all of the outstanding shares of the Corporation entitled to vote on the following matters",
    )
    bylaws_ack = co.get("stockholder_consent_bylaws_acknowledgment")
    bylaws_ack_block = f"\n{bylaws_ack}\n" if bylaws_ack else ""
    mech_suffix = co.get("stockholder_consent_bylaws_mechanics_suffix")
    mech_tail = f" {mech_suffix}" if mech_suffix else ""
    mechanics = f"""**Written Consent Mechanics**
This Written Consent is intended to be delivered to the Corporation and to become effective in accordance with the {_corporation_statute_name(co)} and the Corporation’s bylaws, including any timing requirements applicable to the delivery of consents bearing dated signatures.{mech_tail} The Corporation is authorized and directed to file this Written Consent with the minutes of the proceedings of the stockholders of the Corporation and to give any prompt notice required by applicable law, the certificate of incorporation, and the bylaws (to the extent applicable)."""
    consent_heading = f"({_corp_law_section_ref(co, '228')}; annual meeting of stockholders)"
    if _jurisdiction(co) == "WY":
        consent_intro = (
            f"The undersigned, being the {shareholder_term.lower()} of {display_name}, a {_jurisdiction(co)} corporation (the \"Corporation\"), "
            "and constituting **all** shareholders entitled to vote on the following matters, hereby adopts the following resolutions by **written consent without a meeting** "
            f"pursuant to **{_corp_law_section_ref(co, '228')}** and the {_corporation_statute_name(co)}, to the same extent as if the resolutions were approved at a duly held "
            "shareholders’ meeting by the unanimous vote of all shares entitled to vote thereon. The undersigned directs that this consent be delivered to the Corporation for "
            "inclusion in the minutes or filing with the corporate records as required by law. The undersigned acknowledges that, to the best of the undersigned’s knowledge, "
            f"the articles of incorporation do not prohibit shareholder action by written consent in the manner set forth herein.{bylaws_ack_block}"
        )
    else:
        consent_intro = (
            f"The undersigned, being the {shareholder_term.lower()} of {display_name}, a {_jurisdiction(co)} corporation (the \"Corporation\"), and holding {voting_shares}, "
            f"hereby adopts the following resolutions by written consent of the stockholders pursuant to the {_corporation_statute_name(co)}, effective without a meeting to the "
            "same extent as if adopted at a duly held meeting. The undersigned acknowledges that, to the best of the undersigned’s knowledge, the Corporation’s certificate of "
            "incorporation does not prohibit stockholder action by written consent as contemplated hereby, including action by the holders of outstanding shares of capital stock "
            "having not less than the minimum number of votes that would be necessary to authorize or take such action at a meeting at which all shares entitled to vote thereon "
            f"were present and voted.{bylaws_ack_block}"
        )
    return f"""
**{display_name}**
**Written Consent of {shareholder_term}**
**Action by Written Consent of Stockholders**
{consent_heading}

{consent_intro}
**Approval of Board Actions**
RESOLVED, that all actions taken and resolutions adopted by the Board of Directors of the Corporation at the Annual Meeting of the Board of Directors held on {as_of} (or as otherwise recorded in the minutes of such meeting), including but not limited to the approval of financial statements, budgets, officer actions, and banking authorizations, are hereby ratified, confirmed, and approved in all respects.

**Annual Meeting of Stockholders**
RESOLVED, that the matters set forth in this Written Consent, including ratification of the Board’s actions taken at the annual board meeting referenced above, constitute or supplement, as applicable, the business addressed for purposes of the annual meeting of stockholders for the year {year} to the extent permitted by the {_corporation_statute_name(co)}, the certificate of incorporation, and the bylaws of the Corporation, and the undersigned waives any requirement to convene a separate annual meeting of stockholders for such year solely to duplicate the matters resolved herein.

{mechanics}

**Effective Date**
This Written Consent shall be effective as of {as_of}, and shall be filed with the minutes of the proceedings of the stockholders of the Corporation.

**Stockholder Certification**
{shareholder_term}:
Derek E. Pappas

**Executed by:**
{wet_signing_lines_markdown(co, "Derek E. Pappas", shareholder_term, date)}
---"""


def generate_written_consent(company_name_year: str, year: int, company_name: str):
    """Generate sole stockholder written consent under applicable corporate law (.docx)."""
    co = companies[company_name]
    date = annual_meeting_date_str(co, year)
    content = sole_stockholder_written_consent_markdown(company_name, year)
    output = meeting_filename(company_name, date, "written_consent_in_lieu_of_annual_meeting", ext="docx")
    print(f"Writing Stockholder Written Consent to {output}")
    write_docx_from_minutes(content, output, date, company_name)


def generate_stockholder_side(company_name_year: str, year: int, co_name: str):
    co = companies[co_name]
    kind = co.get("stockholder_meeting", "written_consent")
    if kind == "annual_meeting_stockholders":
        mdate = annual_meeting_date_str(co, year)
        out = meeting_filename(co_name, mdate, "annual_meeting_of_stockholders", ext="docx")
        print(f"Writing Annual Meeting of Stockholders minutes to {out}")
        write_docx_from_minutes(generate_annual_meeting_stockholders(co_name, year), out, mdate, co_name)
        generate_stockholder_waiver_of_notice_annual_meeting(company_name_year, year, co_name)
        generate_notice_of_annual_stockholder_meeting(company_name_year, year, co_name)
        generate_majority_stockholder_written_consent_ratification(company_name_year, year, co_name)
    else:
        generate_written_consent(company_name_year, year, co_name)
    

import argparse


def _company_years_for_calendar(co: dict, years: tuple[int, ...]) -> list[int]:
    start = max(co.get("minutes_start_year", co.get("inc_year", min(years))), min(years))
    return [y for y in years if y >= start]


def write_company_calendars(output_dir: str = "calendars", years: tuple[int, ...] = (2022, 2023, 2024, 2025, 2026)) -> int:
    """
    Produce one .txt file per company with meetings grouped by date.

    Master (all companies): ``{output_dir}/unified_calendar.txt``.

    Overlap audit: same calendar date and identical time string across *different*
    companies is written to ``{output_dir}/conflicts.txt``. Returns the number of
    such conflict slots (0 = clean). This does not parse clock intervals; only
    exact (date, time) equality is checked.

    Format:
    - Company name at top
    - For each date: date on its own line
    - Then indented lines: "<Company> - <Meeting Name> - <Time>"
    - One blank line between dates
    """
    os.makedirs(output_dir, exist_ok=True)

    quarterly_month = {"Q1": 3, "Q2": 6, "Q3": 9, "Q4": 12}

    unified_entries: list[tuple[str, str, str, str]] = []
    # tuple: (date_iso, time_str, company_name, meeting_label)

    for co_name, co in companies.items():
        entries_by_date: dict[str, list[str]] = {}

        for year in _company_years_for_calendar(co, years):
            annual_date = annual_meeting_date_str(co, year)
            special_date = board_special_meeting_date_str(co, year)
            t_stock = scheduled_stockholder_meeting_time(co, year)
            t_board = scheduled_board_agm_time(co, year)
            t_special = scheduled_special_meeting_time(co, year)

            if co.get("stockholder_meeting") == "annual_meeting_stockholders":
                entries_by_date.setdefault(annual_date, []).append(f"{co_name} - Annual Meeting of Stockholders - {t_stock}")
                entries_by_date.setdefault(annual_date, []).append(f"{co_name} - Annual Meeting of the Board of Directors - {t_board}")
                entries_by_date.setdefault(annual_date, []).append(f"{co_name} - Majority Stockholders Written Consent (Ratification) - {t_stock}")

                unified_entries.append((annual_date, t_stock, co_name, "Annual Meeting of Stockholders"))
                unified_entries.append((annual_date, t_board, co_name, "Annual Meeting of the Board of Directors"))
                unified_entries.append((annual_date, t_stock, co_name, "Majority Stockholders Written Consent (Ratification)"))
            else:
                entries_by_date.setdefault(annual_date, []).append(f"{co_name} - Annual Meeting of the Board of Directors - {t_board}")
                entries_by_date.setdefault(annual_date, []).append(f"{co_name} - Stockholder Written Consent - {t_stock}")

                unified_entries.append((annual_date, t_board, co_name, "Annual Meeting of the Board of Directors"))
                unified_entries.append((annual_date, t_stock, co_name, "Stockholder Written Consent"))

            entries_by_date.setdefault(special_date, []).append(
                f"{co_name} - Yearly Special Meeting (Board) - {t_special}"
            )
            unified_entries.append((special_date, t_special, co_name, "Yearly Special Meeting (Board)"))

            org_date = organizational_meeting_date_str(co, year)
            if org_date:
                entries_by_date.setdefault(org_date, []).append(
                    f"{co_name} - Organizational Meeting (Board) - {ORGANIZATIONAL_MEETING_TIME}"
                )
                unified_entries.append((org_date, ORGANIZATIONAL_MEETING_TIME, co_name, "Organizational Meeting (Board)"))

            for q in ("Q1", "Q2", "Q3", "Q4"):
                if not _quarterly_exists_for_year(co, year, q):
                    continue
                q_date = quarterly_meeting_date_str(co, year, q)
                t_q = scheduled_quarterly_meeting_time(co, year, q)
                entries_by_date.setdefault(q_date, []).append(
                    f"{co_name} - Quarterly Meeting (Board) {q} - {t_q}"
                )
                unified_entries.append((q_date, t_q, co_name, f"Quarterly Meeting (Board) {q}"))

        lines: list[str] = [co_name, ""]
        for d in sorted(entries_by_date.keys()):
            lines.append(d)
            for item in entries_by_date[d]:
                lines.append(f"  {item}")
            lines.append("")  # blank line between dates

        out_path = os.path.join(output_dir, f"{sanitize_company_name(co_name)}_calendar.txt")
        with open(out_path, "w", encoding="utf-8") as f:
            f.write("\n".join(lines).rstrip() + "\n")

    # Unified calendar + conflict check across companies (same date + same time, different companies)
    by_date: dict[str, list[tuple[str, str, str]]] = {}
    by_slot: dict[tuple[str, str], list[tuple[str, str]]] = {}
    # by_slot[(date,time)] => [(company, meeting_label), ...]

    for d, t, co_name, label in unified_entries:
        by_date.setdefault(d, []).append((t, co_name, label))
        by_slot.setdefault((d, t), []).append((co_name, label))

    conflicts: list[tuple[str, str, list[tuple[str, str]]]] = []
    for (d, t), items in sorted(by_slot.items()):
        companies_in_slot = {co for co, _ in items}
        if len(companies_in_slot) > 1:
            conflicts.append((d, t, items))

    unified_lines: list[str] = ["Unified Meeting Calendar", ""]
    if conflicts:
        unified_lines.append("Conflict Summary (same date + time across companies)")
        for d, t, items in conflicts:
            unified_lines.append(f"{d} {t}")
            for co_name, label in items:
                unified_lines.append(f"  {co_name} - {label} - {t}")
            unified_lines.append("")
        unified_lines.append("---")
        unified_lines.append("")
    else:
        unified_lines.append("Conflict Summary: none detected (across companies)")
        unified_lines.append("")
        unified_lines.append("---")
        unified_lines.append("")

    for d in sorted(by_date.keys()):
        unified_lines.append(d)
        for t, co_name, label in sorted(by_date[d], key=lambda x: (x[0], x[1], x[2])):
            unified_lines.append(f"  {co_name} - {label} - {t}")
        unified_lines.append("")

    unified_path = os.path.join(output_dir, "unified_calendar.txt")
    with open(unified_path, "w", encoding="utf-8") as f:
        f.write("\n".join(unified_lines).rstrip() + "\n")

    conflicts_path = os.path.join(output_dir, "conflicts.txt")
    with open(conflicts_path, "w", encoding="utf-8") as f:
        if not conflicts:
            f.write("No conflicts detected (same date + time across different companies).\n")
        else:
            f.write("Conflicts detected (same date + time across different companies):\n\n")
            for d, t, items in conflicts:
                f.write(f"{d} {t}\n")
                for co_name, label in items:
                    f.write(f"  {co_name} - {label} - {t}\n")
                f.write("\n")

    return len(conflicts)


def print_schedule(years=(2022, 2023, 2024, 2025, 2026)):
    """Print the computed meeting schedule without generating .docx files."""
    print("Schedule (all times local):")
    for year in years:
        print(f"\nYear {year}")
        for co_name, co in companies.items():
            if year < co.get("minutes_start_year", co.get("inc_year", year)):
                continue
            board_date = annual_meeting_date_str(co, year)
            special_date = board_special_meeting_date_str(co, year)
            t_stock = scheduled_stockholder_meeting_time(co, year)
            t_board = scheduled_board_agm_time(co, year)
            t_special = scheduled_special_meeting_time(co, year)
            if co.get("stockholder_meeting") == "annual_meeting_stockholders":
                stock_date = annual_meeting_date_str(co, year)
                rd = stockholder_annual_record_date_str(co, year)
                print(
                    f"- {co_name}: Board (special; record-date cycle) {special_date} {t_special} "
                    f"(record date {rd}); Stockholders {stock_date} {t_stock}; "
                    f"Board (annual) {board_date} {t_board}"
                )
            else:
                print(
                    f"- {co_name}: Board (special) {special_date} {t_special}; "
                    f"Board (annual) {board_date} {t_board}; Written consent dated {board_date}"
                )


# Between meeting bodies in the per-company compilation: one empty paragraph (“hard” break) in the .docx output.
MEETING_BOOK_SEPARATOR = "\n\n"

# Sentinel line only in `*_all_meetings_book` markdown; `write_docx_from_minutes(..., minute_book_page_breaks=True)`
# and `_write_minute_book_pdf` turn it into a hard page break so each included document prints as a standalone piece.
MEETING_BOOK_PAGE_BREAK_MARKER = "<<<MINUTE_BOOK_PAGE_BREAK>>>"


def _minute_book_line_to_paragraph_xml(line: str) -> str:
    """Escape for ReportLab Paragraph XML; convert **bold** to <b>."""
    s = line.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
    return re.sub(r"\*\*(.+?)\*\*", r"<b>\1</b>", s)


def _write_minute_book_pdf(markdown: str, pdf_path: str) -> None:
    """Letter-size PDF, continuous page numbers in footer (single volume)."""
    try:
        from reportlab.lib import colors
        from reportlab.lib.enums import TA_LEFT
        from reportlab.lib.pagesizes import letter
        from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
        from reportlab.lib.units import inch
        from reportlab.platypus import HRFlowable, PageBreak, Paragraph, SimpleDocTemplate, Spacer
    except ModuleNotFoundError as e:
        raise RuntimeError(
            "minute book PDF requires reportlab (see pyproject.toml). "
            "Install project deps: poetry install"
        ) from e

    styles = getSampleStyleSheet()
    body_style = ParagraphStyle(
        "MinuteBookBody",
        parent=styles["Normal"],
        fontName="Helvetica",
        fontSize=10,
        leading=12,
        spaceAfter=4,
        alignment=TA_LEFT,
    )

    story: list = []
    for raw in markdown.splitlines():
        line = raw.rstrip()
        if line.strip() == MEETING_BOOK_PAGE_BREAK_MARKER:
            story.append(PageBreak())
        elif line.strip() == SIGNATURE_BLOCK_MARKER:
            # Marker used only to help .docx keep signature blocks together.
            continue
        elif not line.strip():
            story.append(Spacer(1, 10))
        elif line.strip() == "---":
            story.append(Spacer(1, 6))
            story.append(
                HRFlowable(
                    width=letter[0] - 1.5 * inch,
                    thickness=0.5,
                    color=colors.HexColor("#bbbbbb"),
                    hAlign="CENTER",
                )
            )
            story.append(Spacer(1, 6))
        else:
            story.append(Paragraph(_minute_book_line_to_paragraph_xml(line), body_style))

    def _page_footer(canvas, doc):
        canvas.saveState()
        canvas.setFont("Helvetica", 9)
        canvas.drawRightString(letter[0] - 0.75 * inch, 0.55 * inch, f"Page {canvas.getPageNumber()}")
        canvas.restoreState()

    doc = SimpleDocTemplate(
        pdf_path,
        pagesize=letter,
        rightMargin=0.75 * inch,
        leftMargin=0.75 * inch,
        topMargin=0.75 * inch,
        bottomMargin=0.85 * inch,
        onFirstPage=_page_footer,
        onLaterPages=_page_footer,
    )
    doc.build(story)


def _write_docx_as_simple_pdf(docx_path: str, pdf_path: str) -> None:
    """Letter-size PDF from a generated .docx (paragraph + table text; for distribution copies)."""
    try:
        from reportlab.lib import colors
        from reportlab.lib.enums import TA_LEFT
        from reportlab.lib.pagesizes import letter
        from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
        from reportlab.lib.units import inch
        from reportlab.platypus import HRFlowable, Paragraph, SimpleDocTemplate, Spacer
    except ModuleNotFoundError as e:
        raise RuntimeError(
            "example PDF export requires reportlab (see pyproject.toml). "
            "Install project deps: poetry install"
        ) from e

    docx_doc = Document(docx_path)
    lines: list[str] = []
    for p in docx_doc.paragraphs:
        lines.append(p.text)
    for table in docx_doc.tables:
        for row in table.rows:
            lines.append("\t".join(cell.text.strip() for cell in row.cells))

    styles = getSampleStyleSheet()
    body_style = ParagraphStyle(
        "DocxExportBody",
        parent=styles["Normal"],
        fontName="Helvetica",
        fontSize=10,
        leading=12,
        spaceAfter=4,
        alignment=TA_LEFT,
    )

    story: list = []
    for raw in lines:
        line = raw.rstrip()
        if not line.strip():
            story.append(Spacer(1, 10))
        elif line.strip() == "---":
            story.append(Spacer(1, 6))
            story.append(
                HRFlowable(
                    width=letter[0] - 1.5 * inch,
                    thickness=0.5,
                    color=colors.HexColor("#bbbbbb"),
                    hAlign="CENTER",
                )
            )
            story.append(Spacer(1, 6))
        else:
            story.append(Paragraph(_minute_book_line_to_paragraph_xml(line), body_style))

    def _page_footer(canvas, doc):
        canvas.saveState()
        canvas.setFont("Helvetica", 9)
        canvas.drawRightString(letter[0] - 0.75 * inch, 0.55 * inch, f"Page {canvas.getPageNumber()}")
        canvas.restoreState()

    os.makedirs(os.path.dirname(pdf_path) or ".", exist_ok=True)
    rl_doc = SimpleDocTemplate(
        pdf_path,
        pagesize=letter,
        rightMargin=0.75 * inch,
        leftMargin=0.75 * inch,
        topMargin=0.75 * inch,
        bottomMargin=0.85 * inch,
        onFirstPage=_page_footer,
        onLaterPages=_page_footer,
    )
    rl_doc.build(story)


def _markdown_chunks_for_calendar_year(company_name_year: str, co_name: str, year: int) -> list[str]:
    """Chronological meeting order for one calendar year in the compiled minute book.

    Board meetings are ordered by date/time (quarterlies first, then annual-cycle special/AGM).
    Stockholder-side instruments (if any) are inserted on the annual meeting date ahead of the board AGM.
    """
    co = companies[co_name]
    chunks: list[str] = []

    # Board meetings in chronological order (includes special + AGM + Q1–Q4).
    rows = _board_meeting_rows_for_year(co, year)

    # Precompute annual-date stockholder pack (if used) so it can be inserted before the board AGM.
    stockholder_pack: list[str] = []
    if co.get("stockholder_meeting") == "annual_meeting_stockholders":
        stockholder_pack = [
            generate_annual_meeting_stockholders(co_name, year).rstrip(),
            stockholder_waiver_of_notice_annual_meeting_markdown(company_name_year, year, co_name).rstrip(),
            notice_of_annual_stockholder_meeting_markdown(company_name_year, year, co_name).rstrip(),
            majority_stockholder_written_consent_ratification_markdown(company_name_year, year, co_name).rstrip(),
        ]

    inserted_stockholders = False
    inserted_written_consent = False

    for d_iso, title, _t, _place in rows:
        low = title.lower()
        if "organizational meeting of the board of directors" in low:
            org = generate_organizational(co_name, year).rstrip()
            if org:
                chunks.append(org)
            continue
        if "quarterly governance meeting" in low:
            # Example: "Quarterly Governance Meeting – 2024 Q2"
            q = title.strip().split()[-1]
            chunks.append(generate_quarterly(co_name, year, q).rstrip())
            continue

        if "special meeting of the board of directors" in low:
            chunks.append(generate_special(co_name, year).rstrip())
            continue

        if "annual meeting of the board of directors" in low:
            # Annual meeting date: insert stockholder meeting pack first (if applicable).
            if stockholder_pack and not inserted_stockholders:
                chunks.extend(stockholder_pack)
                inserted_stockholders = True

            agm = generate_agm(co_name, year).rstrip()
            chunks.append(agm)

            # Operating addendum (if any) should follow the AGM in the compiled book.
            _summary, detail_items = accomplishments_for_year(co_name, year)
            exhibit = co.get("agm_president_report_operating_exhibit_label")
            if detail_items and not exhibit:
                exhibit = "Exhibit B"
            if detail_items and exhibit:
                chunks.append(agm_operating_addendum_markdown(co_name, year, exhibit, detail_items).rstrip())

            # Written-consent companies: attach the sole stockholder consent after the AGM on the annual date.
            if co.get("stockholder_meeting") != "annual_meeting_stockholders" and not inserted_written_consent:
                chunks.append(sole_stockholder_written_consent_markdown(co_name, year).rstrip())
                inserted_written_consent = True
            continue

        # Fallback: preserve unknown row titles by not failing the compilation.
        chunks.append(f"**Unrecognized meeting row:** {title} ({d_iso})\n---")

    # In case the annual meeting row is absent (should not happen), still include stockholder pack and/or AGM.
    if stockholder_pack and not inserted_stockholders:
        chunks.extend(stockholder_pack)
    if co.get("stockholder_meeting") != "annual_meeting_stockholders" and not inserted_written_consent:
        chunks.append(sole_stockholder_written_consent_markdown(co_name, year).rstrip())

    # Waiver of notice (board) is dated at the annual meeting date and is conventionally filed after the year’s minutes.
    chunks.append(board_waiver_of_notice_markdown(company_name_year, year, co_name).rstrip())

    return chunks


def _minute_book_compilation_header_markdown(co_name: str, co: dict, applicable: list[int]) -> str:
    """Cover text for compiled `*_all_meetings_book`; per-company template optional."""
    first_year = applicable[0]
    last_year = applicable[-1]
    display_company = minutes_display_name(co_name)
    tpl = co.get("minute_book_compilation_preamble_markdown")
    if isinstance(tpl, str) and tpl.strip():
        return tpl.format(
            display_company=display_company,
            first_year=first_year,
            last_year=last_year,
        ).strip()
    return f"""**Minute book compilation — all meetings**
**{display_company}**
(single document: meetings generated for calendar years {first_year} through {last_year}.)
(The Corporation may maintain other minutes or records for periods outside this span; this volume is limited to the years listed.)

Exhibits referenced in these minutes (including Exhibit A / Exhibit B) are part of the Corporation’s corporate records. Signed counterparts, annexes, or labeled exhibits may be bound with this compilation or cross-filed as separate instruments in the minute book.

---
""".strip()


def generate_company_all_meetings_book(
    safe_company_name: str,
    co_name: str,
    years: tuple[int, ...],
    company_root: str,
) -> None:
    """Compiled minute book per company: .docx (editable) + .pdf (distribution).

    ``company_root`` is ``{output_root}/<safe_company_name>/`` (the consolidated book is written **here**, not under
    ``meetings/``). Individual meeting minutes live under ``<company_root>/meetings/`` (single folder per company).

    Also writes ``<safe>/cap_tables/…`` and ``<safe>/stock_ledgers/…``, and appends cap / ledger markdown as the final
    sections of the compiled book (and PDF).
    """
    co = companies[co_name]
    start_year = co.get("minutes_start_year", co.get("inc_year", min(years)))
    applicable = [y for y in years if y >= start_year]
    if not applicable:
        return
    parts: list[str] = [_minute_book_compilation_header_markdown(co_name, co, applicable)]
    for y in applicable:
        cny = f"{safe_company_name}_{y}"
        chunks = _markdown_chunks_for_calendar_year(cny, co_name, y)
        for i, ch in enumerate(chunks):
            if i == 0:
                # Put a hard page break between year-sections (documents only, not an extra break at the start).
                if y != applicable[0]:
                    parts.append(MEETING_BOOK_PAGE_BREAK_MARKER)
                parts.append(f"**Calendar year {y}**\n\n{ch}")
            else:
                parts.append(f"{MEETING_BOOK_PAGE_BREAK_MARKER}\n{ch}")
    mdate = annual_meeting_date_str(co, applicable[-1])
    company_root = os.path.abspath(company_root)
    cap_tables_dir = os.path.join(company_root, "cap_tables")
    stock_ledgers_dir = os.path.join(company_root, "stock_ledgers")
    legacy_books = os.path.join(company_root, "books")
    if os.path.isdir(legacy_books):
        for legacy_name in (
            f"{safe_company_name}_cap_table.docx",
            f"{safe_company_name}_stock_ledger.docx",
            f"{safe_company_name}_cap_table_carta_pulley.csv",
            f"{safe_company_name}_all_meetings_book.docx",
            f"{safe_company_name}_all_meetings_book.pdf",
        ):
            lp = os.path.join(legacy_books, legacy_name)
            if os.path.isfile(lp):
                try:
                    os.remove(lp)
                except OSError:
                    pass
    _write_cap_table_and_stock_ledger_docx(
        safe_company_name, co_name, co, cap_tables_dir, stock_ledgers_dir, mdate
    )
    parts.append(MEETING_BOOK_PAGE_BREAK_MARKER)
    parts.append(cap_table_document_markdown(co_name, co))
    parts.append(MEETING_BOOK_PAGE_BREAK_MARKER)
    parts.append(stock_ledger_document_markdown(co_name, co))
    book = MEETING_BOOK_SEPARATOR.join(p for p in parts if p)
    os.makedirs(company_root, exist_ok=True)
    out_docx = os.path.join(company_root, f"{safe_company_name}_all_meetings_book.docx")
    out_pdf = os.path.join(company_root, f"{safe_company_name}_all_meetings_book.pdf")
    print(f"Writing compiled minute book to {out_docx}")
    write_docx_from_minutes(
        book, out_docx, mdate, co_name, minute_book_page_breaks=True
    )
    print(f"Writing compiled minute book PDF to {out_pdf}")
    _write_minute_book_pdf(book, out_pdf)


def generate_master_all_companies_book(
    output_root: str,
    years: tuple[int, ...] = (2022, 2023, 2024, 2025, 2026),
) -> tuple[str, str]:
    """One combined minute book spanning all companies.

    Output:
    - `{output_root}/books/all_companies_all_meetings_book.docx`
    - `{output_root}/books/all_companies_all_meetings_book.pdf`
    """
    start_cwd = os.getcwd()
    root_dir = os.path.join(start_cwd, output_root)
    books_dir = os.path.join(root_dir, "books")
    os.makedirs(books_dir, exist_ok=True)

    parts: list[str] = [
        "**Master minute book — all companies**\n"
        f"(single document: all generated meetings for calendar years {min(years)} through {max(years)}.)\n\n"
        "---"
    ]

    for idx, (co_name, co) in enumerate(companies.items()):
        safe = sanitize_company_name(co_name)
        start_year = co.get("minutes_start_year", co.get("inc_year", min(years)))
        applicable = [y for y in years if y >= start_year]
        if not applicable:
            continue
        if idx != 0:
            parts.append(MEETING_BOOK_PAGE_BREAK_MARKER)
        parts.append(f"**Company: {minutes_display_name(co_name)}**\n{_corporation_parenthetical(co)}\n\n---")
        for y in applicable:
            cny = f"{safe}_{y}"
            for i, ch in enumerate(_markdown_chunks_for_calendar_year(cny, co_name, y)):
                parts.append(MEETING_BOOK_PAGE_BREAK_MARKER)
                if i == 0:
                    parts.append(f"**Calendar year {y}**\n\n{ch}")
                else:
                    parts.append(ch)
        parts.append(MEETING_BOOK_PAGE_BREAK_MARKER)
        parts.append(cap_table_document_markdown(co_name, co))
        parts.append(MEETING_BOOK_PAGE_BREAK_MARKER)
        parts.append(stock_ledger_document_markdown(co_name, co))

    book = MEETING_BOOK_SEPARATOR.join(p for p in parts if p)
    out_docx = os.path.join(books_dir, "all_companies_all_meetings_book.docx")
    out_pdf = os.path.join(books_dir, "all_companies_all_meetings_book.pdf")

    # Use the latest applicable annual meeting date among all companies for utime randomization.
    latest_meeting_date = None
    for co_name, co in companies.items():
        start_year = co.get("minutes_start_year", co.get("inc_year", min(years)))
        applicable = [y for y in years if y >= start_year]
        if not applicable:
            continue
        d = annual_meeting_date_str(co, applicable[-1])
        if latest_meeting_date is None or d > latest_meeting_date:
            latest_meeting_date = d
    latest_meeting_date = latest_meeting_date or f"{max(years)}-12-31"

    print(f"Writing master compiled minute book to {out_docx}")
    write_docx_from_minutes(book, out_docx, latest_meeting_date, list(companies.keys())[0], minute_book_page_breaks=True)
    print(f"Writing master compiled minute book PDF to {out_pdf}")
    _write_minute_book_pdf(book, out_pdf)
    return out_docx, out_pdf


def write_samples_directory(
    output_root: str,
    years: tuple[int, ...] = (2022, 2023, 2024, 2025, 2026),
) -> str:
    """Create ``generated/samples/`` (single folder) with PDF samples of representative generated docs.

    All companies write into the same directory (filenames include ``<safe>_`` prefixes, so names do not collide).
    Cross-company master PDF/CSV copies land in that folder too. Removes any legacy ``generated/examples/`` or
    ``generated/<safe>/examples/`` trees from older generator versions.

    Picks the latest year available for each company, then emits one PDF per category:
    - agm
    - yearly_special_meeting
    - written_consent_in_lieu_of_annual_meeting OR annual_meeting_of_stockholders (+ waiver/notice/ratification where present)
    - waiver_of_notice_board_meetings
    - one quarterly (Q1)
    - cap_table, stock_ledger (from `generated/<safe>/cap_tables/*.docx` and `generated/<safe>/stock_ledgers/*.docx` when present)
    - cap_table_carta_pulley.csv (copied next to those PDFs when present)
    - all_meetings_book (compiled): copies the pre-built `generated/<safe>/<safe>_all_meetings_book.pdf` when present

    Standalone .docx files are converted with ReportLab (same letter layout as compiled book PDFs).
    """
    start_cwd = os.getcwd()
    root_dir = os.path.join(start_cwd, output_root)
    known_safes = {sanitize_company_name(n) for n in companies}
    legacy_global = os.path.join(root_dir, "examples")
    if os.path.isdir(legacy_global):
        shutil.rmtree(legacy_global, ignore_errors=True)
    for safe in known_safes:
        legacy_co = os.path.join(root_dir, safe, "examples")
        if os.path.isdir(legacy_co):
            shutil.rmtree(legacy_co, ignore_errors=True)

    samples_root = os.path.join(root_dir, "samples")
    if os.path.isdir(samples_root):
        shutil.rmtree(samples_root, ignore_errors=True)
    os.makedirs(samples_root, exist_ok=True)

    books_dir = os.path.join(root_dir, "books")
    os.makedirs(books_dir, exist_ok=True)

    def _emit_pdf_from_docx(src_docx: str, dst_pdf: str) -> None:
        os.makedirs(os.path.dirname(dst_pdf), exist_ok=True)
        _write_docx_as_simple_pdf(src_docx, dst_pdf)

    for co_name, co in companies.items():
        safe = sanitize_company_name(co_name)
        co_dir = os.path.join(root_dir, safe)
        if not os.path.isdir(co_dir):
            continue

        start_year = co.get("minutes_start_year", co.get("inc_year", min(years)))
        applicable = [y for y in years if y >= start_year]
        if not applicable:
            continue
        y = applicable[-1]
        annual_date = annual_meeting_date_str(co, y)
        special_date = board_special_meeting_date_str(co, y)
        q1_date = quarterly_meeting_date_str(co, y, "Q1")

        out_dir = samples_root

        # Core docs (always generated) — live under ``meetings/`` (flat per company).
        my_meetings = os.path.join(co_dir, "meetings")
        for src in (
            os.path.join(my_meetings, meeting_filename(co_name, annual_date, "agm", ext="docx")),
            os.path.join(my_meetings, meeting_filename(co_name, special_date, "yearly_special_meeting", ext="docx")),
            os.path.join(my_meetings, meeting_filename(co_name, annual_date, "waiver_of_notice_board_meetings", ext="docx")),
            os.path.join(
                my_meetings,
                meeting_filename(
                    co_name,
                    q1_date,
                    "quarterly",
                    quarter="Q1",
                    ext="docx",
                ),
            ),
        ):
            if os.path.isfile(src):
                stem = os.path.splitext(os.path.basename(src))[0]
                _emit_pdf_from_docx(src, os.path.join(out_dir, f"{stem}.pdf"))

        # Stockholder side varies
        stockholder_kind = co.get("stockholder_meeting", "written_consent")
        if stockholder_kind == "annual_meeting_stockholders":
            for src in (
                os.path.join(my_meetings, meeting_filename(co_name, annual_date, "annual_meeting_of_stockholders", ext="docx")),
                os.path.join(
                    my_meetings,
                    meeting_filename(co_name, annual_date, "waiver_of_notice_annual_stockholder_meeting", ext="docx"),
                ),
                os.path.join(
                    my_meetings,
                    meeting_filename(co_name, annual_date, "notice_of_annual_stockholder_meeting", ext="docx"),
                ),
                os.path.join(
                    my_meetings,
                    meeting_filename(
                        co_name,
                        annual_date,
                        "majority_stockholders_written_consent_ratification_of_annual_board_actions",
                        ext="docx",
                    ),
                ),
            ):
                if os.path.isfile(src):
                    stem = os.path.splitext(os.path.basename(src))[0]
                    _emit_pdf_from_docx(src, os.path.join(out_dir, f"{stem}.pdf"))
        else:
            src = os.path.join(
                my_meetings, meeting_filename(co_name, annual_date, "written_consent_in_lieu_of_annual_meeting", ext="docx")
            )
            if os.path.isfile(src):
                stem = os.path.splitext(os.path.basename(src))[0]
                _emit_pdf_from_docx(src, os.path.join(out_dir, f"{stem}.pdf"))

        cap_tables_dir = os.path.join(co_dir, "cap_tables")
        stock_ledgers_dir = os.path.join(co_dir, "stock_ledgers")
        cap_docx = os.path.join(cap_tables_dir, f"{safe}_cap_table.docx")
        if os.path.isfile(cap_docx):
            _emit_pdf_from_docx(cap_docx, os.path.join(out_dir, f"{safe}_cap_table.pdf"))
        led_docx = os.path.join(stock_ledgers_dir, f"{safe}_stock_ledger.docx")
        if os.path.isfile(led_docx):
            _emit_pdf_from_docx(led_docx, os.path.join(out_dir, f"{safe}_stock_ledger.pdf"))
        csv_src = os.path.join(cap_tables_dir, f"{safe}_cap_table_carta_pulley.csv")
        if os.path.isfile(csv_src):
            shutil.copy2(csv_src, os.path.join(out_dir, f"{safe}_cap_table_carta_pulley.csv"))

        # Compiled book: distribution PDF at company root.
        book_pdf = os.path.join(co_dir, f"{safe}_all_meetings_book.pdf")
        if os.path.isfile(book_pdf):
            dst = os.path.join(out_dir, f"{safe}_all_meetings_book.pdf")
            shutil.copy2(book_pdf, dst)

    master_pdf = os.path.join(books_dir, "all_companies_all_meetings_book.pdf")
    if os.path.isfile(master_pdf):
        shutil.copy2(master_pdf, os.path.join(samples_root, "all_companies_all_meetings_book.pdf"))
    master_csv = os.path.join(books_dir, "all_companies_cap_table_carta_pulley.csv")
    if os.path.isfile(master_csv):
        shutil.copy2(master_csv, os.path.join(samples_root, "all_companies_cap_table_carta_pulley.csv"))

    return samples_root


def write_loki_agm_accomplishments_exhibits_pdf_bundle(output_root: str) -> str:
    """Write `generated/loki_sports_enterprises_agm_minutes_with_accomplishment_exhibits_pdf/` with PDFs only.

    For each calendar year that has accomplishments detail in `audit_reports/all_corp_accomplishments_2021-2025.json`
    (LOKI), copies the AGM minutes and the AGM operating addendum (Exhibit B — detailed accomplishments) from the
    generated company folder into one distribution folder as `.pdf` files (no filename collisions across years).
    """
    co_name = "Loki Sports Enterprises, Inc."
    safe = sanitize_company_name(co_name)
    root_dir = os.path.abspath(os.path.join(os.getcwd(), output_root))
    co_dir = os.path.join(root_dir, safe)
    out_dir = os.path.join(root_dir, f"{safe}_agm_minutes_with_accomplishment_exhibits_pdf")
    os.makedirs(out_dir, exist_ok=True)
    for name in list(os.listdir(out_dir)):
        p = os.path.join(out_dir, name)
        if name.endswith(".pdf") and os.path.isfile(p):
            os.remove(p)

    years_with_detail: list[int] = []
    for y in range(2020, 2032):
        _s, details = accomplishments_for_year(co_name, y)
        if details:
            years_with_detail.append(y)

    if not years_with_detail:
        print(f"No accomplishments detail found for {co_name}; skipping Loki AGM exhibit PDF bundle.")
        return out_dir

    for y in years_with_detail:
        annual_date = annual_meeting_date_str(companies[co_name], y)
        y_meet = os.path.join(co_dir, "meetings")
        agm_docx = os.path.join(y_meet, meeting_filename(co_name, annual_date, "agm", ext="docx"))
        add_docx = os.path.join(y_meet, meeting_filename(co_name, annual_date, "agm_operating_addendum", ext="docx"))
        if os.path.isfile(agm_docx):
            out_pdf = meeting_filename(co_name, annual_date, "agm", ext="pdf")
            _write_docx_as_simple_pdf(agm_docx, os.path.join(out_dir, out_pdf))
        else:
            print(f"skip (missing): {agm_docx}")
        if os.path.isfile(add_docx):
            out_pdf = meeting_filename(co_name, annual_date, "agm_operating_addendum", ext="pdf")
            _write_docx_as_simple_pdf(add_docx, os.path.join(out_dir, out_pdf))
        else:
            print(f"skip (missing addendum for {y}; run full generation after accomplishments are present): {add_docx}")

    print(f"Wrote Loki AGM + accomplishment exhibit PDFs to {out_dir}")
    return out_dir


def generate_all(output_root: str, years=(2022, 2023, 2024, 2025, 2026)):
    start_cwd = os.getcwd()
    print(f"Current working directory: {start_cwd}")
    root_dir = os.path.join(start_cwd, output_root)
    os.makedirs(root_dir, exist_ok=True)
    books_dir = os.path.join(root_dir, "books")
    os.makedirs(books_dir, exist_ok=True)

    try:
        for name in companies.keys():
            co = companies[name]
            _warn_if_non_de_company_has_delaware_snippets(name, co)
            print(f"Company {name} Current working directory: {os.getcwd()}")

            os.chdir(root_dir)
            safe_company_name = sanitize_company_name(name)

            company_dir = f"{safe_company_name}"
            company_root = os.path.join(root_dir, company_dir)
            os.makedirs(company_root, exist_ok=True)
            meetings_root = meetings_dir(company_root)
            os.makedirs(meetings_root, exist_ok=True)
            # Drop legacy ``meetings/<YYYY>/`` subfolders from older generator layouts.
            if os.path.isdir(meetings_root):
                for sub in list(os.listdir(meetings_root)):
                    sp = os.path.join(meetings_root, sub)
                    if os.path.isdir(sp) and sub.isdigit():
                        shutil.rmtree(sp, ignore_errors=True)
            # Stray meeting ``.docx`` accidentally written at company root (remove; keep compiled book).
            for fn in list(os.listdir(company_root)):
                if not fn.endswith(".docx"):
                    continue
                if not fn.startswith(f"{safe_company_name}_"):
                    continue
                if "all_meetings_book" in fn:
                    continue
                try:
                    os.remove(os.path.join(company_root, fn))
                except OSError:
                    pass

            start_year = co.get(
                "minutes_start_year", co.get("inc_year", min(years))
            )

            legacy_books = os.path.join(company_root, "books")
            if os.path.isdir(legacy_books):
                for fn in list(os.listdir(legacy_books)):
                    if not fn.endswith(".docx"):
                        continue
                    if fn.startswith(f"{safe_company_name}_"):
                        try:
                            os.remove(os.path.join(legacy_books, fn))
                        except OSError:
                            pass

            os.chdir(meetings_root)
            year_prefix = re.compile(rf"^{re.escape(safe_company_name)}_(\d{{4}})_")
            for fn in list(os.listdir(".")):
                if not fn.endswith(".docx"):
                    continue
                m = year_prefix.match(fn)
                if m and int(m.group(1)) < start_year:
                    try:
                        os.remove(os.path.join(meetings_root, fn))
                    except OSError:
                        pass

            for year in years:
                if year < co.get("minutes_start_year", co.get("inc_year", year)):
                    continue

                company_name_year = f"{safe_company_name}_{year}"

                # Organizational meeting (post-filing; only for inc_year when a filed date is known).
                org_date = organizational_meeting_date_str(co, year)
                if org_date:
                    content = generate_organizational(name, year)
                    if content.strip():
                        org_docx = meeting_filename(name, org_date, "organizational", ext="docx")
                        print(f"Writing Organizational meeting minutes to {org_docx}")
                        write_docx_from_minutes(content, org_docx, org_date, name)

                generate_annual(company_name_year, name, year)
                generate_special_meeting(company_name_year, name, year)
                generate_stockholder_side(company_name_year, year, name)
                generate_board_waiver_of_notice(company_name_year, year, name)

                for quarter in ["Q1", "Q2", "Q3", "Q4"]:
                    if not _quarterly_exists_for_year(co, year, quarter):
                        continue
                    generate_quarterly_summary(company_name_year, year, quarter, name)

            generate_company_all_meetings_book(safe_company_name, name, years, company_root)
    finally:
        os.chdir(start_cwd)

    write_all_companies_cap_table_carta_pulley_csv(root_dir)
    write_standalone_board_resolution_documents(root_dir)


def main():
    parser = argparse.ArgumentParser(description="Generate corporate meeting minutes (.docx).")
    parser.add_argument(
        "--print-schedule",
        action="store_true",
        help="Print computed annual meeting dates/times without generating documents.",
    )
    parser.add_argument(
        "--write-calendars",
        action="store_true",
        help="Write per-company meeting calendar .txt files (does not generate .docx).",
    )
    parser.add_argument(
        "--strict-calendars",
        action="store_true",
        help=(
            "With --write-calendars: exit with status 1 if unified_calendar audit finds any "
            "same date+time slot shared by more than one company (see calendar-output-dir/conflicts.txt)."
        ),
    )
    parser.add_argument(
        "--output-root",
        default="generated",
        help="Output folder (relative to current working directory). Default: generated",
    )
    parser.add_argument(
        "--calendar-output-dir",
        default="calendars",
        help="Calendar output folder (relative to current working directory). Default: calendars",
    )
    parser.add_argument(
        "--schedule-seed",
        default=None,
        help=(
            "Sets CORPORATE_MINUTES_SCHEDULE_SEED for reproducible date/time jitter "
            "(integer or string; unset = nominal schedule only unless the env var is already set)."
        ),
    )
    parser.add_argument(
        "--extract-audit-text",
        action="store_true",
        help=(
            "After generation, run scripts/extract_audit_text.py so audit_text/*.txt mirrors "
            "generated/**/*.docx (books, cap_tables, stock_ledgers, and compiled books)."
        ),
    )
    parser.add_argument(
        "--write-samples",
        action="store_true",
        help=(
            "Create ``generated/samples/`` with PDF samples (and copied CSVs) of representative generated docs "
            "for all registry companies in one folder, plus cross-company master PDF/CSV copies."
        ),
    )
    parser.add_argument(
        "--write-loki-agm-exhibits-pdf",
        action="store_true",
        help=(
            "Create `generated/loki_sports_enterprises_agm_minutes_with_accomplishment_exhibits_pdf/` with PDFs of "
            "each year’s AGM minutes plus the operating addendum (Exhibit B / accomplishments), for years that have "
            "accomplishments data in audit_reports/all_corp_accomplishments_2021-2025.json."
        ),
    )
    parser.add_argument(
        "--write-master-book",
        action="store_true",
        help=(
            "Write one compiled minute book spanning all companies to <output-root>/books/ "
            "(per-company minutes under <output-root>/<safe>/meetings/; standalone equity/domestication "
            "resolutions there too; cap tables and stock ledgers under <output-root>/<safe>/cap_tables/ and …/stock_ledgers/; "
            "compiled book at <output-root>/<safe>/<safe>_all_meetings_book.*)."
        ),
    )
    args = parser.parse_args()

    if args.schedule_seed is not None:
        os.environ["CORPORATE_MINUTES_SCHEDULE_SEED"] = str(args.schedule_seed)

    if args.print_schedule:
        print_schedule()
        return
    if args.write_calendars:
        conflict_slots = write_company_calendars(output_dir=args.calendar_output_dir)
        if args.strict_calendars and conflict_slots:
            print(
                f"Calendar audit failed: {conflict_slots} conflict slot(s). "
                f"See {args.calendar_output_dir}/conflicts.txt",
                file=sys.stderr,
            )
            sys.exit(1)
        return

    generate_all(output_root=args.output_root)

    if args.write_master_book:
        generate_master_all_companies_book(output_root=args.output_root)

    if args.write_samples:
        path = write_samples_directory(output_root=args.output_root)
        print(f"Wrote samples directory to {path}")

    if args.write_loki_agm_exhibits_pdf:
        write_loki_agm_accomplishments_exhibits_pdf_bundle(output_root=args.output_root)

    if args.extract_audit_text:
        repo_root = os.path.dirname(os.path.abspath(__file__))
        extract_script = os.path.join(repo_root, "scripts", "extract_audit_text.py")
        subprocess.run([sys.executable, extract_script], cwd=repo_root, check=True)


if __name__ == "__main__":
    main()
